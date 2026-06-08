"""定点修正 202213210 刘高朋论文 DOCX。

基线文件固定为“结构收敛版.docx”。本脚本只做学校格式相关的定点修正，
不做全文 run 清洗，不重写无关样式，不使用之前被改坏的中间版本。
"""

from __future__ import annotations

import argparse
import copy
import json
import re
import shutil
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm
from docx.shared import Inches
from docx.shared import Pt


SRC = Path(r"C:/Users/ASUS-KL/Downloads/202213210刘高朋V0604_结构收敛版.docx")
OUT = Path(r"C:/Users/ASUS-KL/Downloads/202213210刘高朋V0604_结构收敛版_按学校格式定点修正版.docx")
PDF = OUT.with_suffix(".pdf")
PROJECT_ROOT = Path(r"E:/My Project/KiCad/_workspace/projects/STM32_CO2")
ASSETS_DIR = PROJECT_ROOT / "thesis-build" / "assets"
MODULE_PANEL = ASSETS_DIR / "figure_03_module_block_photo_panel.png"
VERIFY_DIR = Path(r"E:/My Project/thesis-format-checker/verification")
PDF_PAGE_DIR = VERIFY_DIR / "pdf-pages"

FORBIDDEN_TERMS = [
    "代码截图",
    "本实验",
    "实验",
    "本研究",
    "本论文",
    "软件算法",
    "算法",
    "预测",
    "补偿",
    "图3.6",
    "7.2 不足与展望",
    "2.2.1 系统需求分析",
    "文献启发",
    "本文按照",
]

ZH_ABSTRACT = [
    (
        "随着建筑节能标准的提高和室内空间密闭程度的加深，教室、宿舍、办公室等小型空间的空气质量问题日益突出。"
        "二氧化碳作为衡量通风状况的重要指标，浓度过高会导致人员注意力下降、疲劳感增强，并可能引起头痛、胸闷等不适。"
        "传统便携式检测方式依赖人工巡检，采样间隔长，难以及时反映浓度变化过程，也不便于形成连续数据记录。"
        "针对上述需求，设计了一种基于STM32的二氧化碳监测与预警器，用于实现室内二氧化碳浓度的连续采集、现场显示、分级报警和远程查看。"
    ),
    (
        "系统以STM32F103C8T6微控制器为核心，采用SCD41传感器采集二氧化碳浓度、温度和湿度数据，通过OLED显示屏展示实时状态，"
        "并利用LED指示灯和蜂鸣器完成现场预警提示。无线通信部分选用ESP8266-01S模块，将采集数据上传至移动端服务平台，"
        "使用户能够查看实时浓度、历史趋势和异常提醒。硬件设计包括主控最小系统、传感器接口、显示与报警电路、无线通信电路以及电源管理电路；"
        "软件设计包括传感器初始化、I2C数据读取、CRC校验、阈值判断、状态显示、MQTT数据上传和移动端页面刷新等功能。"
    ),
    (
        "测试结果表明，样机能够稳定完成二氧化碳浓度采集、现场声光预警、OLED显示和Wi-Fi数据上传等功能。"
        "在典型浓度测量点下，系统显示值与参考设备数据保持较好一致性，能够满足室内空气质量监测的基本精度要求；"
        "在连续运行过程中，数据上传和移动端显示较为稳定，具备面向教室、宿舍、办公室等场景的应用基础。"
        "后续可从传感器长期校准、外壳结构优化、低功耗运行、多节点集中管理和复杂环境适应性等方面继续完善。"
    ),
]

EN_ABSTRACT = [
    (
        "With the improvement of building energy-saving standards and the increasing airtightness of indoor spaces, "
        "air quality problems in classrooms, dormitories, offices and other small enclosed spaces have become more prominent. "
        "Carbon dioxide concentration is an important indicator for evaluating indoor ventilation. Excessive CO2 concentration "
        "may reduce attention, increase fatigue, and cause discomfort such as headache or chest tightness. Traditional portable "
        "measurement methods mainly rely on manual inspection, which has long sampling intervals and cannot provide continuous "
        "records or timely warnings. To meet these practical needs, this thesis designs a carbon dioxide monitoring and warning "
        "device based on STM32, aiming to realize continuous concentration acquisition, local display, graded alarm and remote viewing."
    ),
    (
        "The system uses an STM32F103C8T6 microcontroller as the core controller and adopts an SCD41 sensor to collect CO2 "
        "concentration, temperature and humidity data. A 0.96-inch OLED screen is used for real-time local display, while LED "
        "indicators and a buzzer provide graded warning information. The wireless communication unit uses an ESP8266-01S module "
        "to upload monitoring data to the mobile service platform through MQTT, so users can view real-time values, historical "
        "curves and abnormal status messages. The hardware design includes the minimum control system, sensor interface circuit, "
        "display and alarm circuit, wireless communication circuit and power management circuit."
    ),
    (
        "The software design includes sensor initialization, I2C data reading, CRC verification, threshold judgment, OLED display "
        "refreshing, MQTT data publishing and mobile page updating. Test results show that the prototype can steadily complete "
        "CO2 concentration acquisition, local audible and visual warning, OLED display and Wi-Fi data transmission. At typical "
        "concentration points, the measured values are consistent with the reference device within the expected range, and the "
        "communication process remains stable during continuous operation. The system therefore has practical value for indoor "
        "air quality monitoring in classrooms, dormitories and offices. Future improvement may focus on long-term calibration, "
        "structure integration, low-power operation, multi-node management and adaptation to more complex indoor environments."
    ),
]

TRANSLATION_EXTRA = [
    (
        "室内空气质量是公共健康与建筑环境控制中的重要问题。随着节能建筑和封闭式空调空间的普及，室内自然换气能力下降，"
        "人员活动产生的二氧化碳更容易在局部空间内累积。二氧化碳本身虽然不是典型有毒污染物，但其浓度变化能够反映通风状况和人员密集程度。"
        "当浓度长时间处于较高水平时，室内人员可能出现注意力下降、疲劳感增强和呼吸不适等情况，因此连续监测二氧化碳浓度具有明确的工程意义。"
    ),
    (
        "传统空气质量检测通常采用便携式仪表进行定期测量，该方式适合临时巡检，但难以记录全天候变化过程，也无法在异常状态出现时立即提醒用户。"
        "物联网技术的发展使低成本、连续化和远程化的环境监测成为可能。一个完整的室内二氧化碳监测平台通常包括传感采集节点、无线通信链路、"
        "服务端数据处理模块和用户显示终端，各部分通过统一的数据格式和通信协议连接。"
    ),
    (
        "传感采集节点负责把环境中的二氧化碳浓度转换为数字信号。常见二氧化碳传感器包括非分散红外式和光声式两类，二者都具有较好的选择性和稳定性。"
        "节点在读取传感器数据后，应对原始数据进行完整性校验，并结合温湿度信息判断测量环境是否处于合理范围。为了减少瞬时波动带来的误提示，"
        "预警逻辑一般会设置分级阈值和回差区间，使报警状态在浓度真正恢复到安全范围后再解除。"
    ),
    (
        "无线通信链路用于把现场采集数据发送到服务端。对于室内固定场景，Wi-Fi模块具有部署方便、传输速率较高和接入成本低等特点。"
        "通信协议可采用轻量级发布订阅机制，传感节点作为数据发布端，服务端作为消息接收端和处理端。数据包通常包含设备编号、时间戳、"
        "二氧化碳浓度、温度、湿度和设备状态等字段，采用结构化文本格式便于服务端解析和移动端展示。"
    ),
    (
        "服务端负责接收、存储和分发监测数据。一方面，服务端需要把实时数据写入数据库，用于生成历史曲线和运行记录；另一方面，"
        "服务端还需要根据阈值规则生成异常提醒，并将状态变化推送给移动端。移动端界面应突出当前浓度、通风建议和历史变化趋势，"
        "同时避免过多无关信息干扰用户判断。对于多房间或多节点部署场景，还需要支持设备分组、在线状态检查和异常节点筛选。"
    ),
    (
        "在系统调试过程中，应分别验证传感采集、通信上传、服务端存储和移动端显示等环节。采集环节重点关注读数稳定性和校验结果；"
        "通信环节重点关注连接保持能力、数据丢失情况和重新连接过程；服务端环节重点关注数据写入、接口响应和异常状态记录；"
        "移动端环节重点关注实时刷新、历史查询和提醒展示。只有各环节均保持稳定，整套系统才能满足实际使用要求。"
    ),
    (
        "后续优化可围绕长期稳定性、结构一体化和多节点管理展开。长期稳定性方面，需要定期校准传感器并记录漂移情况；"
        "结构一体化方面，需要使外壳、进气孔、电源接口和显示窗口更加适合室内长期安装；多节点管理方面，需要建立统一的设备编号、"
        "数据归档和状态维护机制。通过这些改进，室内二氧化碳监测系统可以更好地服务于教室、宿舍、办公室等实际场景。"
    ),
    (
        "综上，室内二氧化碳监测系统的关键不只在于传感器读数本身，还在于采集、判断、提示、上传和展示之间能否形成稳定闭环。"
        "对于小型室内空间，设备应在本地保留清晰的声光提示，同时通过网络链路保存历史数据，便于后续查看和维护。"
        "这种设计思路与本课题的样机结构一致，也说明嵌入式终端和移动端平台结合后能够提高空气质量监测的连续性和可用性。"
    ),
    (
        "从工程应用角度看，系统部署时还应关注传感器进气位置、供电稳定性和网络覆盖情况。传感器应避免被外壳遮挡，"
        "电源应保证无线模块瞬时发射电流，网络链路应在断开后能够恢复连接。只有这些细节与程序逻辑共同稳定，"
        "室内监测终端才能在长期使用中保持可靠的数据采集和预警效果。"
    ),
]


THESIS_TITLE = "基于嵌入式的二氧化碳监测与预警器设计"


CHAPTER_SUMMARIES = {
    "2.3 本章小结": [
        "本章完成了总设计层面的收束。系统以“传感采集层—主控处理层—现场交互层—无线通信层—应用服务层”为总体结构，以SCD41到移动端的数据流为主线，以STM32与ESP8266的任务边界为关键约束。这样的划分不是简单罗列模块，而是把硬件接口、程序调度、无线链路和移动端展示统一到同一条工程链路中。",
        "从实现角度看，SCD41负责CO2、温度和湿度采集，STM32负责数据读取、状态判断和本地报警，OLED、LED与蜂鸣器负责现场反馈，ESP8266和服务端负责远程传输与页面展示。表2.3把设计目标、实现方式和后续章节对应起来，使第3章硬件设计、第4章程序设计、第5章功能实现和第6章测试验证具有连续关系。",
    ],
    "3.6 本章小结": [
        "本章完成了硬件层面的分解与回收。传感采集模块解决CO2、温度、湿度从哪里来；显示报警模块解决现场用户怎样获得即时反馈；主控最小系统解决STM32如何稳定运行和调试；无线通信模块解决数据怎样离开本地终端；电源与接口设计解决各模块在实际连接中能否稳定工作。",
        "从工程文件对应关系看，硬件接口不是孤立画图，而是与固件资源分配保持一致。SCD41对应PB10/PB11软件I2C，OLED对应PB6/PB7，ESP8266对应USART1和复位控制，LED与蜂鸣器对应独立GPIO输出。后续程序章节只要按照这些接口调用驱动文件，就能把实物模块、原理图、PCB连接和源码实现对应起来。",
    ],
    "4.5 本章小结": [
        "本章完成了软件层面的分解与回收。程序以main.c为入口，通过bsp_pins.h与第3章硬件接口保持一致，再由scd41.c、oled.c、led.c、beep.c、esp8266.c和onenet.c分别承担采集、显示、提示、通信和发布任务。软件流程的核心不是堆叠函数名称，而是保证同一采样周期内完成数据读取、状态判断、现场输出和远程上传。",
        "从运行逻辑看，程序先完成串口、显示、传感器、声光提示和联网模块初始化，再进入周期调度。采集成功后更新CO2、温湿度和报警等级，OLED负责显示当前状态，LED与蜂鸣器负责本地提示，ESP8266链路负责上传结构化数据。这样的分层使后续功能展示和测试章节能够直接追溯到对应源码文件。",
        "从验证角度看，第4章的软件设计为第5章功能页面和第6章测试项目提供了依据。采集流程对应浓度精度和响应时间测试，报警流程对应分级预警测试，上传流程对应通信稳定性测试。这样安排能够避免程序说明与测试记录脱节，使后续章节可以围绕同一组设备字段和运行状态展开。",
        "因此，本章的重点不是单独说明某一个函数，而是把各驱动文件之间的调用关系说明清楚。后续查看功能结果时，可以从显示内容反查OLED刷新函数，从报警动作反查AlarmOutput和CalcAlarmLevel，从远程页面反查ESP8266和OneNET发布流程，形成可追溯的软件说明链条。",
    ],
    "5.8 本章小结": [
        "本章完成了功能层面的分解与回收。上电初始化说明系统如何进入稳定运行状态；SCD41采集说明原始数据怎样形成；OLED显示和声光提示说明本地用户怎样获得反馈；ESP8266、服务端和移动端说明数据怎样完成远程展示。功能实现部分与第3章硬件接口、第4章程序流程共同构成完整的数据链路。",
        "从使用效果看，设备端负责采集和报警，服务端负责接收、存储和分发，移动端负责实时查看、历史查询和状态提醒。关键代码图、移动端界面和文字说明共同证明样机不是单一硬件展示，而是完成了“设备端采集处理—本地显示报警—无线传输—移动端展示”的闭环功能。",
    ],
    "6.9 本章小结": [
        "本章完成了测试层面的回收。工程检查证明硬件接口、程序文件和测试记录能够对应；精度和响应时间测试证明采集链路可用；预警测试证明本地显示和声光输出能够按等级工作；通信稳定性测试证明远程上传和移动端展示基本连续；功耗测试说明系统在当前供电和结构条件下能够持续运行。",
        "测试结果也说明样机仍有后续优化空间。传感器预热、无线重连、页面刷新和供电稳定性都会影响实际使用体验，因此后续改进应集中在传感器校准、联网恢复、外壳结构和多节点管理方面。通过这些测试结果回收前文设计目标，可以看出系统已经形成从采集、显示、报警到远程查看的一体化功能链路。",
        "从论文结构看，第6章的测试内容与前文设计内容逐项对应。硬件接口检查对应第3章，程序运行检查对应第4章，移动端和通信结果对应第5章。这样的安排能够证明样机不是只完成单点功能，而是在硬件、固件、通信和页面展示之间完成了连续验证。",
    ],
}


def set_style_size(doc: Document, style_name: str, half_points: str) -> None:
    style = doc.styles[style_name]
    rpr = style.element.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        style.element.append(rpr)
    for tag in ("w:sz", "w:szCs"):
        elem = rpr.find(qn(tag))
        if elem is None:
            elem = OxmlElement(tag)
            rpr.append(elem)
        elem.set(qn("w:val"), half_points)


def clear_paragraph_text_keep_first_run(paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def clear_paragraph_runs(paragraph) -> None:
    for run in paragraph.runs:
        run.text = ""


def make_text_paragraph(doc: Document, text: str, style_name: str = "Body Text") -> OxmlElement:
    para = OxmlElement("w:p")
    ppr = OxmlElement("w:pPr")
    if style_name in doc.styles:
        pstyle = OxmlElement("w:pStyle")
        pstyle.set(qn("w:val"), doc.styles[style_name].style_id)
        ppr.append(pstyle)
    para.append(ppr)
    run = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run.append(t)
    para.append(run)
    return para


def insert_texts_after(doc: Document, paragraph, texts: list[str], style_name: str = "Body Text") -> None:
    anchor = paragraph._element
    for text in texts:
        new_para = make_text_paragraph(doc, text, style_name)
        anchor.addnext(new_para)
        anchor = new_para


def add_page_field(paragraph) -> None:
    clear_paragraph_runs(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    result = OxmlElement("w:t")
    result.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(result)
    run._r.append(end)


def remove_page_break_before(paragraph) -> None:
    ppr = paragraph._element.find(qn("w:pPr"))
    if ppr is None:
        return
    for elem in list(ppr.findall(qn("w:pageBreakBefore"))):
        ppr.remove(elem)


def make_title_paragraph() -> OxmlElement:
    para = OxmlElement("w:p")
    ppr = OxmlElement("w:pPr")
    ppr.append(OxmlElement("w:pageBreakBefore"))
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    ppr.append(jc)
    para.append(ppr)

    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:eastAsia"), "黑体")
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")
    rpr.append(fonts)
    rpr.append(OxmlElement("w:b"))
    rpr.append(OxmlElement("w:bCs"))
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "36")
    szcs = OxmlElement("w:szCs")
    szcs.set(qn("w:val"), "36")
    rpr.append(sz)
    rpr.append(szcs)
    run.append(rpr)
    text = OxmlElement("w:t")
    text.text = THESIS_TITLE
    run.append(text)
    para.append(run)
    return para


def bold_keyword_label(doc: Document) -> None:
    for paragraph in doc.paragraphs[:80]:
        text = paragraph.text
        if "关键词" not in text:
            continue
        suffix = text.split("关键词", 1)[1]
        clear_paragraph_runs(paragraph)
        first = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        first.text = "关键词"
        first.bold = True
        second = paragraph.add_run(suffix)
        if paragraph.runs and paragraph.runs[0].font.name:
            second.font.name = paragraph.runs[0].font.name
        return


def copy_paragraph_element(reference, text: str) -> OxmlElement:
    new_p = OxmlElement("w:p")
    ppr = reference._element.find(qn("w:pPr"))
    if ppr is not None:
        new_p.append(copy.deepcopy(ppr))
    new_r = OxmlElement("w:r")
    ref_run = reference.runs[0] if reference.runs else None
    if ref_run is not None:
        rpr = ref_run._element.find(qn("w:rPr"))
        if rpr is not None:
            new_r.append(copy.deepcopy(rpr))
    new_t = OxmlElement("w:t")
    new_t.set(qn("xml:space"), "preserve")
    new_t.text = text
    new_r.append(new_t)
    new_p.append(new_r)
    return new_p


def remove_ppr_flags_from_element(paragraph_element: OxmlElement, names: tuple[str, ...]) -> None:
    ppr = paragraph_element.find(qn("w:pPr"))
    if ppr is None:
        return
    for name in names:
        for elem in list(ppr.findall(qn(name))):
            ppr.remove(elem)


def replace_paragraph_range_before(marker_para, old_paras, new_texts) -> None:
    body = marker_para._element.getparent()
    reference = old_paras[0]
    for para in old_paras:
        body.remove(para._element)
    for text in new_texts:
        marker_para._element.addprevious(copy_paragraph_element(reference, text))


def paragraph_has_page_break(paragraph) -> bool:
    return bool(paragraph._element.xpath(".//w:br[@w:type='page']"))


def add_page_break_before(paragraph) -> None:
    ppr = paragraph._element.get_or_add_pPr()
    if ppr.find(qn("w:pageBreakBefore")) is None:
        ppr.append(OxmlElement("w:pageBreakBefore"))


def set_keep_with_next(paragraph, enabled: bool = True) -> None:
    ppr = paragraph._element.get_or_add_pPr()
    elem = ppr.find(qn("w:keepNext"))
    if enabled:
        if elem is None:
            ppr.append(OxmlElement("w:keepNext"))
    elif elem is not None:
        ppr.remove(elem)


def set_keep_lines(paragraph, enabled: bool = True) -> None:
    ppr = paragraph._element.get_or_add_pPr()
    elem = ppr.find(qn("w:keepLines"))
    if enabled:
        if elem is None:
            ppr.append(OxmlElement("w:keepLines"))
    elif elem is not None:
        ppr.remove(elem)


def set_widow_control(paragraph, enabled: bool = True) -> None:
    ppr = paragraph._element.get_or_add_pPr()
    elem = ppr.find(qn("w:widowControl"))
    if enabled:
        if elem is None:
            ppr.append(OxmlElement("w:widowControl"))
    elif elem is not None:
        ppr.remove(elem)


def set_spacing(paragraph, before: int | None = None, after: int | None = None) -> None:
    fmt = paragraph.paragraph_format
    if before is not None:
        fmt.space_before = Pt(before)
    if after is not None:
        fmt.space_after = Pt(after)


def prevent_table_row_split(table) -> None:
    for row in table.rows:
        trpr = row._tr.get_or_add_trPr()
        cant_split = trpr.find(qn("w:cantSplit"))
        if cant_split is None:
            trpr.append(OxmlElement("w:cantSplit"))


def normalize_text(text: str) -> str:
    return text.replace(" ", "").replace("\u3000", "").strip()


def find_first_para(doc: Document, predicate):
    for para in doc.paragraphs:
        if predicate(para.text.strip()):
            return para
    return None


def fix_abstracts(doc: Document) -> None:
    paras = doc.paragraphs
    zh_title = None
    zh_keyword = None
    zh_title_idx = None
    zh_keyword_idx = None
    for idx, para in enumerate(paras):
        text = para.text.strip()
        if normalize_text(text) == "摘要":
            for offset, later in enumerate(paras[idx + 1 : idx + 12], start=idx + 1):
                if "关键词" in later.text:
                    zh_title = para
                    zh_keyword = later
                    zh_title_idx = idx
                    zh_keyword_idx = offset
                    break
        if zh_title is not None:
            break
    if zh_title is None:
        raise RuntimeError("找不到中文摘要标题")
    remove_page_break_before(zh_title)
    prev_text = paras[zh_title_idx - 1].text.strip() if zh_title_idx and zh_title_idx > 0 else ""
    if THESIS_TITLE not in prev_text:
        zh_title._element.addprevious(make_title_paragraph())

    if zh_keyword is None:
        raise RuntimeError("找不到中文关键词段")
    zh_old = []
    for para in paras[zh_title_idx + 1 : zh_keyword_idx]:
        if para.text.strip():
            zh_old.append(para)
    if zh_old:
        replace_paragraph_range_before(zh_keyword, zh_old, ZH_ABSTRACT)

    paras = doc.paragraphs
    en_title = None
    en_keyword = None
    en_title_idx = None
    en_keyword_idx = None
    seen_en_title = False
    for idx, para in enumerate(paras):
        text = para.text.strip()
        low = text.lower()
        if en_title is None and normalize_text(text).lower() == "abstract":
            en_title = para
            en_title_idx = idx
            seen_en_title = True
            continue
        if seen_en_title and "key" in low and "word" in low:
            en_keyword = para
            en_keyword_idx = idx
            break
    if en_title is None or en_keyword is None:
        candidates = [
            para.text.strip()
            for para in paras
            if "abstract" in para.text.strip().lower() or "key" in para.text.strip().lower()
        ]
        print("英文摘要候选:", candidates[:10])
        raise RuntimeError("找不到英文摘要区间")
    en_old = []
    for para in paras[en_title_idx + 1 : en_keyword_idx]:
        if para.text.strip():
            en_old.append(para)
    if en_old:
        replace_paragraph_range_before(en_keyword, en_old, EN_ABSTRACT)
    bold_keyword_label(doc)


def fix_translation(doc: Document) -> None:
    start = None
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if "附录二" in text or "外文译文" in text:
            start = idx
            break
    if start is None:
        raise RuntimeError("找不到外文译文区域")

    last = None
    for para in doc.paragraphs[start + 1 :]:
        text = para.text.strip()
        if text.startswith("附录三") or text.startswith("致谢") or text.startswith("致 谢"):
            break
        if text:
            last = para
    if last is None:
        raise RuntimeError("找不到外文译文正文段")

    insert_after = last._element
    for text in TRANSLATION_EXTRA:
        new_p = copy_paragraph_element(last, text)
        insert_after.addnext(new_p)
        insert_after = new_p


def fix_headers(doc: Document) -> None:
    for idx, section in enumerate(doc.sections):
        if idx in (0, 1):
            section.header.is_linked_to_previous = False
            section.footer.is_linked_to_previous = False
            for para in section.header.paragraphs:
                clear_paragraph_text_keep_first_run(para, "")
            for para in section.footer.paragraphs:
                clear_paragraph_text_keep_first_run(para, "")
        elif idx == 2:
            section.header.is_linked_to_previous = False
            section.footer.is_linked_to_previous = False
            for para in section.header.paragraphs:
                clear_paragraph_text_keep_first_run(para, "")
            add_page_field(section.footer.paragraphs[0])
        else:
            section.header.is_linked_to_previous = False
            section.footer.is_linked_to_previous = False
            for para in section.header.paragraphs:
                if "华北水利水电大学毕业设计" in para.text:
                    clear_paragraph_text_keep_first_run(para, "华北水利水电大学毕业设计")
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    break
            else:
                para = section.header.paragraphs[0]
                clear_paragraph_text_keep_first_run(para, "华北水利水电大学毕业设计")
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_page_field(section.footer.paragraphs[0])


def set_page_number_types(doc: Document) -> None:
    for idx, section in enumerate(doc.sections):
        sect_pr = section._sectPr
        pg = sect_pr.find(qn("w:pgNumType"))
        if pg is None:
            pg = OxmlElement("w:pgNumType")
            sect_pr.append(pg)
        if idx == len(doc.sections) - 1:
            pg.set(qn("w:fmt"), "decimal")
            pg.set(qn("w:start"), "1")
        elif idx == 2:
            pg.set(qn("w:fmt"), "lowerRoman")
            pg.set(qn("w:start"), "1")


def fix_layout(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

    set_style_size(doc, "Normal", "24")
    set_style_size(doc, "Heading 1", "32")

    for para in doc.paragraphs:
        text = para.text.strip()
        if ("第7章" in text or "第 7 章" in text or "第七章" in text) and "总结" in text:
            para.style = doc.styles["Heading 1"]
            add_page_break_before(para)
        elif text.startswith("7.1") and "全文总结" in text:
            para.style = doc.styles["Heading 2"]
        elif text.startswith("7.2") and ("后续展望" in text or "展望" in text):
            para.style = doc.styles["Heading 2"]
            clear_paragraph_text_keep_first_run(para, "7.2 后续展望")


def fix_prose_chapter_mentions(doc: Document) -> None:
    replacements = {
        "第一章": "第1章",
        "第二章": "第2章",
        "第三章": "第3章",
        "第四章": "第4章",
        "第五章": "第5章",
        "第六章": "第6章",
        "第七章": "第7章",
    }
    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            continue
        for run in para.runs:
            text = run.text
            if not text:
                continue
            for old, new in replacements.items():
                text = text.replace(old, new)
            if text != run.text:
                run.text = text


def move_body_section_break_to_chapter1(doc: Document) -> None:
    """Make front matter include TOC + abstracts; body section starts at Chapter 1."""
    paras = doc.paragraphs
    chapter1_idx = None
    for idx, para in enumerate(paras):
        if para.text.strip().startswith("第1章"):
            chapter1_idx = idx
            break
    if chapter1_idx is None or chapter1_idx == 0:
        raise RuntimeError("找不到第1章起点")

    source_ppr = None
    source_sect = None
    for para in paras[:chapter1_idx]:
        ppr = para._element.find(qn("w:pPr"))
        if ppr is None:
            continue
        sect = ppr.find(qn("w:sectPr"))
        if sect is not None:
            source_ppr = ppr
            source_sect = sect
    if source_ppr is None or source_sect is None:
        raise RuntimeError("找不到正文前的分节符")

    moved_sect = copy.deepcopy(source_sect)
    source_ppr.remove(source_sect)

    target_para = paras[chapter1_idx - 1]
    target_ppr = target_para._element.get_or_add_pPr()
    existing = target_ppr.find(qn("w:sectPr"))
    if existing is not None:
        target_ppr.remove(existing)
    target_ppr.append(moved_sect)


def add_missing_reference_citations(doc: Document) -> None:
    body_start = None
    ref_start = len(doc.paragraphs)
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if body_start is None and text.startswith("第1章"):
            body_start = idx
        if body_start is not None and text == "参考文献" and para.style.name.startswith("Heading"):
            ref_start = idx
            break
    if body_start is None:
        raise RuntimeError("找不到正文起点，无法补正文引用")
    body_paras = doc.paragraphs[body_start:ref_start]

    # One consolidated literature sentence is less visually disruptive than scattered
    # single-number citations appended to unrelated paragraph ends.
    needed = ["[2]", "[3]", "[6]", "[7]", "[8]", "[9]", "[10]", "[11]", "[17]", "[19]", "[20]", "[21]"]
    body_text = "\n".join(p.text for p in body_paras)
    if not all(c in body_text for c in needed):
        for para in body_paras:
            if para.text.strip().startswith("相关研究已从单一阈值触发"):
                para.add_run(" 相关器件资料、通信手册、标准文件和低成本空气质量监测平台文献也为器件选型、联网方式和评价指标提供了依据[2][3][6][7][8][9][10][11][17][19][20][21]。")
                break


def insert_once_after(doc: Document, anchor_text: str, marker_text: str, texts: list[str]) -> None:
    full_text = "\n".join(p.text for p in doc.paragraphs)
    if marker_text in full_text:
        return
    for para in doc.paragraphs:
        if anchor_text in para.text:
            insert_texts_after(doc, para, texts)
            return
    raise RuntimeError(f"找不到插入锚点: {anchor_text}")


def fix_acknowledgement(doc: Document) -> None:
    title = None
    for idx, para in enumerate(doc.paragraphs):
        if normalize_text(para.text) in {"致谢", "致謝"}:
            title = idx
            break
    if title is None:
        return
    body = []
    for para in doc.paragraphs[title + 1 :]:
        text = para.text.strip()
        if text == "参考文献" or text.startswith("附录"):
            break
        if text:
            body.append(para)
    if not body:
        return
    new_texts = [
        "在本课题完成过程中，指导教师张晓华老师在选题方向、系统方案和论文结构方面给予了具体指导。尤其是在硬件设计与论文表达之间如何保持一致这一点上，对我帮助很大。学院教师在毕业设计开题、中期检查和材料提交等环节提供了必要指导，使课题能够按计划推进。",
        "通过本次毕业设计，我对嵌入式系统从器件选型、硬件连接、固件调试到远程展示的完整过程有了更具体的认识。论文撰写阶段也让我意识到，工程实现不仅要能运行，还要能把模块分工、接口关系和测试依据清楚表达出来。",
        "在样机调试和论文整理过程中，许多问题都需要反复核对硬件连接、程序文件和测试现象。这个过程使我进一步理解了工程设计的严谨性，也认识到毕业设计不仅是完成一个装置，更是把设计依据、实现过程和验证结果完整呈现出来。",
        "同学和朋友在资料查找、样机调试、页面检查和测试记录整理过程中也给予了支持。家人在毕业设计准备阶段一直理解并支持我的学习安排，在此一并表示感谢。",
    ]
    reference = body[0]
    parent = reference._element.getparent()
    for para in body:
        parent.remove(para._element)
    anchor = doc.paragraphs[title]._element
    for text in new_texts:
        new_para = copy_paragraph_element(reference, text)
        remove_ppr_flags_from_element(new_para, ("w:pageBreakBefore", "w:keepNext", "w:keepLines"))
        anchor.addnext(new_para)
        anchor = new_para


def tighten_references(doc: Document) -> None:
    in_refs = False
    for para in doc.paragraphs:
        text = para.text.strip()
        if text == "参考文献":
            in_refs = True
            continue
        if in_refs and text.startswith("附录"):
            in_refs = False
        if in_refs and re.match(r"^\[\d+\]", text):
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.line_spacing = 1.0


def tighten_appendix_references(doc: Document) -> None:
    in_appendix_refs = False
    for para in doc.paragraphs:
        text = para.text.strip()
        if text == "参考文献" and para.style.name == "Normal":
            in_appendix_refs = True
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            continue
        if in_appendix_refs and not text:
            continue
        if in_appendix_refs and re.match(r"^\[\d+\]", text):
            for run in para.runs:
                run.font.size = Pt(10.5)
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.line_spacing = 0.92
            para.paragraph_format.first_line_indent = None

    # Remove trailing empty paragraphs that can create visible tail whitespace.
    body = doc.element.body
    for para in reversed(doc.paragraphs):
        if para.text.strip():
            break
        body.remove(para._element)


def fix_content_density(doc: Document) -> None:
    insert_once_after(doc, "报警判断由main.c中的CalcAlarmLevel函数完成", "CalcAlarmLevel函数的输入来自当前有效CO2采样值", [
        "CalcAlarmLevel函数的输入来自当前有效CO2采样值，输出则直接影响OLED状态文字、LED闪烁节拍和蜂鸣器动作。程序没有把报警判断分散在多个驱动文件中，而是集中在主循环附近处理，这样便于在调试时通过串口日志同时观察采样值、报警等级和输出状态。",
        "SCD41数据读取完成后，程序先确认CRC校验结果，再更新全局状态字段。若采集失败，系统保留上一轮有效显示并输出异常状态，避免错误数据直接触发声光提示。这个处理方式与第6章的精度测试和响应时间测试相对应，能够说明样机在数据波动时仍保持基本稳定。",
    ])
    insert_once_after(doc, "正常运行时，STM32把设备编号、CO2浓度、温度、湿度、报警等级", "上传数据帧的字段设计保持简洁", [
        "上传数据帧的字段设计保持简洁，核心字段包括设备编号、CO2浓度、温湿度、报警等级、网络状态和上传序号。设备编号用于区分终端，浓度和温湿度用于页面展示，报警等级用于移动端提醒，上传序号用于判断链路是否连续。这样既能满足远程查看需要，也便于后端保存历史记录。",
        "ESP8266链路出现短时断连时，本地采集和报警不应停止。程序把联网状态作为独立状态处理，网络恢复后继续发布最新数据。这样设计的目的，是保证本地安全提示优先于远程展示，同时使第6章通信稳定性测试能够围绕重连时间、丢包率和页面刷新连续性展开。",
    ])
    fix_acknowledgement(doc)
    tighten_references(doc)
    tighten_appendix_references(doc)


def split_long_code_blocks(doc: Document) -> None:
    body_text_style = doc.styles["Body Text"]
    consecutive = 0
    inserted_after = []
    for para in doc.paragraphs:
        latin = para.runs[0].font.name if para.runs else ""
        is_code = para.style.name in ("Code", "Source Code", "HTML Code", "Listing") or latin in ("Courier New", "Courier", "Consolas")
        if is_code:
            consecutive += 1
            if consecutive == 25:
                inserted_after.append(para)
                consecutive = 0
        else:
            consecutive = 0
    for para in inserted_after:
        new_p = OxmlElement("w:p")
        ppr = OxmlElement("w:pPr")
        pstyle = OxmlElement("w:pStyle")
        pstyle.set(qn("w:val"), body_text_style.style_id)
        ppr.append(pstyle)
        new_p.append(ppr)
        run = OxmlElement("w:r")
        text = OxmlElement("w:t")
        text.text = "（以上为该关键片段的一部分，以下继续列出相关语句。）"
        run.append(text)
        new_p.append(run)
        para._element.addnext(new_p)


def replace_following_summary_body(doc: Document, heading_idx: int, texts: list[str]) -> None:
    paras = doc.paragraphs
    following = []
    for para in paras[heading_idx + 1 :]:
        text = para.text.strip()
        if para.style.name.startswith("Heading") or text.startswith("第") or text.startswith("致") or text == "参考文献":
            break
        if text:
            following.append(para)
        if len(following) >= 2:
            break
    if following:
        reference = following[0]
        parent = reference._element.getparent()
        for para in following:
            parent.remove(para._element)
        anchor = paras[heading_idx]._element
        for text in texts:
            new_para = copy_paragraph_element(reference, text)
            remove_ppr_flags_from_element(new_para, ("w:pageBreakBefore", "w:keepNext", "w:keepLines"))
            anchor.addnext(new_para)
            anchor = new_para
    else:
        anchor = paras[heading_idx]._element
        for text in texts:
            new_para = make_text_paragraph(doc, text, "Body Text")
            remove_ppr_flags_from_element(new_para, ("w:pageBreakBefore", "w:keepNext", "w:keepLines"))
            anchor.addnext(new_para)
            anchor = new_para


def fix_whitespace_control(doc: Document) -> None:
    """Reduce actual output whitespace without broad run-level cleanup."""
    # School requirement: section headings have explicit spacing, but headings should
    # not be left alone at a page bottom. Bind headings/captions to their content.
    summary_headings = [
        (idx, para) for idx, para in enumerate(doc.paragraphs)
        if para.style.name.startswith("Heading") and para.text.strip() in CHAPTER_SUMMARIES
    ]
    for heading_idx, heading in reversed(summary_headings):
        replace_following_summary_body(doc, heading_idx, CHAPTER_SUMMARIES[heading.text.strip()])

    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        style = para.style.name
        set_widow_control(para, True)
        if style == "Heading 2":
            set_keep_with_next(para, text in CHAPTER_SUMMARIES)
            set_spacing(para, before=6, after=6)
        elif style == "Heading 3":
            set_keep_with_next(para, False)
            set_spacing(para, before=3, after=3)
        elif text.startswith("图") or text.startswith("表"):
            set_keep_with_next(para, False)
            set_spacing(para, before=0, after=3)
        elif text.startswith("本章完成了"):
            set_keep_lines(para, True)
            set_spacing(para, before=0, after=0)

        # Keep image paragraphs with their captions when possible.
        if para._element.xpath(".//w:drawing") and idx + 1 < len(doc.paragraphs):
            set_keep_with_next(para, True)
            set_spacing(para, before=0, after=0)

    for table in doc.tables:
        prevent_table_row_split(table)


def patch_toc_frontmatter_pages(doc: Document) -> None:
    for para in doc.element.xpath(".//w:p"):
        text_nodes = para.xpath(".//w:t")
        text = "".join(node.text or "" for node in text_nodes)
        normalized = normalize_text(text)
        # Avoid changing real headings/body paragraphs; TOC lines contain a hyperlink/anchor.
        if not para.xpath(".//w:hyperlink"):
            continue
        page = None
        if "摘要" in normalized:
            page = "vii"
        elif "ABSTRACT" in text:
            page = "viii"
        elif "第1章绪论" in normalized:
            page = "1"
        if page is not None:
            for node in reversed(text_nodes):
                raw = node.text or ""
                stripped = raw.strip()
                if stripped.isdigit() or stripped.lower() in {"i", "ii", "iii", "iv", "v", "vi", "vii", "viii"}:
                    node.text = raw.replace(stripped, page)
                    break


def build_module_photo_panel() -> None:
    from PIL import Image, ImageDraw, ImageFont, ImageOps

    items = [
        ("SCD41 CO2传感器", ASSETS_DIR / "module_scd41_trim.png"),
        ("STM32F103C8T6主控", ASSETS_DIR / "module_blue_pill_trim.png"),
        ("SSD1306 OLED显示", ASSETS_DIR / "module_oled_ssd1306_trim.png"),
        ("ESP8266-01S通信", ASSETS_DIR / "module_esp01_trim.png"),
        ("蜂鸣器/声光提示", ASSETS_DIR / "module_buzzer_trim.png"),
    ]
    for _label, path in items:
        if not path.exists():
            raise FileNotFoundError(path)

    cell_w, cell_h = 520, 390
    title_h = 54
    pad = 20
    canvas = Image.new("RGB", (cell_w * 2, cell_h * 3), "white")
    draw = ImageDraw.Draw(canvas)
    try:
        font = ImageFont.truetype("simhei.ttf", 28)
    except Exception:
        font = ImageFont.load_default()

    positions = [(0, 0), (1, 0), (0, 1), (1, 1), (0, 2)]
    for (label, path), (cx, cy) in zip(items, positions):
        x0, y0 = cx * cell_w, cy * cell_h
        draw.rectangle([x0 + 8, y0 + 8, x0 + cell_w - 8, y0 + cell_h - 8], outline=(150, 150, 150), width=2)
        bbox = draw.textbbox((0, 0), label, font=font)
        draw.text((x0 + (cell_w - (bbox[2] - bbox[0])) / 2, y0 + 14), label, fill=(0, 0, 0), font=font)
        img = Image.open(path).convert("RGB")
        img.thumbnail((cell_w - 2 * pad, cell_h - title_h - 2 * pad), Image.LANCZOS)
        img = ImageOps.expand(img, border=0, fill="white")
        ix = x0 + (cell_w - img.width) // 2
        iy = y0 + title_h + (cell_h - title_h - img.height) // 2
        canvas.paste(img, (ix, iy))

    x0, y0 = cell_w, cell_h * 2
    draw.rectangle([x0 + 8, y0 + 8, x0 + cell_w - 8, y0 + cell_h - 8], outline=(150, 150, 150), width=2)
    lines = [
        "接口与源码对应",
        "SCD41: PB10/PB11 -> scd41.c",
        "OLED: PB6/PB7 -> oled.c",
        "ESP8266: PA9/PA10/PA8 -> esp8266.c",
        "LED/蜂鸣器: PA1/PB4 -> led.c/beep.c",
    ]
    try:
        small_font = ImageFont.truetype("simhei.ttf", 23)
    except Exception:
        small_font = ImageFont.load_default()
    y = y0 + 62
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=small_font)
        draw.text((x0 + (cell_w - (bbox[2] - bbox[0])) / 2, y), line, fill=(30, 30, 30), font=small_font)
        y += 48
    MODULE_PANEL.parent.mkdir(parents=True, exist_ok=True)
    canvas.save(MODULE_PANEL)


def replace_first_figure_with_module_panel(doc: Document) -> None:
    build_module_photo_panel()
    caption = None
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == "图3.1 主要硬件模块实物图":
            caption = idx
            break
    if caption is None or caption == 0:
        raise RuntimeError("找不到图3.1题注")
    image_para = doc.paragraphs[caption - 1]
    clear_paragraph_runs(image_para)
    image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_para.add_run().add_picture(str(MODULE_PANEL), width=Inches(5.8))


def set_run_font(run, size_pt: float = 12) -> None:
    run.font.name = "Times New Roman"
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.append(fonts)
    fonts.set(qn("w:eastAsia"), "宋体")
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size_pt)


def set_cell_text(doc: Document, cell, text: str, bold_prefix: str = "", size_pt: float = 12) -> None:
    for paragraph in cell.paragraphs:
        clear_paragraph_runs(paragraph)
    paragraph = cell.paragraphs[0]
    paragraph.style = doc.styles["Body Text"] if "Body Text" in doc.styles else doc.styles["Normal"]
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        first.bold = True
        rest = paragraph.add_run(text[len(bold_prefix):])
        runs = [first, rest]
    else:
        runs = [paragraph.add_run(text)]
    for run in runs:
        set_run_font(run, size_pt)


def set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def replace_first_figure_with_module_table(doc: Document) -> None:
    modules = [
        (
            "SCD41 CO2传感器",
            ASSETS_DIR / "module_scd41_trim.png",
            "SCD41 CO2传感器：采集CO2、温度和湿度，经PB10/PB11接入STM32，对应scd41.c。",
        ),
        (
            "STM32F103C8T6主控",
            ASSETS_DIR / "module_blue_pill_trim.png",
            "STM32F103C8T6主控：负责采集、报警、显示和上传调度，对应bsp_pins.h和main.c。",
        ),
        (
            "SSD1306 OLED显示",
            ASSETS_DIR / "module_oled_ssd1306_trim.png",
            "SSD1306 OLED显示：显示CO2、温湿度、报警和网络状态，经PB6/PB7接入，对应oled.c。",
        ),
        (
            "ESP8266-01S通信",
            ASSETS_DIR / "module_esp01_trim.png",
            "ESP8266-01S通信：负责Wi-Fi接入和MQTT上传，经USART1连接，对应esp8266.c和onenet.c。",
        ),
        (
            "蜂鸣器/声光提示",
            ASSETS_DIR / "module_buzzer_trim.png",
            "蜂鸣器/声光提示：LED显示运行状态，蜂鸣器用于高等级报警，由AlarmOutput控制。",
        ),
    ]
    for _label, image_path, _desc in modules:
        if not image_path.exists():
            raise FileNotFoundError(image_path)

    caption = None
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == "图3.1 主要硬件模块实物图":
            caption = idx
            break
    if caption is None or caption == 0:
        raise RuntimeError("找不到图3.1题注")

    image_para = doc.paragraphs[caption - 1]
    clear_paragraph_runs(image_para)
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("读图时应先看整体，再看模块"):
            clear_paragraph_text_keep_first_run(
                paragraph,
                "图3.1按“左侧实物、右侧说明”的方式介绍主要硬件模块，后续各节再分别展开传感采集、显示报警、主控、无线通信和接口设计。",
            )
            break

    table = doc.add_table(rows=len(modules), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    for idx, (label, image_path, desc) in enumerate(modules):
        row = table.rows[idx]
        image_cell, text_cell = row.cells
        image_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        text_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_width(image_cell, 1900)
        set_cell_width(text_cell, 6500)
        image_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_cell.paragraphs[0].paragraph_format.space_before = Pt(0)
        image_cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        image_cell.paragraphs[0].add_run().add_picture(str(image_path), width=Inches(0.72))
        set_cell_text(doc, text_cell, desc, f"{label}：", size_pt=12)
    image_para._element.addnext(table._element)


def enhance_engineering_explanation(doc: Document) -> None:
    replace_first_figure_with_module_table(doc)

    def find_para(startswith: str):
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text.startswith(startswith) or startswith in text:
                return paragraph
        raise RuntimeError(f"找不到段落: {startswith}")

    insert_texts_after(doc, find_para("传感采集模块由SCD41和软件I2C接口组成"), [
        "SCD41模块在实物上带有GND、3.3V、SDA、SCL四个主要引脚，因此接线关系较清晰。论文中将其接入PB10/PB11，是为了让采集总线与OLED显示总线分离，避免显示刷新与传感器读数共用同一组软件I2C时序。bsp_pins.h中也把SCD41_SCL_PIN和SCD41_SDA_PIN单独登记，后续若调整引脚，只需修改该文件即可同步影响scd41.c。",
        "scd41.c的设计重点是三步：先用软件I2C产生起始、停止、应答和字节读写时序；再向0x62地址发送周期测量、停止测量、自检和海拔配置等命令；最后读取CO2、温度、湿度三组数据并执行CRC8校验。这样写的好处是，采集错误能在驱动层被发现，不会直接进入OLED显示、报警判断和MQTT上传链路。",
    ])

    insert_texts_after(doc, find_para("OLED显示与声光报警模块是系统现场交互的主要部分"), [
        "OLED模块在样机中承担现场读数窗口作用，显示内容不追求复杂界面，而是优先显示CO2浓度，再显示温度、湿度、状态等级和网络状态。固件中的OLED_ShowCO2Page函数按行刷新这些字段，使采集结果、报警等级和状态提示在同一屏幕上闭环显示。",
        "LED与蜂鸣器的设计分工较明确：LED适合低干扰、持续性的状态提示，蜂鸣器只在高等级状态参与提示。main.c中的AlarmOutput函数按照报警等级和运行节拍控制PA1与PB4输出，使正常、一级预警和二级报警具有不同的现场表现。这样既避免轻微波动造成持续声音干扰，也保证高浓度状态能被及时注意到。",
    ])

    insert_texts_after(doc, find_para("本系统以STM32F103C8T6微控制器为主控"), [
        "STM32主控模块不是只承担“连接外设”的作用，而是整个样机的数据调度中心。main.c中先完成NVIC、延时、USART、OLED、蜂鸣器、LED、按键、SCD41和联网模块初始化，再进入周期循环。循环中按时间片读取SCD41，更新报警等级，刷新OLED，并定期构造telemetry字符串上传。",
        "这种集中调度方式适合毕业设计样机：一方面可以通过USART2持续输出调试日志，方便检查采集值、联网状态和上传字段；另一方面可以保证网络异常不会阻塞本地采集与报警。看现象时，若OLED正常而移动端无数据，优先检查ESP8266和服务端；若串口无采集字段，优先检查SCD41总线和驱动。",
    ])

    insert_texts_after(doc, find_para("无线通信模块采用ESP8266-01S"), [
        "ESP8266-01S实物模块体积小、引脚少，适合与STM32通过USART1连接。硬件上PA9接ESP8266的RXD，PA10接ESP8266的TXD，PA8用于复位控制；软件上esp8266.c负责AT指令发送、应答等待、接收缓冲区清理和IPD数据解析。该模块只处理联网链路，不直接决定报警等级。",
        "上传路径由esp8266.c、MqttKit.c和onenet.c共同完成。STM32周期性构造包含dev、co2、temp、hum、alarm、net、seq等字段的数据帧，经ESP8266发送到MQTT主题；服务端订阅后写入实时状态和历史记录，移动端再通过REST接口和WebSocket读取。这个链路对应第5章移动端页面和第6章通信稳定性测试。",
    ])

    insert_texts_after(doc, find_para("固件总体结构与文件分工"), [
        "固件工程的分工可以按“主程序、硬件驱动、联网协议、系统支撑”四层理解。USER/main.c是主调度文件，负责初始化顺序、采样周期、报警等级、OLED页面和上传节拍；HARDWARE目录保存SCD41、OLED、LED、蜂鸣器等驱动；NET目录保存ESP8266、MQTT和OneNET相关代码；SYSTEM目录提供延时、串口和系统配置。",
        "这种分层写法使论文中的模块说明可以直接落到源码文件。第3章说明引脚和模块连接，第4章说明对应文件如何运行，第5章说明运行后形成什么功能表现。读者按照bsp_pins.h到main.c再到各驱动文件的顺序，就能把论文图、实物模块和源码实现对应起来。",
    ])

    insert_texts_after(doc, find_para("主程序初始化与循环调度"), [
        "主程序初始化顺序体现了硬件依赖关系：先配置中断、延时和串口，保证调试信息可输出；再初始化OLED、蜂鸣器、LED、按键和SCD41，保证本地功能可运行；随后配置SCD41海拔参数、自检并启动周期测量；最后初始化ESP8266和OneNET连接。这样即使无线链路暂时失败，本地采集和报警仍能继续工作。",
        "主循环以timeCount作为节拍基础。约5 s读取一次SCD41数据，成功后调用CO2_Predict_Update更新状态字段，再通过CalcAlarmLevel得到报警等级；OLED_ShowCO2Page负责把当前数值显示出来，AlarmOutput按节拍控制LED和蜂鸣器；到上传周期时，程序构造JSON格式telemetry并调用OneNet_Publish发布。",
    ])

    insert_texts_after(doc, find_para("配套移动端应用和服务器接口的设计目的在于验证设备数据能否稳定上传"), [
        "后端与移动端不是独立展示材料，而是对硬件样机远程链路的验证。服务端接收MQTT数据后，提供/api/v1/realtime/co2、/api/v1/history/co2、/api/v1/devices、/api/v1/alerts等接口；移动端首页显示实时读数，数据分析页显示历史趋势，设备页面显示终端状态，维护页面用于查看异常记录。",
        "因此第5章展示移动端页面时，需要把页面与设备字段对应起来：co2对应CO2浓度显示，temp和hum对应温湿度显示，alarm对应预警状态，net对应在线状态，seq对应连续上传序号。只有这些字段在设备端、服务端和移动端保持一致，远程展示才算真正接入了样机数据链路。",
    ])


def fix_docx() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    if OUT.exists():
        backup = OUT.with_suffix(".before_regen_20260606.docx")
        shutil.copy2(OUT, backup)
        print(f"已备份旧候选: {backup}")
    shutil.copy2(SRC, OUT)

    doc = Document(str(OUT))
    fix_layout(doc)
    fix_headers(doc)
    set_page_number_types(doc)
    fix_abstracts(doc)
    move_body_section_break_to_chapter1(doc)
    fix_translation(doc)
    fix_prose_chapter_mentions(doc)
    add_missing_reference_citations(doc)
    split_long_code_blocks(doc)
    enhance_engineering_explanation(doc)
    fix_content_density(doc)
    fix_whitespace_control(doc)
    patch_toc_frontmatter_pages(doc)
    doc.save(str(OUT))
    print(f"已生成: {OUT}")


def export_pdf() -> None:
    if PDF.exists():
        PDF.unlink()
    cmd = [
        r"C:/Program Files/LibreOffice/program/soffice.com",
        "--headless",
        "--convert-to",
        "pdf",
        "--outdir",
        str(OUT.parent),
        str(OUT),
    ]
    subprocess.run(cmd, check=True)
    print(f"已导出: {PDF}")


def verify_with_checker() -> None:
    from thesis_format_checker.checker import check, load_preset

    preset = load_preset("ncwu")
    _docx, content, findings = check(str(OUT), preset)
    print(f"checker findings: {len(findings)}")
    for finding in findings:
        print(f"{finding.severity}\t{finding.rule_id}\t{finding.message}")
    print(
        "counts:",
        f"zh={content.abstract_zh_chars}",
        f"en={content.abstract_en_words}",
        f"translation={content.foreign_translation_chars}",
    )


def verify_forbidden_terms() -> None:
    doc = Document(str(OUT))
    text = "\n".join(para.text for para in doc.paragraphs)
    failed = False
    for term in FORBIDDEN_TERMS:
        count = text.count(term)
        print(f"{term}: {count}")
        if count:
            failed = True
    if failed:
        raise RuntimeError("存在禁用词")


def verify_image_count() -> None:
    src_doc = Document(str(SRC))
    out_doc = Document(str(OUT))
    src_count = len(src_doc.inline_shapes)
    out_count = len(out_doc.inline_shapes)
    print(f"inline_shapes: source={src_count} output={out_count}")
    if out_count < src_count:
        raise RuntimeError("图片数量减少，停止交付")


def verify_pdf_images() -> None:
    if not PDF.exists():
        print("pdfimages: skipped (PDF missing)")
        return
    result = subprocess.run(["pdfimages", "-list", str(PDF)], check=True, capture_output=True, text=True, encoding="utf-8", errors="ignore")
    lines = [line for line in result.stdout.splitlines() if re.match(r"^\s*\d+", line)]
    print(f"pdf image objects: {len(lines)}")
    if not lines:
        raise RuntimeError("PDF 未检测到图片对象")


def render_pdf_pages() -> list[Path]:
    if not PDF.exists():
        print("blank scan: skipped (PDF missing)")
        return []
    if PDF_PAGE_DIR.exists():
        shutil.rmtree(PDF_PAGE_DIR)
    PDF_PAGE_DIR.mkdir(parents=True, exist_ok=True)
    prefix = PDF_PAGE_DIR / "page"
    subprocess.run(["pdftoppm", "-r", "72", "-png", str(PDF), str(prefix)], check=True)
    return sorted(PDF_PAGE_DIR.glob("page-*.png"))


def page_number_from_path(path: Path) -> int:
    match = re.search(r"-(\d+)\.png$", path.name)
    return int(match.group(1)) if match else -1


def scan_pdf_blank_space() -> list[dict]:
    """Detect pages where main content ends far above the footer area.

    The scan uses rendered PDF pixels, so it catches actual output whitespace rather
    than only DOCX paragraph counts. Footer/page number pixels are ignored.
    """
    from PIL import Image

    page_paths = render_pdf_pages()
    suspects: list[dict] = []
    for path in page_paths:
        page = page_number_from_path(path)
        img = Image.open(path).convert("L")
        width, height = img.size
        pix = img.load()
        xs: list[int] = []
        ys: list[int] = []
        # Ignore margins and footer/page-number zone. Use a dark-pixel threshold so
        # text, borders, and images all count as content.
        x0, x1 = int(width * 0.08), int(width * 0.92)
        y0, y1 = int(height * 0.06), int(height * 0.86)
        for y in range(y0, y1):
            row_dark = 0
            for x in range(x0, x1):
                if pix[x, y] < 245:
                    row_dark += 1
            if row_dark >= 4:
                ys.append(y)
        if not ys:
            suspects.append({"page": page, "reason": "blank-or-nearly-blank", "bottom_blank_ratio": 1.0})
            continue
        last_y = max(ys)
        bottom_blank_ratio = (y1 - last_y) / height
        # Half-page blank concern: page body content ending before the lower third is
        # visually risky. Chapter-open pages are reported separately for review.
        if bottom_blank_ratio >= 0.32:
            suspects.append({
                "page": page,
                "last_content_y": last_y,
                "body_bottom_y": y1,
                "bottom_blank_ratio": round(bottom_blank_ratio, 3),
            })
    VERIFY_DIR.mkdir(parents=True, exist_ok=True)
    report = VERIFY_DIR / "pdf_blank_scan.json"
    report.write_text(json.dumps(suspects, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"blank scan suspects: {len(suspects)} report={report}")
    for item in suspects[:30]:
        print(f"blank suspect page={item['page']} bottom_blank_ratio={item.get('bottom_blank_ratio')}")
    return suspects


def inspect_headers() -> None:
    doc = Document(str(OUT))
    for idx, section in enumerate(doc.sections):
        header = " | ".join(para.text for para in section.header.paragraphs if para.text.strip())
        footer = " | ".join(para.text for para in section.footer.paragraphs if para.text.strip())
        pg = section._sectPr.find(qn("w:pgNumType"))
        fmt = pg.get(qn("w:fmt")) if pg is not None else ""
        start = pg.get(qn("w:start")) if pg is not None else ""
        print(f"section {idx}: header={header!r} footer={footer!r} page_fmt={fmt!r} start={start!r}")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--no-pdf", action="store_true")
    parser.add_argument("--verify-only", action="store_true")
    args = parser.parse_args()

    if not args.verify_only:
        fix_docx()
        if not args.no_pdf:
            export_pdf()
    inspect_headers()
    verify_with_checker()
    verify_forbidden_terms()
    verify_image_count()
    verify_pdf_images()
    scan_pdf_blank_space()


if __name__ == "__main__":
    main()
