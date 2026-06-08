"""Single full-test entrypoint for thesis-format-checker.

Run from anywhere:

    python E:/My Project/thesis-format-checker/tests/run_full_tests.py

What it covers:
- Python syntax/bytecode compilation
- built-in unit contracts for rules and preset loading
- DOCX mutation regressions proving key rules fail on real broken documents
- real delivery iteration through delivery/run_delivery.py
- NCWU checker pass on the newly generated delivery version
- first-page cover contract against the school literature-review template
- color-consistency regression check against v011
- latest visual audit and blank-scan sanity checks
"""

from __future__ import annotations

import contextlib
import hashlib
import io
import json
import os
import posixpath
import re
import shutil
import subprocess
import sys
import tempfile
from dataclasses import dataclass
from pathlib import Path
from zipfile import ZipFile
from xml.etree import ElementTree as ET

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
SRC_DIR = ROOT / "src"
DOWNLOADS = Path(r"C:/Users/ASUS-KL/Downloads")
DOCX_STEM = "202213210刘高朋修改迭代版"
VERSIONED_DOCX_RE = re.compile(rf"^{re.escape(DOCX_STEM)}_v(\d{{3}})_(.+)\.docx$")
ORIGINAL = DOWNLOADS / f"{DOCX_STEM}.docx"
V011 = DOWNLOADS / "202213210刘高朋修改迭代版_v011_格式统一交付版.docx"
VERSION_LOG = DOWNLOADS / f"{DOCX_STEM}_版本记录.md"
EXPECTED_HEADER = "华北水利水电大学毕业设计"
COVER_TEMPLATE = DOWNLOADS / "202213210刘高朋_文献综述_标准模板版.docx"
COVER_TITLE = "毕业设计（论文）"
W_URI = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_URI = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
A_URI = "http://schemas.openxmlformats.org/drawingml/2006/main"
WP_URI = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
RELS_URI = "http://schemas.openxmlformats.org/package/2006/relationships"
W_NS = f"{{{W_URI}}}"
R_NS = f"{{{R_URI}}}"
A_NS = f"{{{A_URI}}}"
WP_NS = f"{{{WP_URI}}}"
RELS_NS = f"{{{RELS_URI}}}"

ET.register_namespace("w", W_URI)
ET.register_namespace("r", R_URI)

FORBIDDEN_TERMS = [
    "代码截图",
    "本实验",
    "实验",
    "本研究",
    "本论文",
    "软件算法",
    "算法",
    "预测",
    "补偿",
    "图3.6",
    "7.2 不足与展望",
    "2.2.1 系统需求分析",
    "文献启发",
    "本文按照",
]

ALLOWED_BLANK_PAGES = {2, 3, 23, 37, 56, 57, 73, 80}
ALLOWED_EAST_ASIA_FONTS = {"宋体", "黑体", "楷体", "仿宋_GB2312", "隶书", "Consolas"}
ALLOWED_LATIN_FONTS = {"Times New Roman", "Consolas", "宋体"}

COVER_TABLE_ROWS = [
    ("学    院", "电子工程学院"),
    ("专    业", "电子信息工程"),
    ("姓    名", "刘高朋"),
    ("学    号", "202213210"),
    ("指导教师", "张晓华"),
    ("完成时间", "2026年6月"),
]
BODY_MUTATION_ANCHOR = "本章回答本设计为什么需要开展"
ENGINEERING_PAPER_READY_ROOT = Path(
    r"E:/My Project/毕业设计论文/论文/我的论文/给老师_工程资料包_2026-05-26/06_论文插图与截图/07_paper_ready"
)
ENGINEERING_MIN_IMAGE_HEIGHT_EMU = 1_800_000
KICAD_FULL_SCHEMATIC_SOURCE = ENGINEERING_PAPER_READY_ROOT / "figure_04_kicad_component_schematic.png"
KICAD_LOCAL_FIGURES = {
    "图3.5（a） 电源输入与3.3V稳压局部截图": ENGINEERING_PAPER_READY_ROOT / "figure_04a_kicad_power_regulator_crop.png",
    "图3.5（b） I2C传感器与OLED接口局部截图": ENGINEERING_PAPER_READY_ROOT / "figure_04c_kicad_i2c_sensor_display_crop.png",
    "图3.5（c） STM32主控最小系统局部截图": ENGINEERING_PAPER_READY_ROOT / "figure_04b_kicad_mcu_minimum_crop.png",
    "图3.5（d） ESP8266无线通信接口局部截图": ENGINEERING_PAPER_READY_ROOT / "figure_04d_kicad_esp8266_crop.png",
    "图3.5（e） LED与蜂鸣器报警输出局部截图": ENGINEERING_PAPER_READY_ROOT / "figure_04e_kicad_alarm_output_crop.png",
    "图3.5（f） 调试、复位、时钟与ADC局部截图": ENGINEERING_PAPER_READY_ROOT / "figure_04f_kicad_debug_reset_clock_adc_crop.png",
}
ENGINEERING_FIGURES = {
    "图3.2 系统原理图": ENGINEERING_PAPER_READY_ROOT / "figure_01_system_principle_schematic.png",
    "图3.3 STM32主控接口原理图": ENGINEERING_PAPER_READY_ROOT / "figure_02_pre_bluepill_schematic.png",
    "图3.4 PCB顶层布线检查图": ENGINEERING_PAPER_READY_ROOT / "figure_03a_pre_bluepill_top_copper_review.png",
    **KICAD_LOCAL_FIGURES,
}
ENGINEERING_FIGURE_SOURCE_SHA256 = {
    "图3.5（a） 电源输入与3.3V稳压局部截图": "2675acdf012f975c38e11ecfc5afa4c2ccb001b59d700bbb989c6d24f580c57d",
    "图3.5（b） I2C传感器与OLED接口局部截图": "b129776112c20aa782edfd4e13e389a84720c30b8f4cba10f8ee8ecc69e3efbf",
    "图3.5（c） STM32主控最小系统局部截图": "903041437c1f51d2d4c838ccacca78b5f7c413352ac34ee5cf90a3fc177ccf59",
    "图3.5（d） ESP8266无线通信接口局部截图": "9d10a5992aa893e1102018983a9deaa7dbab04d37f5f11ca9541676f1d68e689",
    "图3.5（e） LED与蜂鸣器报警输出局部截图": "3fb42a2f55ad521f76635375cf4608c4966fd4131e4bd5bc8a6a91472d2912bd",
    "图3.5（f） 调试、复位、时钟与ADC局部截图": "37b008c0d09638e6adfc61a993c04955d607119cd1c9807caf8e6ceab1e95e20",
}
KICAD_FULL_SCHEMATIC_SHA256 = "576705875f1b4ba154fe631efa70e3bcd7d63bab47786d69566e8ac860537526"


@dataclass
class StepResult:
    name: str
    status: str
    detail: str = ""


ACTIVE_DOCX: Path | None = None
ACTIVE_SOURCE_DOCX: Path | None = None
ACTIVE_PDF: Path | None = None
ACTIVE_REPORT: Path | None = None
ACTIVE_BLANK_REPORT: Path | None = None


def run_command(args: list[str], *, cwd: Path = ROOT) -> None:
    subprocess.run(args, cwd=str(cwd), check=True)


def versioned_docx_candidates() -> list[tuple[int, str, Path]]:
    candidates: list[tuple[int, str, Path]] = []
    for path in DOWNLOADS.glob(f"{DOCX_STEM}_v*.docx"):
        match = VERSIONED_DOCX_RE.match(path.name)
        if not match:
            continue
        candidates.append((int(match.group(1)), match.group(2), path))
    return sorted(candidates, key=lambda item: item[0])


def version_info(path: Path) -> tuple[int, str]:
    match = VERSIONED_DOCX_RE.match(path.name)
    if not match:
        raise RuntimeError(f"not a versioned thesis DOCX: {path}")
    return int(match.group(1)), match.group(2)


def related_delivery_paths(docx_path: Path) -> tuple[Path, Path, Path]:
    version, _label = version_info(docx_path)
    pdf = docx_path.with_suffix(".pdf")
    report = DOWNLOADS / f"{DOCX_STEM}_v{version:03d}_格式检测报告.md"
    blank_report = DOWNLOADS / f"{DOCX_STEM}_v{version:03d}_留白扫描.json"
    return pdf, report, blank_report


def set_active_delivery(source_docx: Path, output_docx: Path) -> None:
    global ACTIVE_SOURCE_DOCX, ACTIVE_DOCX, ACTIVE_PDF, ACTIVE_REPORT, ACTIVE_BLANK_REPORT
    ACTIVE_SOURCE_DOCX = source_docx
    ACTIVE_DOCX = output_docx
    ACTIVE_PDF, ACTIVE_REPORT, ACTIVE_BLANK_REPORT = related_delivery_paths(output_docx)


def require_active_source_docx() -> Path:
    if ACTIVE_SOURCE_DOCX is None:
        raise RuntimeError("active source DOCX is not set; regenerate-latest step must run first")
    return ACTIVE_SOURCE_DOCX


def require_active_docx() -> Path:
    if ACTIVE_DOCX is None:
        raise RuntimeError("active delivery DOCX is not set; regenerate-latest step must run first")
    return ACTIVE_DOCX


def require_active_related_paths() -> tuple[Path, Path, Path]:
    if ACTIVE_PDF is None or ACTIVE_REPORT is None or ACTIVE_BLANK_REPORT is None:
        raise RuntimeError("active delivery related paths are not set; select-latest-delivery step must run first")
    return ACTIVE_PDF, ACTIVE_REPORT, ACTIVE_BLANK_REPORT


def document_text(path: Path) -> str:
    doc = Document(str(path))
    parts: list[str] = [paragraph.text for paragraph in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(parts)


def sha256(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def magick_command() -> str:
    command = shutil.which("magick") or r"C:\Program Files\ImageMagick-7.1.2-Q16-HDRI\magick.exe"
    if not Path(command).exists() and shutil.which(command) is None:
        raise RuntimeError("ImageMagick magick command is required to render SVG engineering figures")
    return command


def engineering_render_dir(docx_path: Path) -> Path:
    version, _label = version_info(docx_path)
    return DOWNLOADS / f"{DOCX_STEM}_v{version:03d}_engineering_assets"


def prepare_engineering_figure_assets(docx_path: Path) -> dict[str, Path]:
    render_dir = engineering_render_dir(docx_path)
    render_dir.mkdir(parents=True, exist_ok=True)
    rendered: dict[str, Path] = {}
    for index, (caption, source) in enumerate(ENGINEERING_FIGURES.items(), start=1):
        if not source.exists():
            raise RuntimeError(f"engineering source asset missing for {caption}: {source}")
        expected_source_sha = ENGINEERING_FIGURE_SOURCE_SHA256.get(caption)
        if expected_source_sha:
            actual_source_sha = sha256(source.read_bytes())
            if actual_source_sha != expected_source_sha:
                raise RuntimeError(
                    f"engineering source asset changed for {caption}: "
                    f"{source}={actual_source_sha}, expected={expected_source_sha}"
                )
        target = render_dir / f"figure_{index:02d}_{source.stem}.png"
        if source.suffix.lower() == ".svg":
            subprocess.run(
                [
                    magick_command(),
                    "-density",
                    "220",
                    "-background",
                    "white",
                    str(source),
                    "-alpha",
                    "remove",
                    "-alpha",
                    "off",
                    "-strip",
                    str(target),
                ],
                check=True,
            )
        else:
            shutil.copy2(source, target)
        rendered[caption] = target
    return rendered


def docx_target_zip_name(target: str) -> str:
    if target.startswith("/"):
        return posixpath.normpath(target.lstrip("/"))
    return posixpath.normpath(posixpath.join("word", target))


def xml_paragraph_text(paragraph: ET.Element) -> str:
    return "".join(node.text or "" for node in paragraph.findall(f".//{W_NS}t")).strip()


def engineering_figure_media_targets(docx_path: Path) -> dict[str, str]:
    with ZipFile(docx_path) as archive:
        document_root = ET.fromstring(archive.read("word/document.xml"))
        rels_root = ET.fromstring(archive.read("word/_rels/document.xml.rels"))

    rid_to_target = {
        rel.get("Id"): rel.get("Target")
        for rel in rels_root.findall(f".//{RELS_NS}Relationship")
    }
    body = document_root.find(f"{W_NS}body")
    if body is None:
        raise RuntimeError("DOCX document body missing")

    paragraphs = [child for child in body if child.tag == f"{W_NS}p"]
    media_targets: dict[str, str] = {}
    for index, paragraph in enumerate(paragraphs):
        caption = xml_paragraph_text(paragraph)
        if caption not in ENGINEERING_FIGURES:
            continue
        for previous in reversed(paragraphs[max(0, index - 4):index]):
            blips = previous.findall(f".//{A_NS}blip")
            if not blips:
                continue
            rid = blips[0].get(f"{R_NS}embed")
            target = rid_to_target.get(rid)
            if not target:
                raise RuntimeError(f"cannot resolve image relationship {rid} for {caption}")
            media_targets[caption] = docx_target_zip_name(target)
            break

    missing = sorted(set(ENGINEERING_FIGURES) - set(media_targets))
    if missing:
        raise RuntimeError(f"cannot locate engineering figure media before captions: {missing}")
    return media_targets


def engineering_figure_display_extents(docx_path: Path) -> dict[str, tuple[int, int]]:
    with ZipFile(docx_path) as archive:
        document_root = ET.fromstring(archive.read("word/document.xml"))

    body = document_root.find(f"{W_NS}body")
    if body is None:
        raise RuntimeError("DOCX document body missing")

    paragraphs = [child for child in body if child.tag == f"{W_NS}p"]
    extents: dict[str, tuple[int, int]] = {}
    for index, paragraph in enumerate(paragraphs):
        caption = xml_paragraph_text(paragraph)
        if caption not in ENGINEERING_FIGURES:
            continue
        for previous in reversed(paragraphs[max(0, index - 4):index]):
            if not previous.findall(f".//{A_NS}blip"):
                continue
            extent = previous.find(f".//{WP_NS}extent")
            if extent is None:
                raise RuntimeError(f"cannot locate image display extent for {caption}")
            cx = int(extent.get("cx") or "0")
            cy = int(extent.get("cy") or "0")
            extents[caption] = (cx, cy)
            break

    missing = sorted(set(ENGINEERING_FIGURES) - set(extents))
    if missing:
        raise RuntimeError(f"cannot locate engineering figure display extents before captions: {missing}")
    return extents


def verify_engineering_figure_assets(docx_path: Path) -> dict[str, str]:
    figure_assets = prepare_engineering_figure_assets(docx_path)
    media_targets = engineering_figure_media_targets(docx_path)
    display_extents = engineering_figure_display_extents(docx_path)
    with ZipFile(docx_path) as archive:
        embedded_media_hashes = {
            sha256(archive.read(name))
            for name in archive.namelist()
            if name.startswith("word/media/")
        }
        if KICAD_FULL_SCHEMATIC_SHA256 in embedded_media_hashes:
            raise RuntimeError("full KiCad schematic is still embedded; expected only localized 图3.5 crop screenshots")
        for caption, asset_path in figure_assets.items():
            media_name = media_targets[caption]
            actual = sha256(archive.read(media_name))
            expected = sha256(asset_path.read_bytes())
            if actual != expected:
                raise RuntimeError(
                    f"engineering figure asset mismatch for {caption}: "
                    f"{media_name}={actual}, rendered_source={expected}"
                )
            _cx, cy = display_extents[caption]
            if cy < ENGINEERING_MIN_IMAGE_HEIGHT_EMU:
                raise RuntimeError(
                    f"engineering figure is too small on page for {caption}: "
                    f"height_emu={cy}, expected>={ENGINEERING_MIN_IMAGE_HEIGHT_EMU}"
                )
    text = document_text(docx_path)
    if "图3.5 KiCad元器件连线原理图" in text:
        raise RuntimeError("stale whole-schematic 图3.5 caption remains in DOCX")
    return media_targets


def iter_direct_run_fonts(path: Path):
    with ZipFile(path) as archive:
        for name in archive.namelist():
            if not name.startswith("word/") or not name.endswith(".xml"):
                continue
            root = ET.fromstring(archive.read(name))
            for run in root.findall(f".//{W_NS}r"):
                text = "".join(node.text or "" for node in run.findall(f".//{W_NS}t")).strip()
                if not text:
                    continue
                rpr = run.find(f"{W_NS}rPr")
                if rpr is None:
                    continue
                fonts = rpr.find(f"{W_NS}rFonts")
                if fonts is None:
                    continue
                yield {
                    "source": name,
                    "text": text[:80],
                    "east_asia": fonts.get(f"{W_NS}eastAsia"),
                    "ascii": fonts.get(f"{W_NS}ascii"),
                    "hansi": fonts.get(f"{W_NS}hAnsi"),
                }


def xml_bytes(root: ET.Element) -> bytes:
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def mutate_docx_xml(src: Path, dst: Path, mutator) -> list[str]:
    """Copy a DOCX and mutate selected XML parts inside the ZIP package."""
    changed_parts: list[str] = []
    with ZipFile(src, "r") as zin, ZipFile(dst, "w") as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            new_data, changed = mutator(item.filename, data)
            if changed:
                changed_parts.append(item.filename)
            zout.writestr(item, new_data)
    if not changed_parts:
        raise RuntimeError(f"DOCX mutation changed no XML parts for {dst.name}")
    return changed_parts


def append_docx_paragraph(document_xml: bytes, text: str) -> tuple[bytes, bool]:
    root = ET.fromstring(document_xml)
    body = root.find(f"{W_NS}body")
    if body is None:
        return document_xml, False
    paragraph = ET.Element(f"{W_NS}p")
    run = ET.SubElement(paragraph, f"{W_NS}r")
    text_node = ET.SubElement(run, f"{W_NS}t")
    text_node.text = text
    children = list(body)
    section_index = next((i for i, child in enumerate(children) if child.tag == f"{W_NS}sectPr"), len(children))
    body.insert(section_index, paragraph)
    return xml_bytes(root), True


def make_xml_paragraph(text: str) -> ET.Element:
    paragraph = ET.Element(f"{W_NS}p")
    run = ET.SubElement(paragraph, f"{W_NS}r")
    text_node = ET.SubElement(run, f"{W_NS}t")
    text_node.text = text
    return paragraph


def paragraph_xml_text(paragraph: ET.Element) -> str:
    return "".join(node.text or "" for node in paragraph.findall(f".//{W_NS}t")).strip()


def insert_after_paragraph_prefix(document_xml: bytes, prefix: str, texts: list[str]) -> tuple[bytes, bool]:
    root = ET.fromstring(document_xml)
    body = root.find(f"{W_NS}body")
    if body is None:
        return document_xml, False
    children = list(body)
    for index, child in enumerate(children):
        if child.tag != f"{W_NS}p":
            continue
        if paragraph_xml_text(child).startswith(prefix):
            for offset, text in enumerate(texts, start=1):
                body.insert(index + offset, make_xml_paragraph(text))
            return xml_bytes(root), True
    return document_xml, False


def insert_before_paragraph_prefix(document_xml: bytes, prefix: str, texts: list[str]) -> tuple[bytes, bool]:
    root = ET.fromstring(document_xml)
    body = root.find(f"{W_NS}body")
    if body is None:
        return document_xml, False
    children = list(body)
    for index, child in enumerate(children):
        if child.tag != f"{W_NS}p":
            continue
        if paragraph_xml_text(child).startswith(prefix):
            for offset, text in enumerate(texts):
                body.insert(index + offset, make_xml_paragraph(text))
            return xml_bytes(root), True
    return document_xml, False


def remove_header_borders(name: str, data: bytes) -> tuple[bytes, bool]:
    if not name.startswith("word/header") or not name.endswith(".xml"):
        return data, False
    root = ET.fromstring(data)
    changed = False
    for ppr in root.findall(f".//{W_NS}pPr"):
        for child in list(ppr):
            if child.tag == f"{W_NS}pBdr":
                ppr.remove(child)
                changed = True
    return (xml_bytes(root), True) if changed else (data, False)


def remove_toc_dot_leaders(name: str, data: bytes) -> tuple[bytes, bool]:
    if name != "word/document.xml":
        return data, False
    root = ET.fromstring(data)
    changed = False
    for tab in root.findall(f".//{W_NS}tab"):
        if tab.get(f"{W_NS}leader") == "dot":
            tab.set(f"{W_NS}leader", "none")
            changed = True
    return (xml_bytes(root), True) if changed else (data, False)


def break_body_page_number_format(name: str, data: bytes) -> tuple[bytes, bool]:
    if name != "word/document.xml":
        return data, False
    root = ET.fromstring(data)
    sections = root.findall(f".//{W_NS}sectPr")
    if not sections:
        return data, False
    body_section = sections[-1]
    pg_num_type = body_section.find(f"{W_NS}pgNumType")
    if pg_num_type is None:
        pg_num_type = ET.SubElement(body_section, f"{W_NS}pgNumType")
    pg_num_type.set(f"{W_NS}fmt", "lowerRoman")
    pg_num_type.set(f"{W_NS}start", "3")
    return xml_bytes(root), True


def break_frontmatter_roman_numbering(name: str, data: bytes) -> tuple[bytes, bool]:
    if name != "word/document.xml":
        return data, False
    root = ET.fromstring(data)
    changed = False
    for pg_num_type in root.findall(f".//{W_NS}pgNumType"):
        if pg_num_type.get(f"{W_NS}fmt") == "lowerRoman":
            pg_num_type.set(f"{W_NS}fmt", "decimal")
            changed = True
    return (xml_bytes(root), True) if changed else (data, False)


def remove_first_line_indent(name: str, data: bytes) -> tuple[bytes, bool]:
    if name not in {"word/styles.xml", "word/document.xml"}:
        return data, False
    root = ET.fromstring(data)
    changed = False
    for ind in root.findall(f".//{W_NS}ind"):
        if ind.get(f"{W_NS}firstLine") is not None:
            ind.set(f"{W_NS}firstLine", "0")
            changed = True
    return (xml_bytes(root), True) if changed else (data, False)


def add_cover_footer_page_number(name: str, data: bytes) -> tuple[bytes, bool]:
    if not name.startswith("word/footer") or not name.endswith(".xml"):
        return data, False
    root = ET.fromstring(data)
    paragraph = ET.SubElement(root, f"{W_NS}p")
    ppr = ET.SubElement(paragraph, f"{W_NS}pPr")
    jc = ET.SubElement(ppr, f"{W_NS}jc")
    jc.set(f"{W_NS}val", "center")
    run = ET.SubElement(paragraph, f"{W_NS}r")
    text_node = ET.SubElement(run, f"{W_NS}t")
    text_node.text = "1"
    return xml_bytes(root), True


def add_isolated_figure_caption(name: str, data: bytes) -> tuple[bytes, bool]:
    if name != "word/document.xml":
        return data, False
    return append_docx_paragraph(data, "图 9.99 错误图题")


def add_isolated_table_caption(name: str, data: bytes) -> tuple[bytes, bool]:
    if name != "word/document.xml":
        return data, False
    return append_docx_paragraph(data, "表 9.99 错误表题")


def add_overlong_body_paragraph(name: str, data: bytes) -> tuple[bytes, bool]:
    if name != "word/document.xml":
        return data, False
    text = (
        "这一段用于验证正文阅读节奏规则。"
        "如果论文正文把多个设计依据、接口关系、程序流程、调试现象和测试结论全部塞进同一段，读者会很难快速判断这一段到底在讲哪一个层次。"
        "规范要求正文段落应保持适度长度，超过阈值时应拆成若干逻辑段，并让每一段分别承担背景、实现、验证或结论中的一个功能。"
        "这里故意写成很长的一段，确保检测器能够在真实DOCX中识别段落过密的问题，而不是只在人工构造的字符串上通过测试。"
        "该回归用例同时避免误伤英文摘要和参考文献等特殊区域，只要求正文主体章节保持稳定的讲述节奏。"
        "如果后续修改把这条规则绕开，测试应当立即失败，因为论文迭代的目标不是把材料继续堆长，而是把设计依据、实现过程和验证结论分成清楚可读的段落。"
        "同时，这段文字必须被插入到正文第1章之后，而不是目录、摘要或封面区域，否则规则按正文范围过滤时就不会触发。"
        "这能证明测试入口确实理解论文结构，而不是只靠随意拼接XML制造表面通过的检查。"
    )
    return insert_after_paragraph_prefix(data, BODY_MUTATION_ANCHOR, [text])


def add_unexplained_figure_caption(name: str, data: bytes) -> tuple[bytes, bool]:
    if name != "word/document.xml":
        return data, False
    return insert_after_paragraph_prefix(data, BODY_MUTATION_ANCHOR, ["图9.91 缺少前后说明的测试图"])


def add_figure_caption_without_followup(name: str, data: bytes) -> tuple[bytes, bool]:
    if name != "word/document.xml":
        return data, False
    return insert_before_paragraph_prefix(data, "致 谢", ["图9.92 缺少后续说明的测试图"])


def add_unbalanced_figure_group(name: str, data: bytes) -> tuple[bytes, bool]:
    if name != "word/document.xml":
        return data, False
    captions = [
        "图9.91 连续截图一",
        "图9.92 连续截图二",
        "图9.93 连续截图三",
        "图9.94 连续截图四",
    ]
    return insert_after_paragraph_prefix(data, BODY_MUTATION_ANCHOR, captions)


def remove_module_visual_lead_text(name: str, data: bytes) -> tuple[bytes, bool]:
    if name != "word/document.xml":
        return data, False
    root = ET.fromstring(data)
    changed = False
    for text_node in root.findall(f".//{W_NS}t"):
        text = text_node.text or ""
        patched = (
            text
            .replace("左侧实物", "左侧模块")
            .replace("右侧说明", "右侧文字")
            .replace("源码", "文件")
        )
        if patched != text:
            text_node.text = patched
            changed = True
    return (xml_bytes(root), True) if changed else (data, False)


def set_caption_style_italic(name: str, data: bytes) -> tuple[bytes, bool]:
    if name != "word/styles.xml":
        return data, False
    root = ET.fromstring(data)
    changed = False
    for style in root.findall(f".//{W_NS}style"):
        style_id = style.get(f"{W_NS}styleId") or ""
        name_el = style.find(f"{W_NS}name")
        style_name = name_el.get(f"{W_NS}val") if name_el is not None else ""
        if style_id.lower() != "caption" and style_name.lower() != "caption":
            continue
        rpr = style.find(f"{W_NS}rPr")
        if rpr is None:
            rpr = ET.SubElement(style, f"{W_NS}rPr")
        italic = rpr.find(f"{W_NS}i")
        if italic is None:
            italic = ET.SubElement(rpr, f"{W_NS}i")
        italic.set(f"{W_NS}val", "1")
        changed = True
    return (xml_bytes(root), True) if changed else (data, False)


def remove_engineering_evidence_caption(name: str, data: bytes) -> tuple[bytes, bool]:
    if name != "word/document.xml":
        return data, False
    root = ET.fromstring(data)
    changed = False
    for text_node in root.findall(f".//{W_NS}t"):
        text = text_node.text or ""
        patched = text.replace("图3.2 系统原理图", "图3.2 系统硬件连接与接口关系图")
        if patched != text:
            text_node.text = patched
            changed = True
    return (xml_bytes(root), True) if changed else (data, False)


def remove_table_cant_split(name: str, data: bytes) -> tuple[bytes, bool]:
    if name != "word/document.xml":
        return data, False
    root = ET.fromstring(data)
    changed = False
    for trpr in root.findall(f".//{W_NS}trPr"):
        for child in list(trpr):
            if child.tag == f"{W_NS}cantSplit":
                trpr.remove(child)
                changed = True
    return (xml_bytes(root), True) if changed else (data, False)


def run_size_pt(run) -> float | None:
    if run.font.size is not None:
        return run.font.size.pt
    rpr = run._element.find(qn("w:rPr"))
    if rpr is None:
        return None
    size = rpr.find(qn("w:sz"))
    if size is None:
        return None
    value = size.get(qn("w:val"))
    return int(value) / 2 if value and value.isdigit() else None


def first_visible_run(paragraph):
    for run in paragraph.runs:
        if run.text.strip():
            return run
    return None


def step_compileall() -> str:
    run_command([
        sys.executable,
        "-m",
        "compileall",
        str(ROOT / "src" / "thesis_format_checker"),
        str(ROOT / "delivery" / "build_lgp_docx.py"),
        str(ROOT / "delivery" / "run_delivery.py"),
        str(ROOT / "tests" / "run_full_tests.py"),
    ])
    return "compileall passed"


def step_unit_contracts() -> str:
    if str(SRC_DIR) not in sys.path:
        sys.path.insert(0, str(SRC_DIR))
    from thesis_format_checker.standard.rules import RULES
    from thesis_format_checker.checker import load_preset

    preset = load_preset("ncwu")
    enabled_preset_rules = {
        item["id"]
        for item in preset.get("rules", [])
        if item.get("enabled", True)
    }
    missing = enabled_preset_rules - set(RULES.keys())
    if missing:
        raise RuntimeError(f"missing rule registrations: {sorted(missing)}")

    if preset["preset_id"] != "ncwu":
        raise RuntimeError(f"unexpected preset_id: {preset['preset_id']}")
    if preset["page"]["margin_cm"]["left"] != 3.0:
        raise RuntimeError("unexpected left margin in preset")
    if preset["styles"]["body"]["size_pt"] != 12:
        raise RuntimeError("unexpected body font size in preset")
    if preset["content"]["abstract_zh"]["min_chars"] != 500:
        raise RuntimeError("unexpected zh abstract threshold in preset")
    expected_standard_rules = {
        "readability-paragraph-length",
        "figure-lead-text",
        "figure-followup-text",
        "figure-text-balance",
        "module-visual-block",
        "caption-not-italic",
        "engineering-evidence-chain",
    }
    missing_standard_rules = expected_standard_rules - enabled_preset_rules
    if missing_standard_rules:
        raise RuntimeError(f"readability/visual rules not enabled in preset: {sorted(missing_standard_rules)}")
    readability = preset.get("readability", {})
    if readability.get("max_body_paragraph_chars") != 360:
        raise RuntimeError("unexpected body paragraph rhythm threshold")
    module_block = readability.get("module_block", {})
    if not module_block.get("enabled"):
        raise RuntimeError("module visual block rule is not enabled")
    if "图3.1 主要硬件模块实物图" not in module_block.get("target_captions", []):
        raise RuntimeError("module visual block target caption is missing")
    if preset.get("styles", {}).get("caption", {}).get("italic") is not False:
        raise RuntimeError("caption style contract must explicitly disable italic")
    evidence_chain = readability.get("engineering_evidence_chain", {})
    if not evidence_chain.get("enabled"):
        raise RuntimeError("engineering evidence chain rule is not enabled")
    for caption in ENGINEERING_FIGURES:
        if caption not in evidence_chain.get("required_captions", []):
            raise RuntimeError(f"engineering evidence chain missing required caption: {caption}")

    return f"rule registry and preset contracts passed: enabled_rules={len(enabled_preset_rules)}"


def step_rule_regression_contracts() -> str:
    if str(SRC_DIR) not in sys.path:
        sys.path.insert(0, str(SRC_DIR))
    from thesis_format_checker.checker import check, load_preset

    active_docx = require_active_docx()
    if not active_docx.exists():
        raise RuntimeError(f"active DOCX missing before rule regression tests: {active_docx}")

    preset = load_preset("ncwu")
    regressions = [
        ("header-underline", remove_header_borders),
        ("toc-dotline", remove_toc_dot_leaders),
        ("page-number-format", break_body_page_number_format),
        ("page-number-roman-frontmatter", break_frontmatter_roman_numbering),
        ("body-first-line-indent", remove_first_line_indent),
        ("cover-no-page-number", add_cover_footer_page_number),
        ("figure-caption-position", add_isolated_figure_caption),
        ("readability-paragraph-length", add_overlong_body_paragraph),
        ("figure-lead-text", add_unexplained_figure_caption),
        ("figure-followup-text", add_figure_caption_without_followup),
        ("figure-text-balance", add_unbalanced_figure_group),
        ("module-visual-block", remove_module_visual_lead_text),
        ("caption-not-italic", set_caption_style_italic),
        ("engineering-evidence-chain", remove_engineering_evidence_caption),
        ("table-caption-position", add_isolated_table_caption),
        ("table-no-page-break", remove_table_cant_split),
    ]

    passed: list[str] = []
    with tempfile.TemporaryDirectory(prefix="thesis-rule-regression-") as tmpdir:
        tmp_root = Path(tmpdir)
        for rule_id, mutator in regressions:
            bad_docx = tmp_root / f"bad-{rule_id}.docx"
            changed_parts = mutate_docx_xml(active_docx, bad_docx, mutator)
            _docx, _content, findings = check(bad_docx, preset)
            rule_ids = {finding.rule_id for finding in findings}
            if rule_id not in rule_ids:
                preview = "; ".join(f"{f.rule_id}: {f.message}" for f in findings[:6])
                raise RuntimeError(
                    f"{rule_id} did not trigger on mutated DOCX; "
                    f"changed_parts={changed_parts}; findings={preview}"
                )
            passed.append(rule_id)

    return f"mutated DOCX regressions passed: {', '.join(passed)}"


def step_regenerate_latest() -> str:
    if str(ROOT) not in sys.path:
        sys.path.insert(0, str(ROOT))
    if str(SRC_DIR) not in sys.path:
        sys.path.insert(0, str(SRC_DIR))

    before = versioned_docx_candidates()
    if not before:
        raise RuntimeError(f"no versioned source DOCX found in {DOWNLOADS}")
    source_version, _source_label, source_docx = before[-1]

    from delivery import run_delivery

    output = io.StringIO()
    try:
        with contextlib.redirect_stdout(output), contextlib.redirect_stderr(output):
            run_delivery.main()
    except Exception:
        print(output.getvalue())
        raise

    after = versioned_docx_candidates()
    if not after:
        raise RuntimeError("no versioned DOCX found after generation")
    output_version, _output_label, output_docx = after[-1]
    if output_version != source_version + 1:
        raise RuntimeError(
            f"delivery version did not auto-increment: source=v{source_version:03d}, latest=v{output_version:03d}"
        )
    if output_docx == source_docx:
        raise RuntimeError("delivery builder reused the source DOCX path")
    set_active_delivery(source_docx, output_docx)
    return f"generated v{output_version:03d} from v{source_version:03d}: {output_docx.name}"


def step_delivery_contract() -> str:
    active_docx = require_active_docx()
    source_docx = require_active_source_docx()
    active_pdf, active_report, active_blank_report = require_active_related_paths()
    active_version, active_label = version_info(active_docx)
    source_version, _source_label = version_info(source_docx)

    required = [ORIGINAL, source_docx, active_docx, active_pdf, active_report, active_blank_report, VERSION_LOG]
    missing = [str(path) for path in required if not path.exists()]
    if missing:
        raise RuntimeError(f"missing delivery artifacts: {missing}")
    if active_docx.suffix.lower() != ".docx":
        raise RuntimeError(f"final artifact is not a DOCX: {active_docx}")
    if active_version != source_version + 1:
        raise RuntimeError(f"final DOCX is not next version: source=v{source_version:03d}, active=v{active_version:03d}")
    if ORIGINAL.name == active_docx.name:
        raise RuntimeError("final DOCX overwrote the original filename")

    log_text = VERSION_LOG.read_text(encoding="utf-8")
    if active_docx.name not in log_text or f"v{active_version:03d} - {active_label}" not in log_text:
        raise RuntimeError("version log is missing current delivery facts")
    return f"v{active_version:03d} DOCX/PDF/report/version-log artifacts present"


def step_cover_contract() -> str:
    if not COVER_TEMPLATE.exists():
        raise RuntimeError(f"cover reference template missing: {COVER_TEMPLATE}")
    doc = Document(str(require_active_docx()))
    if len(doc.paragraphs) < 10 or not doc.tables:
        raise RuntimeError("cover structure is missing")

    school = doc.paragraphs[1]
    title = doc.paragraphs[3]
    topic = doc.paragraphs[4]
    footer = doc.paragraphs[9]
    if school.text.strip() != "华北水利水电大学":
        raise RuntimeError(f"cover school name mismatch: {school.text!r}")
    if title.text.strip() != COVER_TITLE:
        raise RuntimeError(f"cover title mismatch: {title.text!r}")
    cover_text = "\n".join(paragraph.text for paragraph in doc.paragraphs[:10])
    if "文献综述" in cover_text or "毕 业 设 计" in cover_text:
        raise RuntimeError(f"cover contains stale title text: {cover_text!r}")

    school_run = first_visible_run(school)
    title_run = first_visible_run(title)
    topic_run = first_visible_run(topic)
    footer_run = first_visible_run(footer)
    if school_run is None or (run_size_pt(school_run) or 0) < 40 or school_run.bold is not True:
        raise RuntimeError("cover school name is not template-scale bold type")
    if title_run is None or (run_size_pt(title_run) or 0) < 28 or title_run.bold is not True:
        raise RuntimeError("cover title is not template-scale bold type")
    if topic_run is None or (run_size_pt(topic_run) or 0) < 17 or topic_run.bold is not True:
        raise RuntimeError("cover topic line is not template-scale bold type")
    if footer_run is None or (run_size_pt(footer_run) or 0) < 13:
        raise RuntimeError("cover footer is not template size")

    table = doc.tables[0]
    for row_index, (label, value) in enumerate(COVER_TABLE_ROWS):
        row = table.rows[row_index]
        actual_label = row.cells[0].text.strip()
        actual_value = row.cells[1].text.strip()
        if actual_label != label or actual_value != value:
            raise RuntimeError(f"cover table row {row_index} mismatch: {(actual_label, actual_value)}")
        for cell in row.cells[:2]:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if not run.text.strip():
                        continue
                    if (run_size_pt(run) or 0) < 13 or run.bold is not True:
                        raise RuntimeError(f"cover table style mismatch at row {row_index}: {run.text!r}")
    return "cover matches historical template contract with main-thesis title"


def step_check_latest() -> str:
    if str(SRC_DIR) not in sys.path:
        sys.path.insert(0, str(SRC_DIR))
    from thesis_format_checker.checker import check, load_preset

    active_docx = require_active_docx()
    preset = load_preset("ncwu")
    _docx, _content, findings = check(active_docx, preset)
    if findings:
        detail = "; ".join(f"{f.rule_id}: {f.message}" for f in findings[:5])
        raise RuntimeError(f"latest checker findings={len(findings)} {detail}")
    return "latest checker findings=0"


def step_content_integrity_contract() -> str:
    if str(SRC_DIR) not in sys.path:
        sys.path.insert(0, str(SRC_DIR))
    from thesis_format_checker.checker import check, load_preset

    active_docx = require_active_docx()
    preset = load_preset("ncwu")
    _docx, content, _findings = check(active_docx, preset)
    if content.abstract_zh_chars < 500:
        raise RuntimeError(f"zh abstract too short: {content.abstract_zh_chars}")
    if content.abstract_en_words < 300:
        raise RuntimeError(f"en abstract too short: {content.abstract_en_words}")
    if content.foreign_translation_chars < 2000:
        raise RuntimeError(f"foreign translation too short: {content.foreign_translation_chars}")

    expected_chapters = {f"第{i}章" for i in range(1, 8)}
    chapter_text = "\n".join(content.chapters)
    missing_chapters = [chapter for chapter in sorted(expected_chapters) if chapter not in chapter_text]
    if missing_chapters:
        raise RuntimeError(f"missing chapters: {missing_chapters}")

    required_markers = ["摘 要", "ABSTRACT", "参考文献", "附录一", "附录二", "附录三"]
    missing_markers = [marker for marker in required_markers if marker not in content.full_text]
    if missing_markers:
        raise RuntimeError(f"missing structural markers: {missing_markers}")
    return "abstracts/translation/chapters/appendices present"


def step_header_contract() -> str:
    doc = Document(str(require_active_docx()))
    headers = []
    for section in doc.sections:
        text = " ".join(paragraph.text.strip() for paragraph in section.header.paragraphs if paragraph.text.strip()).strip()
        if text:
            headers.append(text)
    if not headers:
        raise RuntimeError("no body headers found")
    bad_headers = [header for header in headers if header != EXPECTED_HEADER]
    if bad_headers:
        raise RuntimeError(f"unexpected headers: {bad_headers}")
    if any(("(" in header or ")" in header or "（" in header or "）" in header or "论文" in header) for header in headers):
        raise RuntimeError(f"header still contains parentheses or thesis suffix: {headers}")
    return f"headers exact: {EXPECTED_HEADER}"


def step_font_contract() -> str:
    if str(SRC_DIR) not in sys.path:
        sys.path.insert(0, str(SRC_DIR))
    from thesis_format_checker.docx_inspector import inspect

    active_docx = require_active_docx()
    docx = inspect(active_docx)
    style_expectations = {
        "Normal": ("宋体", "Times New Roman"),
        "Body Text": ("宋体", "Times New Roman"),
    }
    for style_name, (expected_ea, expected_latin) in style_expectations.items():
        style = docx.styles.get(style_name)
        if style is None:
            raise RuntimeError(f"missing style: {style_name}")
        if style.east_asia and style.east_asia != expected_ea:
            raise RuntimeError(f"{style_name} eastAsia={style.east_asia}, expected {expected_ea}")
        if style.latin and style.latin != expected_latin:
            raise RuntimeError(f"{style_name} latin={style.latin}, expected {expected_latin}")

    bad_fonts = []
    for item in iter_direct_run_fonts(active_docx):
        ea = item["east_asia"]
        ascii_font = item["ascii"]
        hansi = item["hansi"]
        if ea and ea not in ALLOWED_EAST_ASIA_FONTS:
            bad_fonts.append(item)
        if ascii_font and ascii_font not in ALLOWED_LATIN_FONTS:
            bad_fonts.append(item)
        if hansi and hansi not in ALLOWED_LATIN_FONTS:
            bad_fonts.append(item)
        if len(bad_fonts) >= 5:
            break
    if bad_fonts:
        raise RuntimeError(f"unexpected direct fonts: {bad_fonts}")
    return "style and direct font families are within allowed thesis set"


def step_forbidden_terms_contract() -> str:
    text = document_text(require_active_docx())
    hits = {term: text.count(term) for term in FORBIDDEN_TERMS if text.count(term)}
    if hits:
        raise RuntimeError(f"forbidden terms present: {hits}")
    return "forbidden prose terms absent"


def step_bridge_paragraph_contract() -> str:
    if str(ROOT) not in sys.path:
        sys.path.insert(0, str(ROOT))
    from delivery import build_lgp_docx

    build_lgp_docx.verify_bridge_paragraphs_once(require_active_docx())
    return "generated bridge paragraphs present exactly once"


def step_image_table_contract() -> str:
    source_docx = require_active_source_docx()
    active_docx = require_active_docx()
    before = Document(str(source_docx))
    after = Document(str(active_docx))
    if len(after.inline_shapes) < len(before.inline_shapes):
        raise RuntimeError(f"inline image count decreased: source={len(before.inline_shapes)} output={len(after.inline_shapes)}")
    if len(after.tables) < len(before.tables):
        raise RuntimeError(f"table count decreased: source={len(before.tables)} output={len(after.tables)}")
    engineering_targets = verify_engineering_figure_assets(active_docx)
    return (
        f"images/tables preserved: images={len(after.inline_shapes)}, tables={len(after.tables)}; "
        f"engineering figures verified={engineering_targets}"
    )


def step_regression_v011_color_rule() -> str:
    if str(SRC_DIR) not in sys.path:
        sys.path.insert(0, str(SRC_DIR))
    from thesis_format_checker.checker import check, load_preset

    preset = load_preset("ncwu")
    _docx, _content, findings = check(V011, preset)
    color_findings = [f for f in findings if f.rule_id == "text-color-consistency"]
    if not color_findings:
        raise RuntimeError("v011 no longer triggers text-color-consistency regression warning")
    return "v011 triggers text-color-consistency as expected"


def step_visual_audit() -> str:
    if str(ROOT) not in sys.path:
        sys.path.insert(0, str(ROOT))
    from delivery import build_lgp_docx

    audit = build_lgp_docx.audit_visual_format()
    if audit["non_black_runs"] or audit["style_non_black"]:
        raise RuntimeError(f"latest color audit failed: {audit}")
    return f"color audit passed: {audit}"


def step_blank_scan_sanity() -> str:
    _active_pdf, _active_report, active_blank_report = require_active_related_paths()
    if not active_blank_report.exists():
        raise RuntimeError(f"blank scan report missing: {active_blank_report}")
    suspects = json.loads(active_blank_report.read_text(encoding="utf-8"))
    pages = {item.get("page") for item in suspects}
    unexpected = pages - ALLOWED_BLANK_PAGES
    if unexpected:
        raise RuntimeError(f"unexpected blank-scan pages: {sorted(unexpected)}")
    if 90 in pages:
        raise RuntimeError("reference orphan tail page returned at page 90")
    if len(suspects) > len(ALLOWED_BLANK_PAGES):
        raise RuntimeError(f"blank suspects increased: {len(suspects)}")
    return f"blank scan sanity passed: suspects={len(suspects)} pages={sorted(pages)}"


def run_step(name: str, fn) -> StepResult:
    print(f"\n== {name} ==")
    try:
        detail = fn()
        print(f"PASS {detail}")
        return StepResult(name, "PASS", detail)
    except Exception as exc:
        print(f"FAIL {exc}")
        return StepResult(name, "FAIL", str(exc))


def main() -> int:
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    if hasattr(sys.stderr, "reconfigure"):
        sys.stderr.reconfigure(encoding="utf-8", errors="replace")
    os.environ["PYTHONIOENCODING"] = "utf-8"
    os.chdir(ROOT)

    steps = [
        ("compileall", step_compileall),
        ("unit-contracts", step_unit_contracts),
        ("regenerate-latest", step_regenerate_latest),
        ("delivery-contract", step_delivery_contract),
        ("cover-contract", step_cover_contract),
        ("check-latest", step_check_latest),
        ("rule-regression-contracts", step_rule_regression_contracts),
        ("content-integrity", step_content_integrity_contract),
        ("header-contract", step_header_contract),
        ("font-contract", step_font_contract),
        ("forbidden-terms", step_forbidden_terms_contract),
        ("bridge-paragraph-contract", step_bridge_paragraph_contract),
        ("image-table-contract", step_image_table_contract),
        ("regression-v011-color-rule", step_regression_v011_color_rule),
        ("visual-audit", step_visual_audit),
        ("blank-scan-sanity", step_blank_scan_sanity),
    ]
    results = [run_step(name, fn) for name, fn in steps]

    print("\n== summary ==")
    for result in results:
        print(f"{result.status:4} {result.name} {result.detail}")

    failed = [result for result in results if result.status == "FAIL"]
    return 1 if failed else 0


if __name__ == "__main__":
    raise SystemExit(main())
