"""DOCX format-property inspection via python-docx + lxml.

Extracts raw formatting facts (margins, fonts, sizes, styles, headers).
Does NOT make judgments — that's the rules engine's job.
"""

from __future__ import annotations

from dataclasses import dataclass, field, asdict
from pathlib import Path
import re
from zipfile import ZipFile
from xml.etree import ElementTree as ET
from typing import Any

from docx import Document
from docx.oxml.ns import qn

W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
EMU_PER_CM = 360000


@dataclass
class StyleInfo:
    name: str
    style_id: str | None = None
    east_asia: str | None = None
    latin: str | None = None
    size_pt: float | None = None
    bold: bool | None = None
    align: str | None = None
    line_spacing: float | None = None
    first_line_indent_twips: int | None = None
    color: str | None = None
    theme_color: str | None = None


@dataclass
class RunColorInfo:
    source: str
    text: str
    color: str | None = None
    theme_color: str | None = None


@dataclass
class SectionInfo:
    index: int
    page_width_cm: float
    page_height_cm: float
    top_margin_cm: float
    bottom_margin_cm: float
    left_margin_cm: float
    right_margin_cm: float
    header_text: str = ""
    header_linked_to_previous: bool = True
    header_has_bottom_border: bool = False
    header_align: str | None = None
    header_size_pt: float | None = None
    footer_text: str = ""
    footer_align: str | None = None
    footer_has_page_number: bool = False
    page_number_format: str | None = None
    page_number_start: int | None = None


@dataclass
class ParagraphInfo:
    index: int
    block_index: int | None
    style_name: str
    text: str
    align: str | None
    page_break_before: bool
    has_page_break_run: bool
    first_line_indent_twips: int | None
    has_drawing: bool
    has_toc_dot_leader: bool
    first_run_east_asia: str | None
    first_run_latin: str | None
    first_run_size_pt: float | None
    first_run_bold: bool | None
    inline_ref_count: int = 0
    inline_ref_not_superscript_count: int = 0


@dataclass
class TableInfo:
    index: int
    block_index: int | None
    rows: int
    cols: int
    has_outer_border: bool
    has_internal_border: bool
    all_rows_cant_split: bool
    rows_missing_cant_split: list[int] = field(default_factory=list)


@dataclass
class InspectResult:
    path: str
    sections: list[SectionInfo] = field(default_factory=list)
    styles: dict[str, StyleInfo] = field(default_factory=dict)
    paragraphs: list[ParagraphInfo] = field(default_factory=list)
    tables: list[TableInfo] = field(default_factory=list)
    non_black_runs: list[RunColorInfo] = field(default_factory=list)
    toc_dot_leader_count: int = 0

    def to_dict(self) -> dict[str, Any]:
        return {
            "path": self.path,
            "sections": [asdict(s) for s in self.sections],
            "styles": {k: asdict(v) for k, v in self.styles.items()},
            "paragraphs": [asdict(p) for p in self.paragraphs],
            "tables": [asdict(t) for t in self.tables],
            "non_black_runs": [asdict(r) for r in self.non_black_runs],
            "toc_dot_leader_count": self.toc_dot_leader_count,
        }


def _emu_to_cm(emu: int | None) -> float:
    if emu is None:
        return 0.0
    return round(emu / EMU_PER_CM, 3)


def _extract_font_from_rpr(rpr) -> tuple[str | None, str | None, float | None, bool | None]:
    if rpr is None:
        return None, None, None, None
    east_asia = None
    latin = None
    size_pt = None
    bold = None
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is not None:
        east_asia = rfonts.get(qn("w:eastAsia"))
        latin = rfonts.get(qn("w:ascii"))
    sz = rpr.find(qn("w:sz"))
    if sz is not None:
        val = sz.get(qn("w:val"))
        if val:
            size_pt = float(val) / 2.0
    b = rpr.find(qn("w:b"))
    if b is not None:
        val = b.get(qn("w:val"))
        bold = val != "0" and val != "false"
    return east_asia, latin, size_pt, bold


def _extract_color_from_rpr(rpr) -> tuple[str | None, str | None]:
    if rpr is None:
        return None, None
    color = rpr.find(qn("w:color"))
    if color is None:
        return None, None
    return color.get(qn("w:val")), color.get(qn("w:themeColor"))


def _extract_alignment(ppr) -> str | None:
    if ppr is None:
        return None
    jc = ppr.find(qn("w:jc"))
    if jc is None:
        return None
    return jc.get(qn("w:val"))


def _extract_line_spacing(ppr) -> float | None:
    if ppr is None:
        return None
    spacing = ppr.find(qn("w:spacing"))
    if spacing is None:
        return None
    line = spacing.get(qn("w:line"))
    line_rule = spacing.get(qn("w:lineRule"))
    if line is None:
        return None
    line_val = float(line)
    if line_rule == "auto" or line_rule is None:
        return line_val / 240.0
    return line_val / 20.0


def _extract_first_line_indent_twips(ppr) -> int | None:
    if ppr is None:
        return None
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        return None
    first_line = ind.get(qn("w:firstLine"))
    if first_line is None:
        return None
    try:
        return int(first_line)
    except ValueError:
        return None


def _extract_has_bottom_border(ppr) -> bool:
    if ppr is None:
        return False
    border = ppr.find(qn("w:pBdr"))
    if border is None:
        return False
    bottom = border.find(qn("w:bottom"))
    if bottom is None:
        return False
    return bottom.get(qn("w:val")) not in (None, "none", "nil")


def _extract_has_toc_dot_leader(ppr) -> bool:
    if ppr is None:
        return False
    tabs = ppr.find(qn("w:tabs"))
    if tabs is None:
        return False
    for tab in tabs.findall(qn("w:tab")):
        if tab.get(qn("w:leader")) == "dot":
            return True
    return False


def _has_drawing(paragraph_element) -> bool:
    return bool(
        paragraph_element.findall(".//" + qn("w:drawing"))
        or paragraph_element.findall(".//" + qn("w:pict"))
    )


def _row_cant_split(row_element) -> bool:
    tr_pr = row_element.find(qn("w:trPr"))
    if tr_pr is None:
        return False
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        return False
    val = cant_split.get(qn("w:val"))
    return val not in ("0", "false")


def _paragraph_text(paragraph_element) -> str:
    return "".join(node.text or "" for node in paragraph_element.findall(".//" + qn("w:t"))).strip()


def _paragraph_has_page_number(paragraph_element) -> bool:
    for fld_simple in paragraph_element.findall(".//" + qn("w:fldSimple")):
        instr = fld_simple.get(qn("w:instr")) or ""
        if "PAGE" in instr.upper():
            return True
    for instr_text in paragraph_element.findall(".//" + qn("w:instrText")):
        if "PAGE" in (instr_text.text or "").upper():
            return True
    text = _paragraph_text(paragraph_element)
    return bool(re.fullmatch(r"\d+|[ivxlcdm]+", text.lower()))


def _first_text_run_size_pt(paragraph) -> float | None:
    for run in paragraph.runs:
        if not (run.text or "").strip():
            continue
        rpr = run._r.find(qn("w:rPr"))
        _ea, _latin, size_pt, _bold = _extract_font_from_rpr(rpr)
        return size_pt
    return None


def _extract_section_page_numbers(sec) -> tuple[str | None, int | None]:
    pg_num_type = sec._sectPr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        return None, None
    fmt = pg_num_type.get(qn("w:fmt"))
    start_raw = pg_num_type.get(qn("w:start"))
    try:
        start = int(start_raw) if start_raw is not None else None
    except ValueError:
        start = None
    return fmt, start


def _extract_block_indices(doc) -> tuple[list[int | None], list[int | None]]:
    paragraph_blocks: list[int | None] = []
    table_blocks: list[int | None] = []
    paragraph_index = 0
    table_index = 0
    for block_index, child in enumerate(doc.element.body.iterchildren()):
        if child.tag == qn("w:p"):
            paragraph_blocks.append(block_index)
            paragraph_index += 1
        elif child.tag == qn("w:tbl"):
            table_blocks.append(block_index)
            table_index += 1
    if paragraph_index != len(doc.paragraphs):
        paragraph_blocks = list(range(len(doc.paragraphs)))
    if table_index != len(doc.tables):
        table_blocks = list(range(len(doc.tables)))
    return paragraph_blocks, table_blocks


def _extract_toc_dot_leader_count_from_xml(path: Path) -> int:
    count = 0
    with ZipFile(path) as archive:
        root = ET.fromstring(archive.read("word/document.xml"))
    for paragraph in root.findall(".//" + W_NS + "p"):
        ppr = paragraph.find(W_NS + "pPr")
        if ppr is None:
            continue
        tabs = ppr.find(W_NS + "tabs")
        if tabs is None:
            continue
        if any(tab.get(W_NS + "leader") == "dot" for tab in tabs.findall(W_NS + "tab")):
            count += 1
    return count


def _extract_styles(doc) -> dict[str, StyleInfo]:
    out: dict[str, StyleInfo] = {}
    for style in doc.styles:
        try:
            element = style.element
        except AttributeError:
            continue
        if element is None:
            continue
        rpr = element.find(qn("w:rPr"))
        ppr = element.find(qn("w:pPr"))
        ea, latin, size_pt, bold = _extract_font_from_rpr(rpr)
        color, theme_color = _extract_color_from_rpr(rpr)
        align = _extract_alignment(ppr)
        line_sp = _extract_line_spacing(ppr)
        first_line_indent = _extract_first_line_indent_twips(ppr)
        try:
            font_size = style.font.size
            if font_size is not None and size_pt is None:
                size_pt = font_size.pt
        except Exception:
            pass
        out[style.name] = StyleInfo(
            name=style.name,
            style_id=getattr(style, "style_id", None),
            east_asia=ea,
            latin=latin,
            size_pt=size_pt,
            bold=bold,
            align=align,
            line_spacing=line_sp,
            first_line_indent_twips=first_line_indent,
            color=color,
            theme_color=theme_color,
        )
    return out


def _is_non_black_color(color: str | None, theme_color: str | None) -> bool:
    if theme_color:
        return True
    if color is None:
        return False
    return color.lower() not in {"000000", "auto"}


def _extract_non_black_runs(path: Path) -> list[RunColorInfo]:
    out: list[RunColorInfo] = []
    with ZipFile(path) as archive:
        xml_files = [
            name for name in archive.namelist()
            if name.startswith("word/") and name.endswith(".xml")
        ]
        for name in xml_files:
            root = ET.fromstring(archive.read(name))
            for run in root.findall(f".//{W_NS}r"):
                text = "".join(node.text or "" for node in run.findall(f".//{W_NS}t")).strip()
                if not text:
                    continue
                rpr = run.find(f"{W_NS}rPr")
                if rpr is None:
                    continue
                color_el = rpr.find(f"{W_NS}color")
                if color_el is None:
                    continue
                color = color_el.get(f"{W_NS}val")
                theme_color = color_el.get(f"{W_NS}themeColor")
                if _is_non_black_color(color, theme_color):
                    out.append(RunColorInfo(
                        source=name,
                        text=text[:80],
                        color=color,
                        theme_color=theme_color,
                    ))
    return out


def _extract_sections(doc) -> list[SectionInfo]:
    sections = []
    for i, sec in enumerate(doc.sections):
        header_text = ""
        header_has_bottom_border = False
        header_align = None
        header_size_pt = None
        try:
            for p in sec.header.paragraphs:
                ppr = p._p.find(qn("w:pPr"))
                if _extract_has_bottom_border(ppr):
                    header_has_bottom_border = True
                if p.text.strip():
                    header_text = p.text.strip()
                    header_align = _extract_alignment(ppr)
                    header_size_pt = _first_text_run_size_pt(p)
                    break
        except Exception:
            pass
        footer_text = ""
        footer_align = None
        footer_has_page_number = False
        try:
            for p in sec.footer.paragraphs:
                if _paragraph_has_page_number(p._p):
                    footer_has_page_number = True
                if p.text.strip() and not footer_text:
                    footer_text = p.text.strip()
                    footer_align = _extract_alignment(p._p.find(qn("w:pPr")))
        except Exception:
            pass
        try:
            linked = sec.header.is_linked_to_previous
        except Exception:
            linked = True
        page_number_format, page_number_start = _extract_section_page_numbers(sec)
        sections.append(SectionInfo(
            index=i,
            page_width_cm=_emu_to_cm(sec.page_width),
            page_height_cm=_emu_to_cm(sec.page_height),
            top_margin_cm=_emu_to_cm(sec.top_margin),
            bottom_margin_cm=_emu_to_cm(sec.bottom_margin),
            left_margin_cm=_emu_to_cm(sec.left_margin),
            right_margin_cm=_emu_to_cm(sec.right_margin),
            header_text=header_text,
            header_linked_to_previous=linked,
            header_has_bottom_border=header_has_bottom_border,
            header_align=header_align,
            header_size_pt=header_size_pt,
            footer_text=footer_text,
            footer_align=footer_align,
            footer_has_page_number=footer_has_page_number,
            page_number_format=page_number_format,
            page_number_start=page_number_start,
        ))
    return sections


def _extract_paragraphs(doc, paragraph_blocks: list[int | None]) -> list[ParagraphInfo]:
    paragraphs = []
    inline_ref_re = re.compile(r"\[\d+\]")
    for i, p in enumerate(doc.paragraphs):
        ppr = p._p.find(qn("w:pPr"))
        align = _extract_alignment(ppr)
        first_line_indent = _extract_first_line_indent_twips(ppr)
        has_toc_dot_leader = _extract_has_toc_dot_leader(ppr)
        page_break_before = False
        if ppr is not None:
            pbb = ppr.find(qn("w:pageBreakBefore"))
            if pbb is not None:
                val = pbb.get(qn("w:val"))
                page_break_before = val != "0" and val != "false"
        has_page_break_run = False
        ea = latin = size_pt = None
        bold = None
        inline_ref_count = 0
        inline_ref_not_superscript_count = 0
        for run in p.runs:
            run_text = run.text or ""
            for br in run._r.findall(qn("w:br")):
                if br.get(qn("w:type")) == "page":
                    has_page_break_run = True
            refs = inline_ref_re.findall(run_text)
            if refs:
                inline_ref_count += len(refs)
                rpr = run._r.find(qn("w:rPr"))
                _ea, _latin, ref_size_pt, _bold = _extract_font_from_rpr(rpr)
                is_superscript = False
                if rpr is not None:
                    vert = rpr.find(qn("w:vertAlign"))
                    if vert is not None and vert.get(qn("w:val")) == "superscript":
                        is_superscript = True
                if ref_size_pt is not None and ref_size_pt < 10:
                    is_superscript = True
                if not is_superscript:
                    inline_ref_not_superscript_count += len(refs)
            if ea is None and run.text.strip():
                rpr = run._r.find(qn("w:rPr"))
                ea, latin, size_pt, bold = _extract_font_from_rpr(rpr)
                if ea is not None or latin is not None or size_pt is not None or bold is not None:
                    break
        paragraphs.append(ParagraphInfo(
            index=i,
            block_index=paragraph_blocks[i] if i < len(paragraph_blocks) else None,
            style_name=p.style.name if p.style else "Normal",
            text=p.text,
            align=align,
            page_break_before=page_break_before,
            has_page_break_run=has_page_break_run,
            first_line_indent_twips=first_line_indent,
            has_drawing=_has_drawing(p._p),
            has_toc_dot_leader=has_toc_dot_leader,
            first_run_east_asia=ea,
            first_run_latin=latin,
            first_run_size_pt=size_pt,
            first_run_bold=bold,
            inline_ref_count=inline_ref_count,
            inline_ref_not_superscript_count=inline_ref_not_superscript_count,
        ))
    return paragraphs


def _extract_tables(doc, table_blocks: list[int | None]) -> list[TableInfo]:
    tables = []
    for i, t in enumerate(doc.tables):
        rows = len(t.rows)
        cols = len(t.columns) if rows > 0 else 0
        tbl_pr = t._tbl.find(qn("w:tblPr"))
        has_outer = False
        has_internal = False
        if tbl_pr is not None:
            borders = tbl_pr.find(qn("w:tblBorders"))
            if borders is not None:
                for tag in ("top", "bottom", "left", "right"):
                    el = borders.find(qn(f"w:{tag}"))
                    if el is not None and el.get(qn("w:val")) not in (None, "none", "nil"):
                        has_outer = True
                for tag in ("insideH", "insideV"):
                    el = borders.find(qn(f"w:{tag}"))
                    if el is not None and el.get(qn("w:val")) not in (None, "none", "nil"):
                        has_internal = True
        rows_missing_cant_split = [
            row_index for row_index, row in enumerate(t._tbl.findall(qn("w:tr")))
            if not _row_cant_split(row)
        ]
        tables.append(TableInfo(
            index=i,
            block_index=table_blocks[i] if i < len(table_blocks) else None,
            rows=rows,
            cols=cols,
            has_outer_border=has_outer, has_internal_border=has_internal,
            all_rows_cant_split=len(rows_missing_cant_split) == 0,
            rows_missing_cant_split=rows_missing_cant_split,
        ))
    return tables


def inspect(path: str | Path) -> InspectResult:
    """Inspect DOCX file and return raw format properties."""
    path = Path(path)
    doc = Document(str(path))
    paragraph_blocks, table_blocks = _extract_block_indices(doc)
    paragraphs = _extract_paragraphs(doc, paragraph_blocks)
    return InspectResult(
        path=str(path),
        sections=_extract_sections(doc),
        styles=_extract_styles(doc),
        paragraphs=paragraphs,
        tables=_extract_tables(doc, table_blocks),
        non_black_runs=_extract_non_black_runs(path),
        toc_dot_leader_count=_extract_toc_dot_leader_count_from_xml(path),
    )
