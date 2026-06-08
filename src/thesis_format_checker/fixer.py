"""Auto-fix engine: apply mechanical corrections to a copy of the DOCX.

Only handles findings with fixable=True. Writes to a NEW file; never overwrites input.
"""

from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from .standard.rules import Finding


CHAPTER_TEXT_RE = re.compile(r"^第\d+章\s+\S+")
SECTION_TEXT_RE = re.compile(r"^\d+\.\d+\s+\S+")
SUBSECTION_TEXT_RE = re.compile(r"^\d+\.\d+\.\d+\s+\S+")


def _set_style_east_asia(style, font_name: str) -> None:
    style.font.name = font_name
    rpr = style.element.find(qn("w:rPr"))
    if rpr is None:
        return
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        from docx.oxml import OxmlElement
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)


def apply_fixes(input_path: str | Path, output_path: str | Path,
                findings: list[Finding], preset: dict) -> tuple[int, list[str]]:
    """Apply fixable findings to input DOCX and write to output path.

    Returns (count_fixed, messages).
    """
    input_path = Path(input_path)
    output_path = Path(output_path)

    if input_path.resolve() == output_path.resolve():
        raise ValueError("output path must differ from input path")

    shutil.copy2(input_path, output_path)
    doc = Document(str(output_path))

    fixed_count = 0
    messages: list[str] = []
    fixed_rule_ids: set[str] = set()

    fixable_ids = {f.rule_id for f in findings if f.fixable}

    if "page-margins" in fixable_ids:
        margins = preset.get("page", {}).get("margin_cm", {})
        for sec in doc.sections:
            if "top" in margins:
                sec.top_margin = Cm(margins["top"])
            if "bottom" in margins:
                sec.bottom_margin = Cm(margins["bottom"])
            if "left" in margins:
                sec.left_margin = Cm(margins["left"])
            if "right" in margins:
                sec.right_margin = Cm(margins["right"])
        messages.append("修正页边距")
        fixed_count += 1
        fixed_rule_ids.add("page-margins")

    if "header-text-match" in fixable_ids:
        expected_header = preset.get("page", {}).get("header_text", "")
        if expected_header:
            for sec in doc.sections:
                paragraphs = list(sec.header.paragraphs)
                if paragraphs:
                    p = paragraphs[0]
                    for run in list(p.runs):
                        run.text = ""
                    p.add_run(expected_header)
                else:
                    sec.header.add_paragraph(expected_header)
            messages.append(f"统一页眉文本为 {expected_header!r}")
            fixed_count += 1
            fixed_rule_ids.add("header-text-match")

    if "body-font-size" in fixable_ids or "body-east-asia-font" in fixable_ids or "body-line-spacing" in fixable_ids:
        body = preset.get("styles", {}).get("body", {})
        normal = doc.styles["Normal"]
        if "size_pt" in body:
            normal.font.size = Pt(body["size_pt"])
        if "east_asia" in body:
            _set_style_east_asia(normal, body["east_asia"])
        if "latin" in body:
            normal.element.find(qn("w:rPr")).find(qn("w:rFonts")).set(qn("w:ascii"), body["latin"])
            normal.element.find(qn("w:rPr")).find(qn("w:rFonts")).set(qn("w:hAnsi"), body["latin"])
        if "line_spacing" in body and abs(body["line_spacing"] - 1.5) < 0.01:
            normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        elif "line_spacing" in body:
            normal.paragraph_format.line_spacing = body["line_spacing"]
        messages.append("修正正文字号/字体/行距 (Normal style)")
        fixed_count += 1

    for level, style_name in [(1, "Heading 1"), (2, "Heading 2"), (3, "Heading 3")]:
        rule_id = f"heading{level}-style"
        if rule_id not in fixable_ids:
            continue
        cfg = preset.get("styles", {}).get(f"heading{level}", {})
        if not cfg:
            continue
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        if "size_pt" in cfg:
            style.font.size = Pt(cfg["size_pt"])
        if "east_asia" in cfg:
            _set_style_east_asia(style, cfg["east_asia"])
        if "bold" in cfg:
            style.font.bold = cfg["bold"]
        messages.append(f"修正 {style_name} 字号/字体")
        fixed_count += 1

    if "heading-style-applied" in fixable_ids:
        reassigned = 0
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue
            target = None
            if CHAPTER_TEXT_RE.match(text):
                target = "Heading 1"
            elif SUBSECTION_TEXT_RE.match(text):
                target = "Heading 3"
            elif SECTION_TEXT_RE.match(text):
                target = "Heading 2"
            if target and p.style.name != target:
                try:
                    p.style = doc.styles[target]
                    reassigned += 1
                except KeyError:
                    pass
        if reassigned:
            messages.append(f"重指派 {reassigned} 个标题段落到正确的 Heading 样式")
            fixed_count += 1

    doc.save(str(output_path))
    return fixed_count, messages
