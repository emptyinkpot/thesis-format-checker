"""Generate the next Liu Gaopeng thesis delivery with format normalization.

Input is the latest versioned delivery in Downloads. This script keeps the
document content chain intact, writes the next vNNN DOCX, then fixes visual
inconsistencies that the school-rule checker does not fully cover:

- first-page cover style copied from the school literature-review template
- style-level colors such as Hyperlink and Pandoc token styles
- direct run colors
- body / heading / caption / TOC base fonts
"""

from __future__ import annotations

import json
import hashlib
import os
import posixpath
import sys
import re
import shutil
import subprocess
from pathlib import Path
from zipfile import ZipFile, ZIP_DEFLATED

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import etree


ROOT = Path(__file__).resolve().parents[1]
DOWNLOADS = Path(r"C:/Users/ASUS-KL/Downloads")
DOCX_STEM = "202213210刘高朋修改迭代版"
OUTPUT_LABEL = "格式规范迭代版"
VERSIONED_DOCX_RE = re.compile(rf"^{re.escape(DOCX_STEM)}_v(\d{{3}})_(.+)\.docx$")
ORIGINAL = DOWNLOADS / "202213210刘高朋修改迭代版.docx"


def versioned_docx_candidates() -> list[tuple[int, str, Path]]:
    candidates: list[tuple[int, str, Path]] = []
    for path in DOWNLOADS.glob(f"{DOCX_STEM}_v*.docx"):
        match = VERSIONED_DOCX_RE.match(path.name)
        if not match:
            continue
        candidates.append((int(match.group(1)), match.group(2), path))
    return sorted(candidates, key=lambda item: item[0])


def latest_source_docx() -> tuple[int, str, Path]:
    candidates = versioned_docx_candidates()
    if candidates:
        return candidates[-1]
    if ORIGINAL.exists():
        return -1, "原始文件", ORIGINAL
    raise FileNotFoundError(f"no source DOCX found in {DOWNLOADS}")


SRC_VERSION, SRC_LABEL, SRC = latest_source_docx()
OUT_VERSION = SRC_VERSION + 1
OUT = DOWNLOADS / f"{DOCX_STEM}_v{OUT_VERSION:03d}_{OUTPUT_LABEL}.docx"
PDF = OUT.with_suffix(".pdf")
REPORT = DOWNLOADS / f"{DOCX_STEM}_v{OUT_VERSION:03d}_格式检测报告.md"
BLANK_REPORT = DOWNLOADS / f"{DOCX_STEM}_v{OUT_VERSION:03d}_留白扫描.json"
PAGE_DIR = DOWNLOADS / f"{DOCX_STEM}_v{OUT_VERSION:03d}_pdf_pages"
VERSION_LOG = DOWNLOADS / f"{DOCX_STEM}_版本记录.md"
EXPECTED_HEADER = "华北水利水电大学毕业设计"
COVER_TEMPLATE = DOWNLOADS / "202213210刘高朋_文献综述_标准模板版.docx"
COVER_TITLE = "毕业设计（论文）"
ENGINEERING_PAPER_READY_ROOT = Path(
    r"E:/My Project/毕业设计论文/论文/我的论文/给老师_工程资料包_2026-05-26/06_论文插图与截图/07_paper_ready"
)
ENGINEERING_RENDER_DIR = DOWNLOADS / f"{DOCX_STEM}_v{OUT_VERSION:03d}_engineering_assets"
ENGINEERING_MIN_IMAGE_HEIGHT_EMU = 1_800_000
KICAD_FULL_SCHEMATIC_SOURCE = ENGINEERING_PAPER_READY_ROOT / "figure_04_kicad_component_schematic.png"
KICAD_LOCAL_FIGURES = [
    (
        "图3.5（a） 电源输入与3.3V稳压实物-原理对应图",
        ENGINEERING_PAPER_READY_ROOT / "figure_04a_power_evidence_pair.png",
        5.8,
    ),
    (
        "图3.5（b） I2C传感器与OLED接口实物-原理对应图",
        ENGINEERING_PAPER_READY_ROOT / "figure_04b_i2c_evidence_pair.png",
        5.8,
    ),
    (
        "图3.5（c） STM32主控最小系统实物-原理对应图",
        ENGINEERING_PAPER_READY_ROOT / "figure_04c_mcu_evidence_pair.png",
        5.8,
    ),
    (
        "图3.5（d） ESP8266无线通信接口实物-原理对应图",
        ENGINEERING_PAPER_READY_ROOT / "figure_04d_esp8266_evidence_pair.png",
        5.8,
    ),
    (
        "图3.5（e） LED与蜂鸣器报警输出实物-原理对应图",
        ENGINEERING_PAPER_READY_ROOT / "figure_04e_alarm_evidence_pair.png",
        5.8,
    ),
    (
        "图3.5（f） 调试、复位、时钟与ADC实物-原理对应图",
        ENGINEERING_PAPER_READY_ROOT / "figure_04f_debug_evidence_pair.png",
        5.8,
    ),
]
KICAD_LOCAL_FIGURE_BY_CAPTION = {
    caption: (source, width) for caption, source, width in KICAD_LOCAL_FIGURES
}
KICAD_STALE_CAPTIONS = {
    "图3.5 供电稳压与调试检查路径图",
    "图3.5 PCB三维装配与模块位置图",
    "图3.5 PCB连线与元件布局图",
    "图3.5 KiCad元器件连线原理图",
    "图3.5 KiCad局部电路截图组",
    "图3.5（a） 电源输入与3.3V稳压局部截图",
    "图3.5（b） I2C传感器与OLED接口局部截图",
    "图3.5（c） STM32主控最小系统局部截图",
    "图3.5（d） ESP8266无线通信接口局部截图",
    "图3.5（e） LED与蜂鸣器报警输出局部截图",
    "图3.5（f） 调试、复位、时钟与ADC局部截图",
    *(caption for caption, _source, _width in KICAD_LOCAL_FIGURES),
}
KICAD_LOCAL_FIGURE_INSERTIONS = [
    (
        "电源系统调试阶段常见问题包括OLED刷新正常但传感器读数波动",
        [
            (
                "图3.5（a） 电源输入与3.3V稳压实物-原理对应图",
                "图3.5（a）左侧显示板上电源输入、稳压器和去耦电容的装配位置，右侧对应5V输入、AMS1117-3.3以及10μF/100nF电容的原理连接，说明3.3V电源先稳定再分配给主控、传感器、显示屏和无线通信模块。",
            ),
        ],
    ),
    (
        "供电设计采用5 V输入并经3.3 V稳压后提供给STM32、SCD41、OLED和ESP8266等模块",
        [
            (
                "图3.5（b） I2C传感器与OLED接口实物-原理对应图",
                "图3.5（b）左侧显示SCD41与OLED模块在板上的接口位置，右侧对应I2C上拉电阻、SCD41四针接口和SSD1306 OLED接口，说明采集总线与显示总线的硬件边界。",
            ),
        ],
    ),
    (
        "本设计以调试稳定和接口清晰为首要目标进行资源分配",
        [
            (
                "图3.5（c） STM32主控最小系统实物-原理对应图",
                "图3.5（c）左侧定位STM32主控及其周边走线，右侧集中展示STM32F103C8T6主控引脚，能够对应PB10/PB11采集、PB6/PB7显示、PA9/PA10通信以及PA1/PB4报警输出等资源分配。",
            ),
            (
                "图3.5（d） ESP8266无线通信接口实物-原理对应图",
                "图3.5（d）左侧定位ESP8266无线模块及相邻供电电容，右侧对应TXD、RXD、RST、CH_PD/EN和3.3V供电端，说明无线模块与STM32串口及复位控制的连接关系。",
            ),
            (
                "图3.5（e） LED与蜂鸣器报警输出实物-原理对应图",
                "图3.5（e）左侧显示蜂鸣器、报警LED和驱动元件在板上的相对位置，右侧对应LED限流电阻、有源蜂鸣器和三极管驱动关系，说明PA1/PB4输出不是简单并联负载，而是分级提示链路的一部分。",
            ),
            (
                "图3.5（f） 调试、复位、时钟与ADC实物-原理对应图",
                "图3.5（f）左侧显示SWD、USART调试接口和ADC预留端在板上的位置，右侧对应复位按键、8MHz晶振、SWD下载口和ADC调试端，说明硬件调试入口与主控稳定运行条件如何支撑后续固件烧录和联调。",
            ),
        ],
    ),
]
KICAD_LOCAL_EXPLANATIONS = {
    explanation
    for _anchor, items in KICAD_LOCAL_FIGURE_INSERTIONS
    for _caption, explanation in items
}
KICAD_STALE_EXPLANATIONS = {
    "图3.5（a）把5V输入、AMS1117-3.3稳压器和10μF/100nF去耦电容放在同一局部中，便于说明3.3V电源怎样先稳定再分配给主控、传感器、显示屏和无线通信模块。",
    "图3.5（b）对应SCD41与OLED的I2C接口分组，图中上拉电阻、SCD41四针接口和SSD1306 OLED接口共同说明采集总线与显示总线的硬件边界。",
    "图3.5（c）集中展示STM32F103C8T6主控引脚，能够直接对应PB10/PB11采集、PB6/PB7显示、PA9/PA10通信以及PA1/PB4报警输出等资源分配。",
    "图3.5（d）单独裁出ESP8266-01S接口，重点显示TXD、RXD、RST、CH_PD/EN和3.3V供电端，便于说明无线模块与STM32串口及复位控制的连接关系。",
    "图3.5（e）对应声光报警输出，LED限流电阻、报警LED、有源蜂鸣器和三极管驱动关系能够说明PA1/PB4输出不是简单并联负载，而是分级提示链路的一部分。",
    "图3.5（f）补充SWD下载接口、复位按键、8MHz晶振和预留ADC调试端，说明硬件调试入口与主控稳定运行条件如何一起支撑后续固件烧录和联调。",
}
REPLACE_EXISTING_ENGINEERING_FIGURES = {
    caption: source
    for caption, source in {
        "图3.2 系统原理图": ENGINEERING_PAPER_READY_ROOT / "figure_01_system_principle_schematic.png",
        "图3.3 STM32主控接口原理图": ENGINEERING_PAPER_READY_ROOT / "figure_02_pre_bluepill_schematic.png",
        "图3.4 PCB顶层布线检查图": ENGINEERING_PAPER_READY_ROOT / "figure_03a_pre_bluepill_top_copper_review.png",
    }.items()
}
ENGINEERING_FIGURES = {
    **REPLACE_EXISTING_ENGINEERING_FIGURES,
    **{caption: source for caption, source, _width in KICAD_LOCAL_FIGURES},
}
ENGINEERING_CAPTION_RENAMES = {
    "图3.2 系统硬件连接与接口关系图": "图3.2 系统原理图",
    "图3.2 系统原理与模块关系图": "图3.2 系统原理图",
    "图3.3 传感、显示与通信接口分配图": "图3.3 STM32主控接口原理图",
    "图3.3 STM32主控接口分配图": "图3.3 STM32主控接口原理图",
    "图3.4 STM32主控引脚分配图": "图3.4 PCB顶层布线检查图",
    "图3.5 供电稳压与调试检查路径图": "图3.5（a） 电源输入与3.3V稳压实物-原理对应图",
    "图3.5 PCB三维装配与模块位置图": "图3.5（a） 电源输入与3.3V稳压实物-原理对应图",
    "图3.5 PCB连线与元件布局图": "图3.5（a） 电源输入与3.3V稳压实物-原理对应图",
    "图3.5 KiCad元器件连线原理图": "图3.5（a） 电源输入与3.3V稳压实物-原理对应图",
    "图3.5 KiCad局部电路截图组": "图3.5（a） 电源输入与3.3V稳压实物-原理对应图",
}
TEXT_REPLACEMENTS = {
    "主控侧资源分配见图3.4和表3.1。": "主控侧资源分配见表3.1，PCB布线检查见图3.4，KiCad局部电路截图见图3.5（a）至图3.5（f）。",
    "无线通信接口关系已经在图3.2和图3.3中体现，软件上传流程见图4.3。": (
        "无线通信接口关系已经在图3.2和图3.3中体现，PCB走线、KiCad局部电路截图和软件上传流程分别见图3.4、图3.5（a）至图3.5（f）和图4.3。"
    ),
    "PCB布线与装配关系见图3.4、图3.5。": "PCB布线检查见图3.4，KiCad局部电路截图见图3.5（a）至图3.5（f）。",
    "PCB布线与模块装配见图3.4、图3.5": (
        "PCB走线见图3.4，KiCad局部电路截图见图3.5（a）至图3.5（f）"
    ),
    "PCB布线与元件布局关系见图3.4、图3.5。": (
        "PCB布线检查见图3.4，KiCad局部电路截图见图3.5（a）至图3.5（f）。"
    ),
    "PCB走线、元件布局见图3.4、图3.5": (
        "PCB走线见图3.4，KiCad局部电路截图见图3.5（a）至图3.5（f）"
    ),
    "KiCad元器件连线原理图见图3.5": (
        "KiCad局部电路截图见图3.5（a）至图3.5（f）"
    ),
    "PCB走线、KiCad元器件连线和软件上传流程分别见图3.4、图3.5和图4.3。": (
        "PCB走线、KiCad局部电路截图和软件上传流程分别见图3.4、图3.5（a）至图3.5（f）和图4.3。"
    ),
    "PCB走线、KiCad局部电路截图和软件上传流程分别见图3.4、图3.5和图4.3。": (
        "PCB走线、KiCad局部电路截图和软件上传流程分别见图3.4、图3.5（a）至图3.5（f）和图4.3。"
    ),
}

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
NS = {"w": W_NS, "r": R_NS, "a": A_NS, "wp": WP_NS, "rel": REL_NS}

FORBIDDEN_TERMS = [
    "代码截图",
    "本实验",
    "实验",
    "本研究",
    "本论文",
    "软件算法",
    "算法",
    "预测",
    "补偿",
    "图3.6",
    "7.2 不足与展望",
    "2.2.1 系统需求分析",
    "文献启发",
    "本文按照",
]

ALLOWED_BLANK_PAGES = {2, 3, 23, 36, 55, 56, 72, 79}
ALLOWED_EAST_ASIA_FONTS = {"宋体", "黑体", "楷体", "仿宋_GB2312", "隶书", "Consolas"}
ALLOWED_LATIN_FONTS = {"Times New Roman", "Consolas", "宋体"}

COVER_TABLE_ROWS = [
    ("学    院", "电子工程学院"),
    ("专    业", "电子信息工程"),
    ("姓    名", "刘高朋"),
    ("学    号", "202213210"),
    ("指导教师", "张晓华"),
    ("完成时间", "2026年6月"),
]


BRIDGE_PARAGRAPHS = {
    "从实现角度看，SCD41负责CO2": [
        "因此，第2章在结构上承担了承前启后的作用：它把第1章提出的监测需求转化为可以落地的模块分工，也为后续硬件接口、程序流程、功能页面和测试项目提供了同一套解释框架。这样安排后，后文各章不再是分散介绍，而是围绕同一条数据链路逐步展开。"
    ],
    "图3.1按“左侧实物、右侧说明”的方式介绍主要硬件模块": [
        "这张图不是单独展示器件外观，而是作为第3章的模块索引：左侧实物帮助读者先识别SCD41、STM32、OLED、ESP8266和报警器件，右侧说明再对应到接口连接、原理图、PCB、固件源码和后续测试项目。这样处理后，实物、原理图、接口、PCB、源码和测试能够落在同一条说明链路上。",
        "图3.2至图3.5继续把该索引落到工程图纸上。图3.2和图3.3采用已经整理好的系统连线图和主控接口分配截图，分别说明系统供电、传感、显示、通信、报警、调试接口以及STM32主控引脚分配；图3.4采用PCB走线检查截图说明顶层布线、焊盘和排针位置，图3.5（a）至图3.5（f）不再缩放整张KiCad原理图，而是分别裁出电源稳压、I2C传感与显示、STM32主控、ESP8266无线通信、报警输出以及调试复位时钟ADC六个局部板块，逐块说明元器件符号、网络连接和接口分组。ERC/DRC检查记录、固件文件main.c、scd41.c、bsp_pins.h以及第6章测试验证一起构成后续说明依据。"
    ],
    "传感器采集功能对应硬件SCD41和软件scd41.c": [
        "图5.2用于说明采集代码和显示刷新之间的衔接关系。读图时应先看SCD41读取结果怎样进入主程序状态，再看OLED刷新和报警判断怎样使用同一组有效采样值。"
    ],
    "该关键代码体现了传感器数据从底层驱动到界面刷新的基本链路": [
        "图5.3进一步展开状态字段的写入方式，重点不是重复展示代码，而是说明co2、temp、hum、alarm、seq、net等字段如何从采样结果转换为显示和上传所需的数据结构。"
    ],
    "报警状态机采用分级阈值与回差共同控制": [
        "图5.4用于对应阈值判断和蜂鸣器输出代码。读图时需要关注一级预警、二级报警和恢复回差三个判断点，以及这些判断怎样同步到LED、蜂鸣器、OLED和上传字段。"
    ],
    "调试时先检查ESP8266是否能响应AT命令和连接路由器": [
        "图5.5用于说明无线通信代码的检查顺序。读图时应先看AT命令发送和应答匹配，再看超时判断如何区分入网失败、连接失败和数据发送失败。"
    ],
    "主循环以timeCount作为节拍变量。": [
        "在实际运行中，主循环并不是把所有任务同时处理，而是按“采集—判断—显示—上传”的顺序逐项完成。这样可以保证CO2数据先经过校验和报警等级计算，再进入OLED显示和远程数据帧；如果Wi-Fi或MQTT暂时异常，程序只记录网络状态，不回退采集和本地预警流程。",
        "这种调度方式也便于后续测试定位问题。采集异常时可优先检查SCD41总线和CRC校验，报警异常时可检查CalcAlarmLevel和AlarmOutput，上传异常时再检查ESP8266和OneNET连接。图4.1正是按照这个排查顺序组织程序流程。"
    ],
    "从论文结构看，第6章的测试内容与前文设计内容逐项对应": [
        "因此，第6章的意义不只是列出测试结果，而是用测试记录把前文的设计链路重新闭合。采集、报警、显示、上传和移动端查看分别对应不同测试项目，任何一项异常都能追溯到具体硬件接口或程序模块，这也为后续优化提供了明确入口。"
    ],
    "图5.12 手机端维护管理页面截图": [
        "从图5.7至图5.12可以看出，移动端页面并非单独展示界面效果，而是围绕实时值、历史趋势、设备状态和维护入口组织。登录页解决访问入口，首页承担快速判断，数据分析页用于追踪变化，设备、场景和维护页面则为后续管理留下扩展位置。这样处理后，截图与前文的服务端接口和设备上传字段能够对应起来。"
    ],
    "由表6.4可知，样机读数与参考仪器变化趋势一致": [
        "该结果说明，样机在毕业设计阶段更适合用于趋势监测和阈值提示，而不是作为计量级检测仪器。后续若要用于长期部署，还需要结合标准气体或校准设备进行周期校准，并在不同温湿度条件下重新记录误差范围。"
    ],
}

SUPERSEDED_PARAGRAPH_PREFIXES = [
    "这张图不是单独展示器件外观，而是作为第3章的模块索引：左侧实物帮助读者先识别",
    "图3.2和图3.3继续把该索引落到工程图纸上。",
    "图3.2至图3.5继续把该索引落到工程图纸上。",
    "图3.5采用局部截图方式展示KiCad原理图中各功能板块。",
    "图3.5（a）至图3.5（f）分别对应电源、I2C接口、主控、无线通信、报警输出和调试接口。",
    "从图3.5（a）至图3.5（f）可以看出",
]


def normalize_text(text: str) -> str:
    return re.sub(r"\s+", "", text or "")


def document_text(path: Path) -> str:
    doc = Document(str(path))
    parts: list[str] = [paragraph.text for paragraph in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(parts)


def iter_direct_run_fonts(path: Path):
    with ZipFile(path) as archive:
        for name in archive.namelist():
            if not name.startswith("word/") or not name.endswith(".xml"):
                continue
            root = etree.fromstring(archive.read(name))
            for run in root.xpath(".//w:r", namespaces=NS):
                text = "".join(t.text or "" for t in run.xpath(".//w:t", namespaces=NS)).strip()
                if not text:
                    continue
                rpr = run.find("w:rPr", namespaces=NS)
                if rpr is None:
                    continue
                fonts = rpr.find("w:rFonts", namespaces=NS)
                if fonts is None:
                    continue
                yield {
                    "source": name,
                    "text": text[:80],
                    "east_asia": fonts.get(f"{{{W_NS}}}eastAsia"),
                    "ascii": fonts.get(f"{{{W_NS}}}ascii"),
                    "hansi": fonts.get(f"{{{W_NS}}}hAnsi"),
                }


def ensure_child(parent, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def set_xml_color_black(rpr) -> None:
    color = ensure_child(rpr, "w:color")
    color.set(qn("w:val"), "000000")
    for attr in ("w:themeColor", "w:themeTint", "w:themeShade"):
        color.attrib.pop(qn(attr), None)


def set_xml_fonts(rpr, east_asia: str, latin: str) -> None:
    fonts = ensure_child(rpr, "w:rFonts")
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:ascii"), latin)
    fonts.set(qn("w:hAnsi"), latin)
    fonts.set(qn("w:cs"), latin)


def set_xml_size(rpr, size_pt: float) -> None:
    half_points = str(int(round(size_pt * 2)))
    for tag in ("w:sz", "w:szCs"):
        elem = ensure_child(rpr, tag)
        elem.set(qn("w:val"), half_points)


def remove_xml_italic(rpr) -> None:
    for tag in ("w:i", "w:iCs"):
        elem = rpr.find(qn(tag))
        if elem is not None:
            rpr.remove(elem)


def set_run_visual(run, east_asia: str, latin: str, size_pt: float | None, bold: bool | None = None) -> None:
    rpr = run._element.get_or_add_rPr()
    set_xml_color_black(rpr)
    set_xml_fonts(rpr, east_asia, latin)
    if size_pt is not None:
        set_xml_size(rpr, size_pt)
    if bold is not None:
        run.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def is_inline_citation_run(run) -> bool:
    text = run.text or ""
    if not re.fullmatch(r"\s*(\[\d+\]\s*)+", text):
        return False
    rpr = run._element.find(qn("w:rPr"))
    if rpr is None:
        return False
    vert = rpr.find(qn("w:vertAlign"))
    return vert is not None and vert.get(qn("w:val")) == "superscript"


def paragraph_kind(paragraph) -> tuple[str, str, float, bool | None]:
    style = paragraph.style.name if paragraph.style is not None else ""
    text = paragraph.text.strip()
    first_font = paragraph.runs[0].font.name if paragraph.runs else ""
    code_like = (
        style in {"Code", "Source Code", "HTML Code", "Listing"}
        or first_font in {"Consolas", "Courier New", "Courier"}
        or bool(re.match(r"^\d{4}\s+", text))
    )
    if code_like:
        return "Consolas", "Consolas", 10.5, None
    if re.match(r"^\[\d+\]", text):
        return "宋体", "Times New Roman", 10.5, None
    if style == "Heading 1":
        return "黑体", "Times New Roman", 16, None
    if style == "Heading 2":
        return "黑体", "Times New Roman", 14, None
    if style == "Heading 3":
        return "黑体", "Times New Roman", 12, None
    if style.lower().startswith("toc"):
        return "宋体", "Times New Roman", 12, None
    if style == "caption" or text.startswith(("图", "表")):
        return "宋体", "Times New Roman", 10.5, None
    return "宋体", "Times New Roman", 12, None


def is_caption_paragraph(paragraph) -> bool:
    return paragraph.text.strip().startswith(("图", "表")) or (
        paragraph.style is not None and paragraph.style.name.lower() == "caption"
    )


def tighten_reference_paragraph(paragraph) -> None:
    text = paragraph.text.strip()
    if not re.match(r"^\[\d+\]", text):
        return
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 0.92
    fmt.first_line_indent = None


def remove_superseded_paragraphs(doc: Document) -> None:
    generated_prefixes = [
        *SUPERSEDED_PARAGRAPH_PREFIXES,
        *(text for texts in BRIDGE_PARAGRAPHS.values() for text in texts),
    ]
    normalized_prefixes = [normalize_text(prefix) for prefix in generated_prefixes]
    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        normalized = normalize_text(text)
        if any(normalized.startswith(prefix) for prefix in normalized_prefixes):
            paragraph._element.getparent().remove(paragraph._element)


def make_body_paragraph(doc: Document, text: str) -> OxmlElement:
    para = OxmlElement("w:p")
    ppr = OxmlElement("w:pPr")
    if "Body Text" in doc.styles:
        pstyle = OxmlElement("w:pStyle")
        pstyle.set(qn("w:val"), doc.styles["Body Text"].style_id)
        ppr.append(pstyle)
    para.append(ppr)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    set_xml_color_black(rpr)
    set_xml_fonts(rpr, "宋体", "Times New Roman")
    set_xml_size(rpr, 12)
    run.append(rpr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run.append(t)
    para.append(run)
    return para


def paragraph_contains_after(paragraph, prefix: str) -> bool:
    text = paragraph.text.strip()
    return text.startswith(prefix) or normalize_text(text).startswith(normalize_text(prefix))


def insert_bridge_paragraphs(doc: Document) -> None:
    for anchor, texts in BRIDGE_PARAGRAPHS.items():
        for paragraph in doc.paragraphs:
            if paragraph_contains_after(paragraph, anchor):
                existing_after = "\n".join(p.text for p in doc.paragraphs)
                if all(text in existing_after for text in texts):
                    break
                current = paragraph._element
                for text in texts:
                    new_para = make_body_paragraph(doc, text)
                    current.addnext(new_para)
                    current = new_para
                break
        else:
            raise RuntimeError(f"找不到插入锚点: {anchor}")


def replace_paragraph_text(paragraph, text: str):
    for run in list(paragraph.runs):
        paragraph._element.remove(run._element)
    return paragraph.add_run(text)


def rename_engineering_captions(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        replacement = ENGINEERING_CAPTION_RENAMES.get(text)
        if replacement:
            run = replace_paragraph_text(paragraph, replacement)
            set_run_visual(run, "宋体", "Times New Roman", 10.5, False)
            run.italic = False
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def apply_text_replacements(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        text = paragraph.text
        patched = text
        for old, new in TEXT_REPLACEMENTS.items():
            patched = patched.replace(old, new)
        if patched != text:
            run = replace_paragraph_text(paragraph, patched)
            east_asia, latin, size_pt, bold = paragraph_kind(paragraph)
            set_run_visual(run, east_asia, latin, size_pt, bold)


def normalize_caption_run(run) -> None:
    run.italic = False
    run.font.italic = False
    remove_xml_italic(run._element.get_or_add_rPr())


def normalize_caption_style(doc: Document) -> None:
    for style_name in ("Caption", "caption"):
        if style_name not in doc.styles:
            continue
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(10.5)
        style.font.bold = False
        style.font.italic = False
        style.font.color.rgb = RGBColor(0, 0, 0)
        rpr = style.element.get_or_add_rPr()
        set_xml_fonts(rpr, "宋体", "Times New Roman")
        set_xml_size(rpr, 10.5)
        remove_xml_italic(rpr)


def _paragraph_has_drawing(paragraph) -> bool:
    return bool(paragraph._element.xpath(".//w:drawing"))


def _remove_paragraph(paragraph) -> None:
    parent = paragraph._element.getparent()
    if parent is not None:
        parent.remove(paragraph._element)


def remove_stale_kicad_figure_blocks(doc: Document) -> None:
    """Remove old whole-schematic 图3.5 blocks before inserting local crops."""
    paragraphs = list(doc.paragraphs)
    for index, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        if text in KICAD_LOCAL_EXPLANATIONS or text in KICAD_STALE_EXPLANATIONS:
            _remove_paragraph(paragraph)
            continue
        if text not in KICAD_STALE_CAPTIONS:
            continue

        for previous in reversed(paragraphs[max(0, index - 4):index]):
            if _paragraph_has_drawing(previous):
                _remove_paragraph(previous)
                break
        _remove_paragraph(paragraph)


def add_body_paragraph_after(doc: Document, anchor, text: str):
    paragraph = doc.add_paragraph()
    if "Body Text" in doc.styles:
        paragraph.style = doc.styles["Body Text"]
    run = paragraph.add_run(text)
    east_asia, latin, size_pt, bold = paragraph_kind(paragraph)
    set_run_visual(run, east_asia, latin, size_pt, bold)
    anchor._element.addnext(paragraph._element)
    return paragraph


def move_last_paragraph_after(doc: Document, anchor):
    paragraph = doc.paragraphs[-1]
    anchor._element.addnext(paragraph._element)
    return paragraph


def add_picture_after(doc: Document, anchor, image_path: Path, width_inches: float):
    if not image_path.exists():
        raise FileNotFoundError(image_path)
    picture_paragraph = doc.add_paragraph()
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_paragraph.add_run().add_picture(str(image_path), width=Inches(width_inches))
    moved = move_last_paragraph_after(doc, anchor)
    return moved


def add_caption_after(doc: Document, anchor, caption: str):
    caption_paragraph = doc.add_paragraph()
    if "Caption" in doc.styles:
        caption_paragraph.style = doc.styles["Caption"]
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption_paragraph.add_run(caption)
    set_run_visual(run, "宋体", "Times New Roman", 10.5, False)
    normalize_caption_run(run)
    moved = move_last_paragraph_after(doc, anchor)
    return moved


def insert_kicad_local_figures(doc: Document) -> None:
    remove_stale_kicad_figure_blocks(doc)
    inserted: set[str] = set()

    for anchor_prefix, items in KICAD_LOCAL_FIGURE_INSERTIONS:
        anchor = next((p for p in doc.paragraphs if paragraph_contains_after(p, anchor_prefix)), None)
        if anchor is None:
            raise RuntimeError(f"找不到图3.5局部截图插入锚点: {anchor_prefix}")

        current = anchor
        for caption, explanation in items:
            source, width = KICAD_LOCAL_FIGURE_BY_CAPTION[caption]
            picture = add_picture_after(doc, current, source, width)
            caption_paragraph = add_caption_after(doc, picture, caption)
            current = add_body_paragraph_after(doc, caption_paragraph, explanation)
            inserted.add(caption)

    expected = set(KICAD_LOCAL_FIGURE_BY_CAPTION)
    if inserted != expected:
        raise RuntimeError(f"图3.5局部截图插入不完整: inserted={sorted(inserted)}, expected={sorted(expected)}")


def _image_size(path: Path) -> tuple[int, int]:
    from PIL import Image

    with Image.open(path) as image:
        return image.size


def magick_command() -> str:
    command = shutil.which("magick") or r"C:\Program Files\ImageMagick-7.1.2-Q16-HDRI\magick.exe"
    if not Path(command).exists() and shutil.which(command) is None:
        raise RuntimeError("ImageMagick magick command is required to render SVG engineering figures")
    return command


def render_engineering_source(source: Path, target: Path) -> Path:
    if source.suffix.lower() == ".svg":
        subprocess.run(
            [
                magick_command(),
                "-density",
                "220",
                "-background",
                "white",
                str(source),
                "-alpha",
                "remove",
                "-alpha",
                "off",
                "-strip",
                str(target),
            ],
            check=True,
        )
    else:
        shutil.copy2(source, target)
    return target


def prepare_engineering_figure_assets(figures: dict[str, Path] | None = None) -> dict[str, Path]:
    ENGINEERING_RENDER_DIR.mkdir(parents=True, exist_ok=True)
    rendered: dict[str, Path] = {}
    selected_figures = figures or ENGINEERING_FIGURES
    for index, (caption, source) in enumerate(selected_figures.items(), start=1):
        if not source.exists():
            raise FileNotFoundError(f"engineering figure missing for {caption}: {source}")
        target = ENGINEERING_RENDER_DIR / f"figure_{index:02d}_{source.stem}.png"
        render_engineering_source(source, target)
        rendered[caption] = target
    return rendered


def _target_zip_name(target: str) -> str:
    if target.startswith("/"):
        return posixpath.normpath(target.lstrip("/"))
    return posixpath.normpath(posixpath.join("word", target))


def _caption_text(paragraph) -> str:
    return "".join(paragraph.xpath(".//w:t/text()", namespaces=NS)).strip()


def _caption_previous_blips(document_root) -> dict[str, list[tuple[str, object]]]:
    body = document_root.find("w:body", namespaces=NS)
    if body is None:
        return {}
    paragraphs = [child for child in body if child.tag == f"{{{W_NS}}}p"]
    out: dict[str, list[tuple[str, object]]] = {}
    for index, paragraph in enumerate(paragraphs):
        caption = _caption_text(paragraph)
        if caption not in ENGINEERING_FIGURES:
            continue
        for prev in reversed(paragraphs[max(0, index - 4):index]):
            blips = []
            for blip in prev.xpath(".//a:blip", namespaces=NS):
                rid = blip.get(f"{{{R_NS}}}embed")
                if rid:
                    blips.append((rid, blip))
            if blips:
                out[caption] = blips
                break
    return out


def _set_inline_image_aspect(blip, image_path: Path) -> None:
    inline = blip
    while inline is not None and inline.tag != f"{{{WP_NS}}}inline":
        inline = inline.getparent()
    if inline is None:
        return
    extent = inline.find("wp:extent", namespaces=NS)
    if extent is None:
        return
    cx = int(extent.get("cx") or "0")
    width_px, height_px = _image_size(image_path)
    if cx <= 0 or width_px <= 0:
        return
    cy = int(round(cx * height_px / width_px))
    extent.set("cy", str(cy))
    pic_extents = inline.xpath(".//a:xfrm/a:ext", namespaces=NS)
    for pic_extent in pic_extents:
        pic_extent.set("cx", str(cx))
        pic_extent.set("cy", str(cy))


def _inline_extent_for_blip(blip) -> tuple[int, int] | None:
    inline = blip
    while inline is not None and inline.tag != f"{{{WP_NS}}}inline":
        inline = inline.getparent()
    if inline is None:
        return None
    extent = inline.find("wp:extent", namespaces=NS)
    if extent is None:
        return None
    return int(extent.get("cx") or "0"), int(extent.get("cy") or "0")


def replace_engineering_figure_media() -> None:
    figure_assets = prepare_engineering_figure_assets(REPLACE_EXISTING_ENGINEERING_FIGURES)

    tmp = OUT.with_suffix(".figures.tmp.docx")
    replaced: set[str] = set()
    with ZipFile(OUT, "r") as zin:
        document_root = etree.fromstring(zin.read("word/document.xml"))
        rels_root = etree.fromstring(zin.read("word/_rels/document.xml.rels"))
        rid_to_target = {
            rel.get("Id"): rel.get("Target")
            for rel in rels_root.xpath(".//rel:Relationship", namespaces=NS)
        }
        caption_blips = _caption_previous_blips(document_root)
        media_replacements: dict[str, bytes] = {}
        for caption, figure in figure_assets.items():
            blips = caption_blips.get(caption)
            if not blips:
                raise RuntimeError(f"cannot locate image before caption: {caption}")
            rid, blip = blips[0]
            target = rid_to_target.get(rid)
            if not target:
                raise RuntimeError(f"cannot resolve image relationship {rid} for {caption}")
            media_replacements[_target_zip_name(target)] = figure.read_bytes()
            _set_inline_image_aspect(blip, figure)
            replaced.add(caption)

        if replaced != set(REPLACE_EXISTING_ENGINEERING_FIGURES):
            raise RuntimeError(f"not all engineering figures replaced: {sorted(replaced)}")

        patched_document = etree.tostring(document_root, xml_declaration=True, encoding="UTF-8", standalone=True)
        with ZipFile(tmp, "w", ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/document.xml":
                    data = patched_document
                elif item.filename in media_replacements:
                    data = media_replacements[item.filename]
                zout.writestr(item, data)
    tmp.replace(OUT)


def _relationship_target_zip_name(rels_name: str, target: str) -> str:
    if target.startswith("/") or "://" in target:
        return _target_zip_name(target)
    source_dir = posixpath.dirname(posixpath.dirname(rels_name))
    return posixpath.normpath(posixpath.join(source_dir, target))


def remove_unused_document_media() -> None:
    """Drop image relationships/files that were left behind after deleting stale figure paragraphs."""
    tmp = OUT.with_suffix(".media-clean.tmp.docx")
    with ZipFile(OUT, "r") as zin:
        document_root = etree.fromstring(zin.read("word/document.xml"))
        rels_root = etree.fromstring(zin.read("word/_rels/document.xml.rels"))
        used_rids = {
            blip.get(f"{{{R_NS}}}embed")
            for blip in document_root.xpath(".//a:blip", namespaces=NS)
            if blip.get(f"{{{R_NS}}}embed")
        }
        removed_targets: set[str] = set()
        changed = False
        for rel in list(rels_root.xpath(".//rel:Relationship", namespaces=NS)):
            rel_id = rel.get("Id")
            rel_type = rel.get("Type") or ""
            target = rel.get("Target") or ""
            if "/image" not in rel_type or rel_id in used_rids:
                continue
            removed_targets.add(_relationship_target_zip_name("word/_rels/document.xml.rels", target))
            rel.getparent().remove(rel)
            changed = True

        if not changed:
            return

        kept_targets: set[str] = set()
        rels_payloads: dict[str, bytes] = {
            "word/_rels/document.xml.rels": etree.tostring(
                rels_root,
                xml_declaration=True,
                encoding="UTF-8",
                standalone=True,
            )
        }
        for item in zin.infolist():
            if not item.filename.endswith(".rels") or item.filename == "word/_rels/document.xml.rels":
                continue
            root = etree.fromstring(zin.read(item.filename))
            for rel in root.xpath(".//rel:Relationship", namespaces=NS):
                rel_type = rel.get("Type") or ""
                target = rel.get("Target") or ""
                if "/image" in rel_type and target:
                    kept_targets.add(_relationship_target_zip_name(item.filename, target))
        for rel in rels_root.xpath(".//rel:Relationship", namespaces=NS):
            rel_type = rel.get("Type") or ""
            target = rel.get("Target") or ""
            if "/image" in rel_type and target:
                kept_targets.add(_relationship_target_zip_name("word/_rels/document.xml.rels", target))

        removable_targets = removed_targets - kept_targets
        with ZipFile(tmp, "w", ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                if item.filename in removable_targets:
                    continue
                data = rels_payloads.get(item.filename) or zin.read(item.filename)
                zout.writestr(item, data)
    tmp.replace(OUT)


def _engineering_figure_media_targets(docx_path: Path) -> dict[str, str]:
    with ZipFile(docx_path, "r") as archive:
        document_root = etree.fromstring(archive.read("word/document.xml"))
        rels_root = etree.fromstring(archive.read("word/_rels/document.xml.rels"))
        rid_to_target = {
            rel.get("Id"): rel.get("Target")
            for rel in rels_root.xpath(".//rel:Relationship", namespaces=NS)
        }
        caption_blips = _caption_previous_blips(document_root)

    targets: dict[str, str] = {}
    for caption in ENGINEERING_FIGURES:
        blips = caption_blips.get(caption)
        if not blips:
            raise RuntimeError(f"cannot locate image before caption: {caption}")
        rid, _blip = blips[0]
        target = rid_to_target.get(rid)
        if not target:
            raise RuntimeError(f"cannot resolve image relationship {rid} for {caption}")
        targets[caption] = _target_zip_name(target)
    return targets


def _engineering_figure_display_extents(docx_path: Path) -> dict[str, tuple[int, int]]:
    with ZipFile(docx_path, "r") as archive:
        document_root = etree.fromstring(archive.read("word/document.xml"))
        caption_blips = _caption_previous_blips(document_root)

    extents: dict[str, tuple[int, int]] = {}
    for caption in ENGINEERING_FIGURES:
        blips = caption_blips.get(caption)
        if not blips:
            raise RuntimeError(f"cannot locate image before caption: {caption}")
        _rid, blip = blips[0]
        extent = _inline_extent_for_blip(blip)
        if extent is None:
            raise RuntimeError(f"cannot locate display extent for engineering figure: {caption}")
        extents[caption] = extent
    return extents


def _sha256(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def verify_engineering_figure_assets(docx_path: Path | None = None) -> dict[str, str]:
    """Verify that required engineering captions point to the real source assets."""
    target_docx = docx_path or OUT
    figure_assets = prepare_engineering_figure_assets()
    full_schematic_hash = _sha256(KICAD_FULL_SCHEMATIC_SOURCE.read_bytes())

    media_targets = _engineering_figure_media_targets(target_docx)
    display_extents = _engineering_figure_display_extents(target_docx)
    with ZipFile(target_docx, "r") as archive:
        embedded_media_hashes = {
            _sha256(archive.read(name))
            for name in archive.namelist()
            if name.startswith("word/media/")
        }
        if full_schematic_hash in embedded_media_hashes:
            raise RuntimeError("full KiCad schematic is still embedded; expected localized 图3.5 crop screenshots only")
        for caption, figure in figure_assets.items():
            media_name = media_targets.get(caption)
            if not media_name:
                raise RuntimeError(f"engineering figure media target missing for {caption}")
            actual_hash = _sha256(archive.read(media_name))
            expected_hash = _sha256(figure.read_bytes())
            if actual_hash != expected_hash:
                raise RuntimeError(
                    f"engineering figure mismatch for {caption}: "
                    f"{media_name} sha256={actual_hash}, expected={expected_hash}"
                )
            _cx, cy = display_extents[caption]
            if cy < ENGINEERING_MIN_IMAGE_HEIGHT_EMU:
                raise RuntimeError(
                    f"engineering figure is too small on page for {caption}: "
                    f"height_emu={cy}, expected>={ENGINEERING_MIN_IMAGE_HEIGHT_EMU}"
                )
    if "图3.5 KiCad元器件连线原理图" in document_text(target_docx):
        raise RuntimeError("stale whole-schematic 图3.5 caption remains in DOCX")
    return media_targets


def set_cover_paragraph(paragraph, text: str, alignment, east_asia: str, latin: str, size_pt: float, bold: bool | None) -> None:
    run = replace_paragraph_text(paragraph, text)
    paragraph.alignment = alignment
    set_run_visual(run, east_asia, latin, size_pt, bold)


def restore_cover_template_style(doc: Document) -> None:
    """Restore the first-page cover contract from the school template.

    The body normalizer intentionally touches all visible text, so the cover is
    repaired afterwards as a separate front-matter layer. The reference is the
    literature-review cover template; only the middle title changes to the
    main-thesis title.
    """
    if not COVER_TEMPLATE.exists():
        raise FileNotFoundError(COVER_TEMPLATE)
    if len(doc.paragraphs) < 10 or not doc.tables:
        raise RuntimeError("document does not have the expected cover structure")

    set_cover_paragraph(doc.paragraphs[0], "存档编号                 ", WD_ALIGN_PARAGRAPH.RIGHT, "宋体", "Times New Roman", 12, True)
    set_cover_paragraph(doc.paragraphs[1], "华北水利水电大学", WD_ALIGN_PARAGRAPH.CENTER, "隶书", "Times New Roman", 42, True)
    set_cover_paragraph(
        doc.paragraphs[2],
        "North China University of Water Resources and Electric Power",
        WD_ALIGN_PARAGRAPH.CENTER,
        "宋体",
        "Times New Roman",
        12,
        None,
    )
    set_cover_paragraph(doc.paragraphs[3], COVER_TITLE, WD_ALIGN_PARAGRAPH.CENTER, "楷体", "Times New Roman", 30, True)
    set_cover_paragraph(
        doc.paragraphs[4],
        "题目    基于嵌入式的二氧化碳监测与预警器设计",
        None,
        "宋体",
        "Times New Roman",
        18,
        True,
    )
    set_cover_paragraph(doc.paragraphs[9], "教务处制", WD_ALIGN_PARAGRAPH.CENTER, "宋体", "Times New Roman", 14, None)

    table = doc.tables[0]
    if len(table.rows) < len(COVER_TABLE_ROWS):
        raise RuntimeError("cover table is missing expected rows")
    for row_index, (label, value) in enumerate(COVER_TABLE_ROWS):
        row = table.rows[row_index]
        row.cells[0].text = label
        row.cells[1].text = value
        for cell_index, cell in enumerate(row.cells[:2]):
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if cell_index == 0 else WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    if run.text:
                        set_run_visual(run, "宋体", "Times New Roman", 14, True)


def normalize_docx_with_python_docx() -> None:
    shutil.copy2(SRC, OUT)
    doc = Document(str(OUT))

    remove_superseded_paragraphs(doc)
    insert_bridge_paragraphs(doc)
    rename_engineering_captions(doc)
    apply_text_replacements(doc)
    insert_kicad_local_figures(doc)
    normalize_caption_style(doc)

    for style in doc.styles:
        if not hasattr(style, "font"):
            continue
        try:
            style.font.color.rgb = RGBColor(0, 0, 0)
            if style.name == "Hyperlink":
                style.font.underline = False
        except Exception:
            pass

    for section in doc.sections:
        for container in (section.header, section.footer):
            for paragraph in container.paragraphs:
                for run in paragraph.runs:
                    set_run_visual(run, "宋体", "Times New Roman", 9)

    for paragraph in doc.paragraphs:
        east_asia, latin, size_pt, bold = paragraph_kind(paragraph)
        if paragraph.style.name.lower().startswith("toc"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        tighten_reference_paragraph(paragraph)
        for run in paragraph.runs:
            if not run.text:
                continue
            if is_inline_citation_run(run):
                set_run_visual(run, "宋体", "Times New Roman", 9, None)
                vert = ensure_child(run._element.get_or_add_rPr(), "w:vertAlign")
                vert.set(qn("w:val"), "superscript")
                continue
            set_run_visual(run, east_asia, latin, size_pt, bold)
            if is_caption_paragraph(paragraph):
                normalize_caption_run(run)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    east_asia, latin, size_pt, bold = paragraph_kind(paragraph)
                    for run in paragraph.runs:
                        if run.text:
                            set_run_visual(run, east_asia, latin, size_pt, bold)
                            if is_caption_paragraph(paragraph):
                                normalize_caption_run(run)

    restore_cover_template_style(doc)
    doc.save(str(OUT))
    remove_unused_document_media()
    replace_engineering_figure_media()


def patch_xml_colors_and_styles() -> None:
    tmp = OUT.with_suffix(".tmp.docx")
    with ZipFile(OUT, "r") as zin, ZipFile(tmp, "w", ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename.startswith("word/") and item.filename.endswith(".xml"):
                root = etree.fromstring(data)
                changed = False
                for color in root.xpath(".//w:color", namespaces=NS):
                    color.set(f"{{{W_NS}}}val", "000000")
                    for attr in ("themeColor", "themeTint", "themeShade"):
                        color.attrib.pop(f"{{{W_NS}}}{attr}", None)
                    changed = True
                for highlight in root.xpath(".//w:highlight", namespaces=NS):
                    val = highlight.get(f"{{{W_NS}}}val")
                    if val and val.lower() not in {"none", "clear"}:
                        highlight.set(f"{{{W_NS}}}val", "none")
                        changed = True
                if item.filename == "word/styles.xml":
                    changed = patch_styles_xml(root) or changed
                if changed:
                    data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
            zout.writestr(item, data)
    tmp.replace(OUT)


def patch_styles_xml(root) -> bool:
    changed = False
    for style in root.xpath(".//w:style", namespaces=NS):
        sid = style.get(f"{{{W_NS}}}styleId") or ""
        name_elem = style.find("w:name", namespaces=NS)
        name = name_elem.get(f"{{{W_NS}}}val") if name_elem is not None else sid
        rpr = style.find("w:rPr", namespaces=NS)
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            style.append(rpr)
        set_xml_color_black(rpr)
        lower_name = (name or "").lower()
        if sid == "1" or lower_name == "heading 1":
            set_xml_fonts(rpr, "黑体", "Times New Roman")
            set_xml_size(rpr, 16)
        elif sid == "2" or lower_name == "heading 2":
            set_xml_fonts(rpr, "黑体", "Times New Roman")
            set_xml_size(rpr, 14)
        elif sid == "3" or lower_name == "heading 3":
            set_xml_fonts(rpr, "黑体", "Times New Roman")
            set_xml_size(rpr, 12)
        elif lower_name.startswith("toc"):
            set_xml_fonts(rpr, "宋体", "Times New Roman")
            set_xml_size(rpr, 12)
        elif name == "Hyperlink":
            set_xml_fonts(rpr, "宋体", "Times New Roman")
            underline = rpr.find(qn("w:u"))
            if underline is not None:
                underline.set(qn("w:val"), "none")
        elif "tok" in sid.lower() or "verbatim" in sid.lower() or name in {"Code", "Source Code", "HTML Code", "Listing"}:
            set_xml_fonts(rpr, "Consolas", "Consolas")
            set_xml_size(rpr, 10.5)
        elif lower_name in {"normal", "body text", "caption"} or name in {"Normal", "Body Text", "caption"}:
            set_xml_fonts(rpr, "宋体", "Times New Roman")
        if lower_name == "caption" or name == "caption":
            set_xml_size(rpr, 10.5)
            remove_xml_italic(rpr)
        changed = True
    return changed


def export_pdf() -> None:
    if PDF.exists():
        PDF.unlink()
    subprocess.run(
        [
            r"C:/Program Files/LibreOffice/program/soffice.com",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(OUT.parent),
            str(OUT),
        ],
        check=True,
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
    )


def run_checker() -> None:
    env = dict(os.environ)
    env["PYTHONIOENCODING"] = "utf-8"
    env["PYTHONPATH"] = f"{ROOT}{os.pathsep}{ROOT / 'src'}{os.pathsep}{env.get('PYTHONPATH', '')}"
    subprocess.run(
        [
            sys.executable,
            "-m",
            "thesis_format_checker.cli",
            "check",
            str(OUT),
            "--preset",
            "ncwu",
            "--md",
            str(REPORT),
        ],
        check=True,
        env=env,
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
    )


def render_pdf_pages() -> list[Path]:
    if PAGE_DIR.exists():
        shutil.rmtree(PAGE_DIR)
    PAGE_DIR.mkdir(parents=True, exist_ok=True)
    prefix = PAGE_DIR / "page"
    subprocess.run(["pdftoppm", "-r", "72", "-png", str(PDF), str(prefix)], check=True)
    return sorted(PAGE_DIR.glob("page-*.png"))


def page_number_from_path(path: Path) -> int:
    match = re.search(r"-(\d+)\.png$", path.name)
    return int(match.group(1)) if match else -1


def scan_pdf_blank_space() -> list[dict]:
    from PIL import Image

    suspects: list[dict] = []
    for path in render_pdf_pages():
        page = page_number_from_path(path)
        img = Image.open(path).convert("L")
        width, height = img.size
        pix = img.load()
        y0, y1 = int(height * 0.06), int(height * 0.86)
        x0, x1 = int(width * 0.08), int(width * 0.92)
        ys: list[int] = []
        for y in range(y0, y1):
            row_dark = 0
            for x in range(x0, x1):
                if pix[x, y] < 245:
                    row_dark += 1
            if row_dark >= 4:
                ys.append(y)
        if not ys:
            suspects.append({"page": page, "reason": "blank-or-nearly-blank", "bottom_blank_ratio": 1.0})
            continue
        last_y = max(ys)
        bottom_blank_ratio = (y1 - last_y) / height
        if bottom_blank_ratio >= 0.32:
            suspects.append(
                {
                    "page": page,
                    "last_content_y": last_y,
                    "body_bottom_y": y1,
                    "bottom_blank_ratio": round(bottom_blank_ratio, 3),
                }
            )
    BLANK_REPORT.write_text(json.dumps(suspects, ensure_ascii=False, indent=2), encoding="utf-8")
    return suspects


def audit_visual_format() -> dict:
    counters = {
        "visible_runs": 0,
        "non_black_runs": 0,
        "style_non_black": 0,
    }
    with ZipFile(OUT) as z:
        for filename in [n for n in z.namelist() if n.startswith("word/") and n.endswith(".xml")]:
            root = etree.fromstring(z.read(filename))
            for run in root.xpath(".//w:r", namespaces=NS):
                text = "".join(t.text or "" for t in run.xpath(".//w:t", namespaces=NS)).strip()
                if not text:
                    continue
                counters["visible_runs"] += 1
                color = run.find("w:rPr/w:color", namespaces=NS)
                if color is not None and color.get(f"{{{W_NS}}}val") not in {None, "000000", "auto"}:
                    counters["non_black_runs"] += 1
        styles = etree.fromstring(z.read("word/styles.xml"))
        for color in styles.xpath(".//w:style/w:rPr/w:color", namespaces=NS):
            if color.get(f"{{{W_NS}}}val") not in {None, "000000", "auto"}:
                counters["style_non_black"] += 1
    return counters


def verify_delivery_contract() -> None:
    required = [ORIGINAL, SRC, OUT, PDF, REPORT, BLANK_REPORT]
    missing = [str(path) for path in required if not path.exists()]
    if missing:
        raise RuntimeError(f"missing delivery artifacts: {missing}")
    expected_version_token = f"_v{OUT_VERSION:03d}_"
    if OUT.suffix.lower() != ".docx" or expected_version_token not in OUT.name:
        raise RuntimeError(f"final output is not the expected versioned DOCX: {OUT}")
    if ORIGINAL.name == OUT.name:
        raise RuntimeError("final output overwrote the original filename")
    if SRC_VERSION >= 0 and OUT_VERSION != SRC_VERSION + 1:
        raise RuntimeError(f"output version did not increment: source=v{SRC_VERSION:03d}, output=v{OUT_VERSION:03d}")


def run_size_pt(run) -> float | None:
    if run.font.size is not None:
        return run.font.size.pt
    rpr = run._element.find(qn("w:rPr"))
    if rpr is None:
        return None
    size = rpr.find(qn("w:sz"))
    if size is None:
        return None
    value = size.get(qn("w:val"))
    return int(value) / 2 if value and value.isdigit() else None


def first_visible_run(paragraph):
    for run in paragraph.runs:
        if run.text.strip():
            return run
    return None


def verify_cover_contract() -> None:
    doc = Document(str(OUT))
    if len(doc.paragraphs) < 10 or not doc.tables:
        raise RuntimeError("cover structure is missing")

    school_run = first_visible_run(doc.paragraphs[1])
    title_run = first_visible_run(doc.paragraphs[3])
    topic_run = first_visible_run(doc.paragraphs[4])
    footer_run = first_visible_run(doc.paragraphs[9])
    if doc.paragraphs[1].text.strip() != "华北水利水电大学":
        raise RuntimeError(f"cover school name mismatch: {doc.paragraphs[1].text!r}")
    if doc.paragraphs[3].text.strip() != COVER_TITLE:
        raise RuntimeError(f"cover title mismatch: {doc.paragraphs[3].text!r}")
    cover_text = "\n".join(p.text for p in doc.paragraphs[:10])
    if "文献综述" in cover_text or "毕 业 设 计" in cover_text:
        raise RuntimeError(f"cover contains stale title text: {cover_text!r}")
    if school_run is None or (run_size_pt(school_run) or 0) < 40 or school_run.bold is not True:
        raise RuntimeError("cover school name is not restored to template-scale bold type")
    if title_run is None or (run_size_pt(title_run) or 0) < 28 or title_run.bold is not True:
        raise RuntimeError("cover title is not restored to template-scale bold type")
    if topic_run is None or (run_size_pt(topic_run) or 0) < 17 or topic_run.bold is not True:
        raise RuntimeError("cover topic line is not restored to template-scale bold type")
    if footer_run is None or (run_size_pt(footer_run) or 0) < 13:
        raise RuntimeError("cover footer is not restored to template size")

    table = doc.tables[0]
    for row_index, (label, value) in enumerate(COVER_TABLE_ROWS):
        row = table.rows[row_index]
        actual_label = row.cells[0].text.strip()
        actual_value = row.cells[1].text.strip()
        if actual_label != label or actual_value != value:
            raise RuntimeError(f"cover table row {row_index} mismatch: {(actual_label, actual_value)}")
        for cell in row.cells[:2]:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if not run.text.strip():
                        continue
                    if (run_size_pt(run) or 0) < 13 or run.bold is not True:
                        raise RuntimeError(f"cover table style mismatch at row {row_index}: {run.text!r}")


def verify_content_integrity() -> None:
    src_dir = ROOT / "src"
    if str(ROOT) not in sys.path:
        sys.path.insert(0, str(ROOT))
    if str(src_dir) not in sys.path:
        sys.path.insert(0, str(src_dir))
    from thesis_format_checker.checker import check, load_preset

    preset = load_preset("ncwu")
    _docx, content, findings = check(OUT, preset)
    if findings:
        detail = "; ".join(f"{finding.rule_id}: {finding.message}" for finding in findings[:5])
        raise RuntimeError(f"checker findings={len(findings)} {detail}")
    if content.abstract_zh_chars < 500:
        raise RuntimeError(f"zh abstract too short: {content.abstract_zh_chars}")
    if content.abstract_en_words < 300:
        raise RuntimeError(f"en abstract too short: {content.abstract_en_words}")
    if content.foreign_translation_chars < 2000:
        raise RuntimeError(f"foreign translation too short: {content.foreign_translation_chars}")
    chapter_text = "\n".join(content.chapters)
    missing_chapters = [f"第{i}章" for i in range(1, 8) if f"第{i}章" not in chapter_text]
    if missing_chapters:
        raise RuntimeError(f"missing chapters: {missing_chapters}")
    missing_markers = [
        marker for marker in ["摘 要", "ABSTRACT", "参考文献", "附录一", "附录二", "附录三"]
        if marker not in content.full_text
    ]
    if missing_markers:
        raise RuntimeError(f"missing structural markers: {missing_markers}")


def verify_headers_exact() -> None:
    doc = Document(str(OUT))
    headers = []
    for section in doc.sections:
        text = " ".join(paragraph.text.strip() for paragraph in section.header.paragraphs if paragraph.text.strip()).strip()
        if text:
            headers.append(text)
    if not headers:
        raise RuntimeError("no body headers found")
    bad_headers = [header for header in headers if header != EXPECTED_HEADER]
    if bad_headers:
        raise RuntimeError(f"unexpected headers: {bad_headers}")
    if any(("(" in header or ")" in header or "（" in header or "）" in header or "论文" in header) for header in headers):
        raise RuntimeError(f"header still contains parentheses or thesis suffix: {headers}")


def verify_font_contract() -> None:
    src_dir = ROOT / "src"
    if str(ROOT) not in sys.path:
        sys.path.insert(0, str(ROOT))
    if str(src_dir) not in sys.path:
        sys.path.insert(0, str(src_dir))
    from thesis_format_checker.docx_inspector import inspect

    docx = inspect(OUT)
    for style_name, expected_ea, expected_latin in [
        ("Normal", "宋体", "Times New Roman"),
        ("Body Text", "宋体", "Times New Roman"),
    ]:
        style = docx.styles.get(style_name)
        if style is None:
            raise RuntimeError(f"missing style: {style_name}")
        if style.east_asia and style.east_asia != expected_ea:
            raise RuntimeError(f"{style_name} eastAsia={style.east_asia}, expected {expected_ea}")
        if style.latin and style.latin != expected_latin:
            raise RuntimeError(f"{style_name} latin={style.latin}, expected {expected_latin}")

    bad_fonts = []
    for item in iter_direct_run_fonts(OUT):
        ea = item["east_asia"]
        ascii_font = item["ascii"]
        hansi = item["hansi"]
        if ea and ea not in ALLOWED_EAST_ASIA_FONTS:
            bad_fonts.append(item)
        if ascii_font and ascii_font not in ALLOWED_LATIN_FONTS:
            bad_fonts.append(item)
        if hansi and hansi not in ALLOWED_LATIN_FONTS:
            bad_fonts.append(item)
        if len(bad_fonts) >= 5:
            break
    if bad_fonts:
        raise RuntimeError(f"unexpected direct fonts: {bad_fonts}")


def verify_forbidden_terms() -> None:
    text = document_text(OUT)
    hits = {term: text.count(term) for term in FORBIDDEN_TERMS if text.count(term)}
    if hits:
        raise RuntimeError(f"forbidden terms present: {hits}")


def verify_bridge_paragraphs_once(path: Path | None = None) -> None:
    target = path or OUT
    doc = Document(str(target))
    paragraphs = [normalize_text(paragraph.text) for paragraph in doc.paragraphs if paragraph.text.strip()]
    missing: list[str] = []
    repeated: dict[str, int] = {}
    for text in (item for texts in BRIDGE_PARAGRAPHS.values() for item in texts):
        normalized = normalize_text(text)
        count = sum(1 for paragraph in paragraphs if paragraph == normalized)
        if count == 0:
            missing.append(text[:40])
        elif count > 1:
            repeated[text[:40]] = count
    if missing or repeated:
        raise RuntimeError(f"bridge paragraph contract failed: missing={missing}, repeated={repeated}")


def verify_image_table_count() -> None:
    before = Document(str(SRC))
    after = Document(str(OUT))
    if len(after.inline_shapes) < len(before.inline_shapes):
        raise RuntimeError(f"inline image count decreased: source={len(before.inline_shapes)} output={len(after.inline_shapes)}")
    if len(after.tables) < len(before.tables):
        raise RuntimeError(f"table count decreased: source={len(before.tables)} output={len(after.tables)}")


def verify_blank_scan_contract(blank_suspects: list[dict]) -> None:
    pages = {item.get("page") for item in blank_suspects}
    unexpected = pages - ALLOWED_BLANK_PAGES
    if unexpected:
        raise RuntimeError(f"unexpected blank-scan pages: {sorted(unexpected)}")
    if 90 in pages:
        raise RuntimeError("reference orphan tail page returned at page 90")
    if len(blank_suspects) > len(ALLOWED_BLANK_PAGES):
        raise RuntimeError(f"blank suspects increased: {len(blank_suspects)}")


def verify_visual_audit(audit: dict) -> None:
    if audit["non_black_runs"] or audit["style_non_black"]:
        raise RuntimeError(f"visual color audit failed: {audit}")


def update_version_log(blank_suspects: list[dict], audit: dict) -> None:
    source_label = f"v{SRC_VERSION:03d} - {SRC_LABEL}" if SRC_VERSION >= 0 else SRC_LABEL
    version_title = f"v{OUT_VERSION:03d} - {OUTPUT_LABEL}"
    entry = f"""

## {version_title}

- 文件: `{OUT}`
- PDF: `{PDF}`
- 检测报告: `{REPORT}`
- 留白扫描: `{BLANK_REPORT}`，可疑页 {len(blank_suspects)} 页
- 处理内容:
  - 基于 {source_label} 继续迭代，不覆盖原始文件和上一版交付件。
  - 恢复封面首页为学校文献综述模板的版式尺度。
  - 封面主标题固定为“毕业设计（论文）”，不保留“文献综述”，也不退回“毕 业 设 计”。
  - 将封面作为独立 front-matter 契约处理，避免被正文 12pt 全篇统一逻辑覆盖。
  - 保持上一版的阅读节奏补充段、页眉、字体、颜色、图片、表格和参考文献紧凑排版规则。
- 颜色审计: visible_runs={audit['visible_runs']}, non_black_runs={audit['non_black_runs']}, style_non_black={audit['style_non_black']}
"""
    if VERSION_LOG.exists():
        text = VERSION_LOG.read_text(encoding="utf-8")
    else:
        text = "# 202213210刘高朋修改迭代版 版本记录\n"
    marker = f"## {version_title}"
    start = text.find(marker)
    if start == -1:
        updated = text.rstrip() + entry + "\n"
    else:
        next_match = re.search(r"\n## v\d+", text[start + len(marker):])
        if next_match:
            end = start + len(marker) + next_match.start()
            updated = text[:start].rstrip() + entry + text[end:]
        else:
            updated = text[:start].rstrip() + entry + "\n"
    VERSION_LOG.write_text(updated, encoding="utf-8")


def main() -> None:
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    if hasattr(sys.stderr, "reconfigure"):
        sys.stderr.reconfigure(encoding="utf-8", errors="replace")
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    normalize_docx_with_python_docx()
    patch_xml_colors_and_styles()
    export_pdf()
    run_checker()
    verify_content_integrity()
    verify_headers_exact()
    verify_cover_contract()
    verify_font_contract()
    verify_forbidden_terms()
    verify_bridge_paragraphs_once()
    verify_image_table_count()
    verify_engineering_figure_assets()
    blank_suspects = scan_pdf_blank_space()
    audit = audit_visual_format()
    verify_delivery_contract()
    verify_blank_scan_contract(blank_suspects)
    verify_visual_audit(audit)
    update_version_log(blank_suspects, audit)
    print(f"OUT={OUT}")
    print(f"PDF={PDF}")
    print(f"REPORT={REPORT}")
    print(f"BLANK_REPORT={BLANK_REPORT}")
    print(f"blank_suspects={len(blank_suspects)} {blank_suspects}")
    print(f"audit={audit}")


if __name__ == "__main__":
    main()
