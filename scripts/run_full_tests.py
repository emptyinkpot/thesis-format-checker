"""Single full-test entrypoint for thesis-format-checker.

Run from anywhere:

    python E:/My Project/thesis-format-checker/scripts/run_full_tests.py

What it covers:
- Python syntax/bytecode compilation
- built-in unit contracts for rules and preset loading
- v014 DOCX regeneration through the canonical delivery builder
- NCWU checker pass on v014
- first-page cover contract against the school literature-review template
- color-consistency regression check against v011
- v014 visual audit and blank-scan sanity checks
"""

from __future__ import annotations

import contextlib
import io
import json
import os
import subprocess
import sys
from dataclasses import dataclass
from pathlib import Path
from zipfile import ZipFile
from xml.etree import ElementTree as ET

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
SRC_DIR = ROOT / "src"
DOWNLOADS = Path(r"C:/Users/ASUS-KL/Downloads")
ORIGINAL = DOWNLOADS / "202213210刘高朋修改迭代版.docx"
V011 = DOWNLOADS / "202213210刘高朋修改迭代版_v011_格式统一交付版.docx"
V012 = DOWNLOADS / "202213210刘高朋修改迭代版_v012_全篇黑色字体统一版.docx"
V013 = DOWNLOADS / "202213210刘高朋修改迭代版_v013_阅读节奏优化版.docx"
V014 = DOWNLOADS / "202213210刘高朋修改迭代版_v014_封面模板修正版.docx"
V014_PDF = DOWNLOADS / "202213210刘高朋修改迭代版_v014_封面模板修正版.pdf"
V014_REPORT = DOWNLOADS / "202213210刘高朋修改迭代版_v014_格式检测报告.md"
V014_BLANK_REPORT = DOWNLOADS / "202213210刘高朋修改迭代版_v014_留白扫描.json"
VERSION_LOG = DOWNLOADS / "202213210刘高朋修改迭代版_版本记录.md"
EXPECTED_HEADER = "华北水利水电大学毕业设计"
COVER_TEMPLATE = DOWNLOADS / "202213210刘高朋_文献综述_标准模板版.docx"
COVER_TITLE = "毕业设计（论文）"
W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

FORBIDDEN_TERMS = [
    "代码截图",
    "本实验",
    "实验",
    "本研究",
    "本论文",
    "软件算法",
    "算法",
    "预测",
    "补偿",
    "图3.6",
    "7.2 不足与展望",
    "2.2.1 系统需求分析",
    "文献启发",
    "本文按照",
]

ALLOWED_BLANK_PAGES = {2, 3, 23, 52, 53, 69, 76}
ALLOWED_EAST_ASIA_FONTS = {"宋体", "黑体", "楷体", "仿宋_GB2312", "隶书", "Consolas"}
ALLOWED_LATIN_FONTS = {"Times New Roman", "Consolas", "宋体"}

COVER_TABLE_ROWS = [
    ("学    院", "电子工程学院"),
    ("专    业", "电子信息工程"),
    ("姓    名", "刘高朋"),
    ("学    号", "202213210"),
    ("指导教师", "张晓华"),
    ("完成时间", "2026年6月"),
]


@dataclass
class StepResult:
    name: str
    status: str
    detail: str = ""


def run_command(args: list[str], *, cwd: Path = ROOT) -> None:
    subprocess.run(args, cwd=str(cwd), check=True)


def document_text(path: Path) -> str:
    doc = Document(str(path))
    parts: list[str] = [paragraph.text for paragraph in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(parts)


def iter_direct_run_fonts(path: Path):
    with ZipFile(path) as archive:
        for name in archive.namelist():
            if not name.startswith("word/") or not name.endswith(".xml"):
                continue
            root = ET.fromstring(archive.read(name))
            for run in root.findall(f".//{W_NS}r"):
                text = "".join(node.text or "" for node in run.findall(f".//{W_NS}t")).strip()
                if not text:
                    continue
                rpr = run.find(f"{W_NS}rPr")
                if rpr is None:
                    continue
                fonts = rpr.find(f"{W_NS}rFonts")
                if fonts is None:
                    continue
                yield {
                    "source": name,
                    "text": text[:80],
                    "east_asia": fonts.get(f"{W_NS}eastAsia"),
                    "ascii": fonts.get(f"{W_NS}ascii"),
                    "hansi": fonts.get(f"{W_NS}hAnsi"),
                }


def run_size_pt(run) -> float | None:
    if run.font.size is not None:
        return run.font.size.pt
    rpr = run._element.find(qn("w:rPr"))
    if rpr is None:
        return None
    size = rpr.find(qn("w:sz"))
    if size is None:
        return None
    value = size.get(qn("w:val"))
    return int(value) / 2 if value and value.isdigit() else None


def first_visible_run(paragraph):
    for run in paragraph.runs:
        if run.text.strip():
            return run
    return None


def step_compileall() -> str:
    run_command([
        sys.executable,
        "-m",
        "compileall",
        str(ROOT / "src" / "thesis_format_checker"),
        str(ROOT / "delivery" / "build_lgp_docx.py"),
        str(ROOT / "scripts" / "run_full_tests.py"),
    ])
    return "compileall passed"


def step_unit_contracts() -> str:
    if str(SRC_DIR) not in sys.path:
        sys.path.insert(0, str(SRC_DIR))
    from thesis_format_checker.rules import RULES
    from thesis_format_checker.checker import load_preset

    expected_rules = {
        "page-margins", "header-text-match", "header-on-all-sections",
        "body-font-size", "body-east-asia-font", "text-color-consistency",
        "body-line-spacing", "heading1-style", "heading2-style",
        "heading3-style", "heading-style-applied", "chapter-page-break",
        "abstract-zh-length", "abstract-en-length", "foreign-translation-length",
        "toc-present", "cover-fields",
    }
    missing = expected_rules - set(RULES.keys())
    if missing:
        raise RuntimeError(f"missing rule registrations: {sorted(missing)}")

    preset = load_preset("ncwu")
    if preset["preset_id"] != "ncwu":
        raise RuntimeError(f"unexpected preset_id: {preset['preset_id']}")
    if preset["page"]["margin_cm"]["left"] != 3.0:
        raise RuntimeError("unexpected left margin in preset")
    if preset["styles"]["body"]["size_pt"] != 12:
        raise RuntimeError("unexpected body font size in preset")
    if preset["content"]["abstract_zh"]["min_chars"] != 500:
        raise RuntimeError("unexpected zh abstract threshold in preset")

    return "rule registry and preset contracts passed"


def step_regenerate_v014() -> str:
    if str(ROOT) not in sys.path:
        sys.path.insert(0, str(ROOT))
    if str(SRC_DIR) not in sys.path:
        sys.path.insert(0, str(SRC_DIR))
    from delivery import build_lgp_docx

    output = io.StringIO()
    try:
        with contextlib.redirect_stdout(output), contextlib.redirect_stderr(output):
            build_lgp_docx.main()
    except Exception:
        print(output.getvalue())
        raise
    if not V014.exists():
        raise RuntimeError(f"v014 DOCX missing after generation: {V014}")
    return "generated v014 DOCX/PDF/report/blank-scan"


def step_delivery_contract() -> str:
    required = [ORIGINAL, V013, V014, V014_PDF, V014_REPORT, V014_BLANK_REPORT, VERSION_LOG]
    missing = [str(path) for path in required if not path.exists()]
    if missing:
        raise RuntimeError(f"missing delivery artifacts: {missing}")
    if V014.suffix.lower() != ".docx":
        raise RuntimeError(f"final artifact is not a DOCX: {V014}")
    if "_v014_" not in V014.name:
        raise RuntimeError(f"final DOCX is not versioned as v014: {V014.name}")
    if ORIGINAL.name == V014.name:
        raise RuntimeError("final DOCX overwrote the original filename")

    log_text = VERSION_LOG.read_text(encoding="utf-8")
    if V014.name not in log_text or "v014 - 封面模板修正版" not in log_text:
        raise RuntimeError("version log is missing current v014 delivery facts")
    return "v014 DOCX/PDF/report/version-log artifacts present"


def step_cover_contract() -> str:
    if not COVER_TEMPLATE.exists():
        raise RuntimeError(f"cover reference template missing: {COVER_TEMPLATE}")
    doc = Document(str(V014))
    if len(doc.paragraphs) < 10 or not doc.tables:
        raise RuntimeError("cover structure is missing")

    school = doc.paragraphs[1]
    title = doc.paragraphs[3]
    topic = doc.paragraphs[4]
    footer = doc.paragraphs[9]
    if school.text.strip() != "华北水利水电大学":
        raise RuntimeError(f"cover school name mismatch: {school.text!r}")
    if title.text.strip() != COVER_TITLE:
        raise RuntimeError(f"cover title mismatch: {title.text!r}")
    cover_text = "\n".join(paragraph.text for paragraph in doc.paragraphs[:10])
    if "文献综述" in cover_text or "毕 业 设 计" in cover_text:
        raise RuntimeError(f"cover contains stale title text: {cover_text!r}")

    school_run = first_visible_run(school)
    title_run = first_visible_run(title)
    topic_run = first_visible_run(topic)
    footer_run = first_visible_run(footer)
    if school_run is None or (run_size_pt(school_run) or 0) < 40 or school_run.bold is not True:
        raise RuntimeError("cover school name is not template-scale bold type")
    if title_run is None or (run_size_pt(title_run) or 0) < 28 or title_run.bold is not True:
        raise RuntimeError("cover title is not template-scale bold type")
    if topic_run is None or (run_size_pt(topic_run) or 0) < 17 or topic_run.bold is not True:
        raise RuntimeError("cover topic line is not template-scale bold type")
    if footer_run is None or (run_size_pt(footer_run) or 0) < 13:
        raise RuntimeError("cover footer is not template size")

    table = doc.tables[0]
    for row_index, (label, value) in enumerate(COVER_TABLE_ROWS):
        row = table.rows[row_index]
        actual_label = row.cells[0].text.strip()
        actual_value = row.cells[1].text.strip()
        if actual_label != label or actual_value != value:
            raise RuntimeError(f"cover table row {row_index} mismatch: {(actual_label, actual_value)}")
        for cell in row.cells[:2]:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if not run.text.strip():
                        continue
                    if (run_size_pt(run) or 0) < 13 or run.bold is not True:
                        raise RuntimeError(f"cover table style mismatch at row {row_index}: {run.text!r}")
    return "cover matches historical template contract with main-thesis title"


def step_check_v014() -> str:
    if str(SRC_DIR) not in sys.path:
        sys.path.insert(0, str(SRC_DIR))
    from thesis_format_checker.checker import check, load_preset

    preset = load_preset("ncwu")
    _docx, _content, findings = check(V014, preset)
    if findings:
        detail = "; ".join(f"{f.rule_id}: {f.message}" for f in findings[:5])
        raise RuntimeError(f"v014 checker findings={len(findings)} {detail}")
    return "v014 checker findings=0"


def step_content_integrity_contract() -> str:
    if str(SRC_DIR) not in sys.path:
        sys.path.insert(0, str(SRC_DIR))
    from thesis_format_checker.checker import check, load_preset

    preset = load_preset("ncwu")
    _docx, content, _findings = check(V014, preset)
    if content.abstract_zh_chars < 500:
        raise RuntimeError(f"zh abstract too short: {content.abstract_zh_chars}")
    if content.abstract_en_words < 300:
        raise RuntimeError(f"en abstract too short: {content.abstract_en_words}")
    if content.foreign_translation_chars < 2000:
        raise RuntimeError(f"foreign translation too short: {content.foreign_translation_chars}")

    expected_chapters = {f"第{i}章" for i in range(1, 8)}
    chapter_text = "\n".join(content.chapters)
    missing_chapters = [chapter for chapter in sorted(expected_chapters) if chapter not in chapter_text]
    if missing_chapters:
        raise RuntimeError(f"missing chapters: {missing_chapters}")

    required_markers = ["摘 要", "ABSTRACT", "参考文献", "附录一", "附录二", "附录三"]
    missing_markers = [marker for marker in required_markers if marker not in content.full_text]
    if missing_markers:
        raise RuntimeError(f"missing structural markers: {missing_markers}")
    return "abstracts/translation/chapters/appendices present"


def step_header_contract() -> str:
    doc = Document(str(V014))
    headers = []
    for section in doc.sections:
        text = " ".join(paragraph.text.strip() for paragraph in section.header.paragraphs if paragraph.text.strip()).strip()
        if text:
            headers.append(text)
    if not headers:
        raise RuntimeError("no body headers found")
    bad_headers = [header for header in headers if header != EXPECTED_HEADER]
    if bad_headers:
        raise RuntimeError(f"unexpected headers: {bad_headers}")
    if any(("(" in header or ")" in header or "（" in header or "）" in header or "论文" in header) for header in headers):
        raise RuntimeError(f"header still contains parentheses or thesis suffix: {headers}")
    return f"headers exact: {EXPECTED_HEADER}"


def step_font_contract() -> str:
    if str(SRC_DIR) not in sys.path:
        sys.path.insert(0, str(SRC_DIR))
    from thesis_format_checker.docx_inspector import inspect

    docx = inspect(V014)
    style_expectations = {
        "Normal": ("宋体", "Times New Roman"),
        "Body Text": ("宋体", "Times New Roman"),
    }
    for style_name, (expected_ea, expected_latin) in style_expectations.items():
        style = docx.styles.get(style_name)
        if style is None:
            raise RuntimeError(f"missing style: {style_name}")
        if style.east_asia and style.east_asia != expected_ea:
            raise RuntimeError(f"{style_name} eastAsia={style.east_asia}, expected {expected_ea}")
        if style.latin and style.latin != expected_latin:
            raise RuntimeError(f"{style_name} latin={style.latin}, expected {expected_latin}")

    bad_fonts = []
    for item in iter_direct_run_fonts(V014):
        ea = item["east_asia"]
        ascii_font = item["ascii"]
        hansi = item["hansi"]
        if ea and ea not in ALLOWED_EAST_ASIA_FONTS:
            bad_fonts.append(item)
        if ascii_font and ascii_font not in ALLOWED_LATIN_FONTS:
            bad_fonts.append(item)
        if hansi and hansi not in ALLOWED_LATIN_FONTS:
            bad_fonts.append(item)
        if len(bad_fonts) >= 5:
            break
    if bad_fonts:
        raise RuntimeError(f"unexpected direct fonts: {bad_fonts}")
    return "style and direct font families are within allowed thesis set"


def step_forbidden_terms_contract() -> str:
    text = document_text(V014)
    hits = {term: text.count(term) for term in FORBIDDEN_TERMS if text.count(term)}
    if hits:
        raise RuntimeError(f"forbidden terms present: {hits}")
    return "forbidden prose terms absent"


def step_image_table_contract() -> str:
    before = Document(str(V013))
    after = Document(str(V014))
    if len(after.inline_shapes) < len(before.inline_shapes):
        raise RuntimeError(f"inline image count decreased: v013={len(before.inline_shapes)} v014={len(after.inline_shapes)}")
    if len(after.tables) < len(before.tables):
        raise RuntimeError(f"table count decreased: v013={len(before.tables)} v014={len(after.tables)}")
    return f"images/tables preserved: images={len(after.inline_shapes)}, tables={len(after.tables)}"


def step_regression_v011_color_rule() -> str:
    if str(SRC_DIR) not in sys.path:
        sys.path.insert(0, str(SRC_DIR))
    from thesis_format_checker.checker import check, load_preset

    preset = load_preset("ncwu")
    _docx, _content, findings = check(V011, preset)
    color_findings = [f for f in findings if f.rule_id == "text-color-consistency"]
    if not color_findings:
        raise RuntimeError("v011 no longer triggers text-color-consistency regression warning")
    return "v011 triggers text-color-consistency as expected"


def step_visual_audit() -> str:
    if str(ROOT) not in sys.path:
        sys.path.insert(0, str(ROOT))
    from delivery import build_lgp_docx

    audit = build_lgp_docx.audit_visual_format()
    if audit["non_black_runs"] or audit["style_non_black"]:
        raise RuntimeError(f"v014 color audit failed: {audit}")
    return f"color audit passed: {audit}"


def step_blank_scan_sanity() -> str:
    if not V014_BLANK_REPORT.exists():
        raise RuntimeError(f"blank scan report missing: {V014_BLANK_REPORT}")
    suspects = json.loads(V014_BLANK_REPORT.read_text(encoding="utf-8"))
    pages = {item.get("page") for item in suspects}
    unexpected = pages - ALLOWED_BLANK_PAGES
    if unexpected:
        raise RuntimeError(f"unexpected blank-scan pages: {sorted(unexpected)}")
    if 90 in pages:
        raise RuntimeError("reference orphan tail page returned at page 90")
    if len(suspects) > 7:
        raise RuntimeError(f"blank suspects increased: {len(suspects)}")
    return f"blank scan sanity passed: suspects={len(suspects)} pages={sorted(pages)}"


def run_step(name: str, fn) -> StepResult:
    print(f"\n== {name} ==")
    try:
        detail = fn()
        print(f"PASS {detail}")
        return StepResult(name, "PASS", detail)
    except Exception as exc:
        print(f"FAIL {exc}")
        return StepResult(name, "FAIL", str(exc))


def main() -> int:
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    if hasattr(sys.stderr, "reconfigure"):
        sys.stderr.reconfigure(encoding="utf-8", errors="replace")
    os.environ["PYTHONIOENCODING"] = "utf-8"
    os.chdir(ROOT)

    steps = [
        ("compileall", step_compileall),
        ("unit-contracts", step_unit_contracts),
        ("regenerate-v014", step_regenerate_v014),
        ("delivery-contract", step_delivery_contract),
        ("cover-contract", step_cover_contract),
        ("check-v014", step_check_v014),
        ("content-integrity", step_content_integrity_contract),
        ("header-contract", step_header_contract),
        ("font-contract", step_font_contract),
        ("forbidden-terms", step_forbidden_terms_contract),
        ("image-table-contract", step_image_table_contract),
        ("regression-v011-color-rule", step_regression_v011_color_rule),
        ("visual-audit", step_visual_audit),
        ("blank-scan-sanity", step_blank_scan_sanity),
    ]
    results = [run_step(name, fn) for name, fn in steps]

    print("\n== summary ==")
    for result in results:
        print(f"{result.status:4} {result.name} {result.detail}")

    failed = [result for result in results if result.status == "FAIL"]
    return 1 if failed else 0


if __name__ == "__main__":
    raise SystemExit(main())
